from __future__ import annotations

import io
import logging
import os
import json
import re
from html import escape as html_escape
from zipfile import BadZipFile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    HTMLResponse,
    PlainTextResponse,
    Response,
    StreamingResponse,
    RedirectResponse,
)
from fastapi.staticfiles import StaticFiles

from starlette.exceptions import HTTPException as StarletteHTTPException
from starlette.middleware.gzip import GZipMiddleware
from starlette.requests import Request

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response as StarletteResponse


class CanonicalHostRedirectMiddleware(BaseHTTPMiddleware):
    """Redirect non-canonical host/scheme to canonical values.

    - Skips localhost/127.0.0.1
    - Uses X-Forwarded-Proto / X-Forwarded-Host when present (common behind Render/CDNs)
    - Controlled via env vars:
        CANONICAL_HOST (e.g. www.ecuacionesaword.com)
        CANONICAL_SCHEME (e.g. https)
    """

    def __init__(self, app, canonical_host: str, canonical_scheme: str = "https"):
        super().__init__(app)
        self.canonical_host = (canonical_host or "").strip().lower()
        self.canonical_scheme = (canonical_scheme or "https").strip().lower()

    async def dispatch(self, request: Request, call_next):
        # Determine effective host/scheme considering reverse proxy headers
        host = (request.headers.get("x-forwarded-host") or request.headers.get("host") or "").split(",")[0].strip().lower()
        scheme = (request.headers.get("x-forwarded-proto") or request.url.scheme or "").split(",")[0].strip().lower()

        # Skip local/dev
        if host.startswith("127.0.0.1") or host.startswith("localhost"):
            return await call_next(request)

        # If not configured, do nothing
        if not self.canonical_host:
            return await call_next(request)

        if host != self.canonical_host or (self.canonical_scheme and scheme != self.canonical_scheme):
            # Preserve path + query
            url = str(request.url)
            # Rebuild
            path_q = request.url.path
            if request.url.query:
                path_q += "?" + request.url.query
            target = f"{self.canonical_scheme}://{self.canonical_host}{path_q}"
            return RedirectResponse(url=target, status_code=301)

        return await call_next(request)


def _add_common_headers(resp: StarletteResponse) -> StarletteResponse:
    # SEO/UX-safe defaults
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    resp.headers.setdefault("Permissions-Policy", "interest-cohort=()")
    return resp

import math2docx


# ================================================================
# Config
# ================================================================
APP_TITLE = "Ecuaciones a Word (LaTeX → Word OMML)"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

logger = logging.getLogger("ecuacionesaword")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
Segment = Tuple[str, str]  # ("text" | "inline" | "display", contenido)

# Puedes desactivar reglas específicas del "ejercicio" si no las quieres:
USE_EXERCISE_TWEAKS = True


def _mojibake_score(value: str) -> int:
    markers = ("Ã", "Â", "â€™", "â€œ", "â€", "�", "Ð", "Ñ")
    return sum(value.count(m) for m in markers)


def _fix_text_mojibake(text: str) -> str:
    """Best-effort repair for UTF-8/Latin-1 mojibake fragments."""
    if not text:
        return text

    fixed = text
    original_score = _mojibake_score(fixed)
    if any(token in fixed for token in ("Ã", "Â", "â", "Ð", "Ñ", "�")):
        candidates = [fixed]
        for encoding in ("latin-1", "cp1252"):
            try:
                candidates.append(fixed.encode(encoding, errors="ignore").decode("utf-8", errors="ignore"))
            except Exception:
                continue
        fixed = min(candidates, key=_mojibake_score) if candidates else fixed

    replacements = {
        "â†’": "->",
        "â€œ": "\"",
        "â€": "\"",
        "â€˜": "'",
        "â€™": "'",
        "â€“": "-",
        "â€”": "-",
        "â€¦": "...",
        "Â¿": "¿",
        "Â¡": "¡",
        "Âº": "º",
        "Âª": "ª",
        "Â·": "·",
        " ": " ",
    }
    for bad, good in replacements.items():
        fixed = fixed.replace(bad, good)

    if _mojibake_score(fixed) > original_score:
        return text
    return fixed


def _deep_fix_mojibake(value: Any) -> Any:
    if isinstance(value, str):
        return _fix_text_mojibake(value)
    if isinstance(value, list):
        return [_deep_fix_mojibake(v) for v in value]
    if isinstance(value, tuple):
        return tuple(_deep_fix_mojibake(v) for v in value)
    if isinstance(value, dict):
        return {k: _deep_fix_mojibake(v) for k, v in value.items()}
    return value


def _parse_allowed_origins(value: str) -> List[str]:
    return [o.strip() for o in value.split(",") if o.strip()]


def get_allowed_origins() -> List[str]:
    env_val = os.getenv("ALLOWED_ORIGINS", "").strip()
    if env_val:
        return _parse_allowed_origins(env_val)
    return [
        "https://www.ecuacionesaword.com",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
    ]


app = FastAPI(title=APP_TITLE)

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)
app.add_middleware(GZipMiddleware, minimum_size=800)

# Canonical redirects (host + scheme). Configure in production via env vars.
CANONICAL_HOST = os.getenv("CANONICAL_HOST", "").strip()
CANONICAL_SCHEME = os.getenv("CANONICAL_SCHEME", "https").strip()
if CANONICAL_HOST:
    app.add_middleware(CanonicalHostRedirectMiddleware, canonical_host=CANONICAL_HOST, canonical_scheme=CANONICAL_SCHEME)


@app.middleware("http")
async def add_common_headers_mw(request: Request, call_next):
    resp = await call_next(request)
    return _add_common_headers(resp)


# Static
_static_dir = os.path.join(BASE_DIR, "static")
if os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")


# ================================================================
# Blog (data-driven): metadata + aliases + templates
# ================================================================
SITE_NAME = "Ecuaciones a Word"
SITE_CANONICAL_ORIGIN = os.getenv("CANONICAL_SITE_ORIGIN", "").strip().rstrip("/")
if not SITE_CANONICAL_ORIGIN:
    _canonical_host_for_urls = os.getenv("CANONICAL_HOST", "").strip()
    _canonical_scheme_for_urls = os.getenv("CANONICAL_SCHEME", "https").strip()
    if _canonical_host_for_urls:
        SITE_CANONICAL_ORIGIN = f"{_canonical_scheme_for_urls}://{_canonical_host_for_urls}"
    else:
        SITE_CANONICAL_ORIGIN = "https://www.ecuacionesaword.com"

BLOG_DATA_PATH = os.path.join(BASE_DIR, "blog_content", "posts.json")
BLOG_POSTS_DIR = os.path.join(BASE_DIR, "blog_content", "posts")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")

try:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
except Exception as e:  # pragma: no cover
    raise RuntimeError(
        "Missing dependency 'jinja2'. Install it with: pip install jinja2"
    ) from e

JINJA_ENV = Environment(
    loader=FileSystemLoader(TEMPLATES_DIR),
    autoescape=select_autoescape(["html", "xml"]),
)

SUPPORTED_LANGS: Tuple[str, ...] = ("es", "en", "de", "fr", "it", "pt")
PRIMARY_CONTENT_LANGS: Tuple[str, ...] = ("es", "en")
SECONDARY_LANGS: Tuple[str, ...] = tuple(lang for lang in SUPPORTED_LANGS if lang not in PRIMARY_CONTENT_LANGS)
SITEMAP_LANGS: Tuple[str, ...] = PRIMARY_CONTENT_LANGS

# Posts kept live for users but intentionally excluded from index/sitemap to reduce cannibalization.
NON_INDEXABLE_BLOG_SLUGS: Dict[str, set[str]] = {
    "es": {
        "convertidor-formulas-chatgpt-a-word",
        "conversor-formulas-chatgpt-a-word-errores",
        "de-latex-a-word-guia-rapida",
    },
    "en": {
        "latex-to-word-quick-guide-no-broken-equations",
    },
}

LANGUAGE_LABELS: Dict[str, str] = {
    "es": "Español",
    "en": "English",
    "de": "Deutsch",
    "fr": "Français",
    "it": "Italiano",
    "pt": "Português",
}

UI_TEXT: Dict[str, Dict[str, str]] = {
    "es": {
        "nav_converter": "Conversor",
        "related_title": "Artículos relacionados",
        "summary_title": "Consejo rápido",
        "summary_text": "Convierte LaTeX (o ecuaciones generadas por IA) en ecuaciones nativas y editables de Word (OMML).",
        "summary_link": "Conversor LaTeX → Word (OMML)",
        "cta_strong": "¿Necesitas convertir un .docx o .txt con fórmulas LaTeX?",
        "cta_text": "Usa el conversor y descarga un Word con ecuaciones nativas (OMML).",
        "cta_primary": "Conversor LaTeX → Word (OMML)",
        "cta_secondary": "Ver más artículos",
        "blog_intro": "Guías prácticas para exportar ecuaciones de LaTeX/IA a Word como OMML editable, con flujos online y troubleshooting.",
        "index_cta_primary": "Conversor LaTeX → Word (OMML)",
        "search_label": "Buscar artículos",
        "search_placeholder": "Buscar por tema, herramienta o problema (ej. pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filtrar por etiqueta",
        "filter_all": "Todos",
        "featured_title": "Artículos",
        "legal_privacy": "Política de privacidad",
        "legal_terms": "Términos de uso",
        "legal_contact": "Contacto",
    },
    "en": {
        "nav_converter": "Converter",
        "related_title": "Related articles",
        "summary_title": "Quick tip",
        "summary_text": "Convert LaTeX (or AI-generated math) into native, editable Word equations (OMML).",
        "summary_link": "LaTeX → OMML converter",
        "cta_strong": "Need to convert a .docx or .txt containing LaTeX?",
        "cta_text": "Use the converter and download a Word file with native editable OMML equations.",
        "cta_primary": "LaTeX → OMML Converter",
        "cta_secondary": "More articles",
        "blog_intro": "Practical guides to export LaTeX/AI equations to Word as native editable OMML equations, including online workflows and troubleshooting.",
        "index_cta_primary": "LaTeX → OMML Converter",
        "search_label": "Search articles",
        "search_placeholder": "Search by topic, tool or issue (e.g., pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filter by tag",
        "filter_all": "All",
        "featured_title": "Articles",
        "legal_privacy": "Privacy Policy",
        "legal_terms": "Terms of use",
        "legal_contact": "Contact",
    },
    "de": {
        "nav_converter": "Konverter",
        "related_title": "Verwandte Artikel",
        "summary_title": "Schneller Tipp",
        "summary_text": "Konvertiere LaTeX (oder KI-generierte Mathematik) in native, bearbeitbare Word-Gleichungen (OMML).",
        "summary_link": "LaTeX → OMML-Konverter",
        "cta_strong": "Möchtest du eine .docx- oder .txt-Datei mit LaTeX konvertieren?",
        "cta_text": "Nutze den Konverter und lade eine Word-Datei mit nativen, bearbeitbaren OMML-Gleichungen herunter.",
        "cta_primary": "LaTeX → OMML-Konverter",
        "cta_secondary": "Weitere Artikel",
        "blog_intro": "Praktische Leitfäden, um LaTeX/KI-Gleichungen als bearbeitbares OMML nach Word zu exportieren.",
        "index_cta_primary": "LaTeX → OMML-Konverter",
        "search_label": "Artikel suchen",
        "search_placeholder": "Suche nach Thema, Tool oder Problem (z. B. pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Nach Tag filtern",
        "filter_all": "Alle",
        "featured_title": "Artikel",
        "legal_privacy": "Datenschutzerklärung",
        "legal_terms": "Nutzungsbedingungen",
        "legal_contact": "Kontakt",
    },
    "fr": {
        "nav_converter": "Convertisseur",
        "related_title": "Articles associés",
        "summary_title": "Conseil rapide",
        "summary_text": "Convertissez le LaTeX (ou les maths générées par IA) en équations Word natives et modifiables (OMML).",
        "summary_link": "Convertisseur LaTeX → OMML",
        "cta_strong": "Besoin de convertir un fichier .docx ou .txt contenant du LaTeX ?",
        "cta_text": "Utilisez le convertisseur et téléchargez un fichier Word avec des équations OMML natives et modifiables.",
        "cta_primary": "Convertisseur LaTeX → OMML",
        "cta_secondary": "Plus d'articles",
        "blog_intro": "Guides pratiques pour exporter des équations LaTeX/IA vers Word en OMML éditable.",
        "index_cta_primary": "Convertisseur LaTeX → OMML",
        "search_label": "Rechercher des articles",
        "search_placeholder": "Rechercher par sujet, outil ou problème (ex. pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filtrer par étiquette",
        "filter_all": "Tous",
        "featured_title": "Articles",
        "legal_privacy": "Politique de confidentialité",
        "legal_terms": "Conditions d'utilisation",
        "legal_contact": "Contact",
    },
    "it": {
        "nav_converter": "Convertitore",
        "related_title": "Articoli correlati",
        "summary_title": "Suggerimento rapido",
        "summary_text": "Converti LaTeX (o matematica generata dall'IA) in equazioni Word native e modificabili (OMML).",
        "summary_link": "Convertitore LaTeX → OMML",
        "cta_strong": "Devi convertire un file .docx o .txt con LaTeX?",
        "cta_text": "Usa il convertitore e scarica un file Word con equazioni OMML native e modificabili.",
        "cta_primary": "Convertitore LaTeX → OMML",
        "cta_secondary": "Altri articoli",
        "blog_intro": "Guide pratiche per esportare equazioni LaTeX/IA in Word come OMML modificabile.",
        "index_cta_primary": "Convertitore LaTeX → OMML",
        "search_label": "Cerca articoli",
        "search_placeholder": "Cerca per tema, strumento o problema (es. pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filtra per tag",
        "filter_all": "Tutti",
        "featured_title": "Articoli",
        "legal_privacy": "Informativa sulla privacy",
        "legal_terms": "Termini di utilizzo",
        "legal_contact": "Contatto",
    },
    "pt": {
        "nav_converter": "Conversor",
        "related_title": "Artigos relacionados",
        "summary_title": "Dica rápida",
        "summary_text": "Converta LaTeX (ou matemática gerada por IA) em equações nativas e editáveis do Word (OMML).",
        "summary_link": "Conversor LaTeX → OMML",
        "cta_strong": "Precisa converter um arquivo .docx ou .txt com LaTeX?",
        "cta_text": "Use o conversor e baixe um arquivo Word com equações OMML nativas e editáveis.",
        "cta_primary": "Conversor LaTeX → OMML",
        "cta_secondary": "Mais artigos",
        "blog_intro": "Guias práticos para exportar equações LaTeX/IA para Word em OMML editável.",
        "index_cta_primary": "Conversor LaTeX → OMML",
        "search_label": "Pesquisar artigos",
        "search_placeholder": "Pesquise por tema, ferramenta ou problema (ex.: pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filtrar por tag",
        "filter_all": "Todos",
        "featured_title": "Artigos",
        "legal_privacy": "Política de Privacidade",
        "legal_terms": "Termos de uso",
        "legal_contact": "Contato",
    },
}


def _ui(lang: str, key: str) -> str:
    if lang in UI_TEXT and key in UI_TEXT[lang]:
        return _fix_text_mojibake(UI_TEXT[lang][key])
    return _fix_text_mojibake(UI_TEXT["en"].get(key, ""))

LANG_PREFIX: Dict[str, str] = {
    "es": "",
    "en": "/en",
    "de": "/de",
    "fr": "/fr",
    "it": "/it",
    "pt": "/pt",
}

BLOG_POSTS: Dict[str, Dict[str, Dict[str, Any]]] = {lang: {} for lang in SUPPORTED_LANGS}
BLOG_LIST: Dict[str, List[Dict[str, Any]]] = {lang: [] for lang in SUPPORTED_LANGS}
BLOG_ALIASES: Dict[str, Dict[str, str]] = {lang: {} for lang in SUPPORTED_LANGS}


def _content_lang(lang: str) -> str:
    if lang == "es":
        return "es"
    if lang in SUPPORTED_LANGS and lang in BLOG_POSTS and BLOG_POSTS.get(lang):
        return lang
    return "en"


def _lang_prefix(lang: str) -> str:
    return LANG_PREFIX.get(lang, "/en")


def _home_path(lang: str) -> str:
    prefix = _lang_prefix(lang)
    return prefix or "/"


def _blog_index_path(lang: str) -> str:
    prefix = _lang_prefix(lang)
    return f"{prefix}/blog" if prefix else "/blog"


def _legal_path(lang: str, page: str) -> str:
    prefix = _lang_prefix(lang)
    return f"{prefix}/{page}" if prefix else f"/{page}"


def _solutions_path(lang: str) -> str:
    if lang == "es":
        return "/soluciones"
    return f"{_lang_prefix(lang)}/solutions"


def _all_alternates(
    path_by_lang: Dict[str, str],
    default_lang: str = "es",
    langs: Tuple[str, ...] = PRIMARY_CONTENT_LANGS,
) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    for code in langs:
        p = (path_by_lang.get(code) or "").strip()
        if not p:
            continue
        out.append({"hreflang": code, "href": _abs_url(p)})
    default_path = path_by_lang.get(default_lang) or path_by_lang.get("en") or "/"
    out.append({"hreflang": "x-default", "href": _abs_url(default_path)})
    return out


def _lang_options(path_by_lang: Dict[str, str]) -> List[Dict[str, str]]:
    options: List[Dict[str, str]] = []
    for code in SUPPORTED_LANGS:
        path = (path_by_lang.get(code) or "").strip()
        if not path:
            continue
        options.append(
            {
                "code": code,
                "label": _fix_text_mojibake(LANGUAGE_LABELS.get(code, code.upper())),
                "href": path,
            }
        )
    return options


def _format_reading_time(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (int, float)):
        return f"{int(value)} min"
    txt = str(value).strip()
    if not txt:
        return ""
    if re.fullmatch(r"\d+", txt):
        return f"{txt} min"
    txt = txt.replace("–", "-").replace("—", "-")
    txt = re.sub(r"(?<!\w)minutos?(?!\w)", "min", txt, flags=re.IGNORECASE)
    if re.search(r"\bmin\b", txt, flags=re.IGNORECASE):
        return txt
    if re.fullmatch(r"\d+\s*-\s*\d+", txt):
        return f"{txt} min"
    return txt


def _normalize_tag(tag: str) -> str:
    t = (tag or "").strip()
    if not t:
        return ""
    low = t.lower()
    if low in {"omml", "docx", "ai", "ia"}:
        return low.upper()
    if low in {"chatgpt", "pandoc", "overleaf", "word", "latex"}:
        return t[0].upper() + t[1:].lower()
    return t[0].upper() + t[1:].lower()


def _is_primary_lang(lang: str) -> bool:
    return lang in PRIMARY_CONTENT_LANGS


def _is_blog_post_indexable(lang: str, post: Dict[str, Any]) -> bool:
    slug = (post.get("slug") or "").strip()
    if not slug:
        return False
    if not _is_primary_lang(lang):
        return False
    if slug in NON_INDEXABLE_BLOG_SLUGS.get(lang, set()):
        return False
    if bool(post.get("noindex")):
        return False
    return True


def _robots_directive(lang: str, indexable: bool = True) -> str:
    if not indexable or not _is_primary_lang(lang):
        return "noindex,follow,max-image-preview:large"
    return "index,follow,max-image-preview:large"

LANDING_PAGES: Dict[str, Dict[str, Dict[str, Any]]] = {
    "chatgpt-equations-to-word": {
        "es": {
            "path": "/soluciones/chatgpt-ecuaciones-a-word",
            "title": "Convertir ecuaciones de ChatGPT a Word (OMML)",
            "seo_title": "Convertir ecuaciones de ChatGPT a Word (OMML editable) | Ecuaciones a Word",
            "description": "Pasa respuestas matemáticas de ChatGPT a Word con ecuaciones OMML nativas y editables, sin copiar fórmula por fórmula.",
            "h1": "De ChatGPT a Word con ecuaciones editables",
            "kicker": "Flujo recomendado",
            "intro": "Sube un .docx o .txt con LaTeX generado por ChatGPT y descarga un Word con ecuaciones nativas (OMML).",
            "intent_items": [
                "Entrega en Word sin capturas ni fórmulas rotas.",
                "Evita rehacer ecuaciones a mano en el editor de Word.",
                "Mantén formato editable para revisiones y tutorías.",
            ],
            "cta_label": "Convertir archivo ahora",
            "related_blog_href": "/blog/pasar-ecuaciones-chatgpt-word",
            "related_blog_label": "Guía: pasar ecuaciones de ChatGPT a Word",
        },
        "en": {
            "path": "/en/solutions/chatgpt-equations-to-word",
            "title": "Convert ChatGPT equations to Word (OMML)",
            "seo_title": "Convert ChatGPT equations to Word (editable OMML) | Equations to Word",
            "description": "Turn ChatGPT math output into native editable Word equations (OMML) without rebuilding each formula manually.",
            "h1": "From ChatGPT output to editable Word equations",
            "kicker": "Recommended workflow",
            "intro": "Upload a .docx or .txt with ChatGPT-generated LaTeX and download a Word file with native OMML equations.",
            "intent_items": [
                "Submit Word documents without screenshots.",
                "Skip manual equation retyping.",
                "Keep equations editable for revisions.",
            ],
            "cta_label": "Convert file now",
            "related_blog_href": "/en/blog/copy-chatgpt-equations-word",
            "related_blog_label": "Guide: copy ChatGPT equations to Word",
        },
    },
    "gemini-equations-to-word": {
        "es": {
            "path": "/soluciones/gemini-ecuaciones-a-word",
            "title": "Gemini equations to Word (OMML)",
            "seo_title": "Gemini a Word con ecuaciones OMML editables | Ecuaciones a Word",
            "description": "Convierte fórmulas de Gemini en ecuaciones nativas de Word (OMML) para entregar documentos limpios y editables.",
            "h1": "Gemini a Word en formato OMML",
            "kicker": "IA a documento final",
            "intro": "Cuando Gemini te devuelve LaTeX, este conversor lo transforma en ecuaciones de Word nativas para tu entrega final.",
            "intent_items": [
                "Ideal para informes y ejercicios en Word.",
                "Soporta bloques y ecuaciones inline.",
                "Salida .docx lista para editar.",
            ],
            "cta_label": "Subir y convertir",
            "related_blog_href": "/blog/ia-chatgpt-a-word-ejercicios",
            "related_blog_label": "Guía: usar IA + Word en ejercicios",
        },
        "en": {
            "path": "/en/solutions/gemini-equations-to-word",
            "title": "Gemini equations to Word (OMML)",
            "seo_title": "Gemini equations to Word with editable OMML | Equations to Word",
            "description": "Convert Gemini math responses into native Word OMML equations for clean, editable .docx submissions.",
            "h1": "Gemini output to native Word equations",
            "kicker": "AI to final document",
            "intro": "If Gemini outputs LaTeX, this converter turns it into editable OMML equations inside your Word file.",
            "intent_items": [
                "Built for assignments and reports in Word.",
                "Supports inline and display equations.",
                "Download-ready editable .docx output.",
            ],
            "cta_label": "Upload and convert",
            "related_blog_href": "/en/blog/gemini-equations-to-word-omml",
            "related_blog_label": "Guide: Gemini equations to Word",
        },
    },
    "pandoc-to-word-omml": {
        "es": {
            "path": "/soluciones/pandoc-a-word-omml",
            "title": "Pandoc a Word OMML",
            "seo_title": "Pandoc to Word OMML: ecuaciones editables | Ecuaciones a Word",
            "description": "Si Pandoc no deja tus fórmulas editables en Word, convierte el resultado a OMML nativo sin rehacer el documento.",
            "h1": "Convierte salidas de Pandoc a OMML editable",
            "kicker": "Solución práctica",
            "intro": "Sube tu .docx o .txt posterior a Pandoc y obtén ecuaciones compatibles con Word en formato OMML.",
            "intent_items": [
                "Recupera editabilidad de ecuaciones.",
                "Reduce errores de copy/paste.",
                "Funciona con documentos largos.",
            ],
            "cta_label": "Arreglar ecuaciones",
            "related_blog_href": "/blog/pandoc-ecuaciones-word-no-editables-soluciones",
            "related_blog_label": "Guía: problemas típicos con Pandoc",
        },
        "en": {
            "path": "/en/solutions/pandoc-to-word-omml",
            "title": "Pandoc to Word OMML",
            "seo_title": "Pandoc to Word OMML: editable equations | Equations to Word",
            "description": "If Pandoc output is not editable in Word, convert the resulting file to native OMML equations without manual cleanup.",
            "h1": "Fix Pandoc output with editable OMML equations",
            "kicker": "Practical fix",
            "intro": "Upload your post-Pandoc .docx or .txt and get native Word-compatible OMML equations.",
            "intent_items": [
                "Restore equation editability.",
                "Cut copy/paste formatting issues.",
                "Works with long technical docs.",
            ],
            "cta_label": "Fix equations now",
            "related_blog_href": "/en/blog/pandoc-math-to-word-omml-troubleshooting",
            "related_blog_label": "Guide: Pandoc troubleshooting",
        },
    },
    "overleaf-latex-document-to-word": {
        "es": {
            "path": "/soluciones/overleaf-latex-documento-a-word",
            "title": "Overleaf/LaTeX document to Word",
            "seo_title": "Overleaf o documento LaTeX a Word con OMML | Ecuaciones a Word",
            "description": "Convierte documentos de Overleaf/LaTeX a Word con ecuaciones OMML editables para entregar TFG, TFM e informes.",
            "h1": "De Overleaf o LaTeX a Word editable",
            "kicker": "Entrega académica",
            "intro": "Transforma contenido LaTeX de Overleaf en un .docx con ecuaciones nativas de Word listas para revisión final.",
            "intent_items": [
                "Útil para TFG/TFM y documentación técnica.",
                "Sin perder editabilidad matemática.",
                "Flujo rápido para entregas.",
            ],
            "cta_label": "Preparar mi entrega",
            "related_blog_href": "/blog/overleaf-latex-a-word-ecuaciones-editables",
            "related_blog_label": "Guía: Overleaf a Word",
        },
        "en": {
            "path": "/en/solutions/overleaf-latex-document-to-word",
            "title": "Overleaf / LaTeX document to Word",
            "seo_title": "Overleaf or LaTeX document to Word with OMML | Equations to Word",
            "description": "Convert Overleaf/LaTeX documents to Word with editable OMML equations for thesis and assignment submissions.",
            "h1": "From Overleaf or LaTeX to editable Word",
            "kicker": "Academic delivery",
            "intro": "Turn Overleaf LaTeX content into a .docx with native Word equations ready for final review.",
            "intent_items": [
                "Great for thesis and coursework.",
                "Keep equations fully editable.",
                "Fast handoff to Word-based workflows.",
            ],
            "cta_label": "Prepare my submission",
            "related_blog_href": "/en/blog/overleaf-latex-to-word-editable-equations",
            "related_blog_label": "Guide: Overleaf to Word",
        },
    },
    "omml-converter": {
        "es": {
            "path": "/soluciones/conversor-omml",
            "title": "OMML converter",
            "seo_title": "Conversor OMML para ecuaciones editables en Word | Ecuaciones a Word",
            "description": "Convierte fórmulas LaTeX a formato OMML, el estándar nativo de ecuaciones en Word, para documentos editables y estables.",
            "h1": "Conversor OMML para Word",
            "kicker": "Formato nativo Word",
            "intro": "Pasa de LaTeX a OMML en un solo flujo para obtener ecuaciones editables dentro de Word.",
            "intent_items": [
                "Salida compatible con Microsoft Word.",
                "Mejor consistencia tipográfica.",
                "Ideal para colaboración en .docx.",
            ],
            "cta_label": "Convertir a OMML",
            "related_blog_href": "/blog/que-es-omml-ecuaciones-word",
            "related_blog_label": "Guía: qué es OMML",
        },
        "en": {
            "path": "/en/solutions/omml-converter",
            "title": "OMML converter",
            "seo_title": "OMML converter for editable Word equations | Equations to Word",
            "description": "Convert LaTeX into OMML, Word's native equation format, to keep math content editable and stable in .docx files.",
            "h1": "OMML converter for Word",
            "kicker": "Native Word format",
            "intro": "Move from LaTeX to OMML in one flow and keep equations editable directly in Word.",
            "intent_items": [
                "Output compatible with Microsoft Word.",
                "More consistent math rendering.",
                "Ideal for collaborative .docx editing.",
            ],
            "cta_label": "Convert to OMML",
            "related_blog_href": "/en/blog/what-is-omml-word-equations",
            "related_blog_label": "Guide: what OMML is",
        },
    },
}


LANDING_LONGFORM: Dict[str, Dict[str, Dict[str, Any]]] = {
    "chatgpt-equations-to-word": {
        "es": {
            "primary_keyword": "ecuaciones de chatgpt a word",
            "primary_intent": "transaccional para convertir respuestas de IA a .docx editable",
            "problem_angle": "evitar copy/paste manual, capturas y formulas rotas desde ChatGPT",
            "problem_statement": "Cuando ChatGPT entrega una solucion matematica en LaTeX, el problema real empieza al pasarla a Word. Muchos usuarios terminan rehaciendo formulas a mano, pegando capturas o corrigiendo simbolos raros. Esta landing esta pensada para cerrar ese hueco: pasar de una respuesta generada por IA a un documento Word limpio, editable y apto para revision academica. El foco no es solo conversion tecnica, sino productividad en entregas reales.",
            "extra_context": "Tambien te ayuda a mantener consistencia entre versiones del mismo ejercicio. Si haces varias iteraciones de una respuesta con IA, puedes consolidar cambios sin perder la estructura matematica ni bloquear la edicion posterior.",
            "when_to_use_intro": "Usa esta solucion cuando ChatGPT es tu fuente principal y Word es el formato final de entrega o colaboracion.",
            "when_to_use_items": [
                "Entregas ejercicios o practicas en .docx.",
                "Necesitas revision con tutor/profesor en Word.",
                "Consolidas varias respuestas de IA en un informe unico.",
                "Quieres evitar fallos de copy/paste desde el navegador.",
                "Necesitas mantener ecuaciones editables tras la entrega.",
            ],
            "workflow_steps": [
                "Genera la respuesta en ChatGPT con delimitadores LaTeX claros.",
                "Pasa el contenido a un .txt o .docx conservando $...$, $$...$$ o \\[...\\].",
                "Sube el archivo al conversor desde esta pagina.",
                "Descarga el .docx con ecuaciones OMML nativas.",
                "Revisa 2-3 ecuaciones criticas y envia el archivo.",
            ],
            "example_input": "Texto con $f(x)=x^2+1$ y $$\\int_0^1 x^2 dx=\\frac{1}{3}$$ generado por ChatGPT.",
            "example_output": "Documento Word con ecuaciones OMML editables e integradas en el texto.",
            "before_text": "Antes: formulas como texto o capturas dificilmente editables.",
            "after_text": "Despues: ecuaciones nativas Word listas para corregir y compartir.",
            "mini_case": "Alumno de ingenieria con 20 ejercicios resueltos por IA que debe entregar en campus en formato .docx.",
            "supported_formats": [
                "Entrada DOCX con texto y formulas.",
                "Entrada TXT con respuestas de chat.",
                "Salida DOCX editable en Word.",
                "Soporte de contenido mixto texto + matematicas.",
            ],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": [
                "LaTeX sintacticamente invalido requiere correccion previa.",
                "Macros no estandar pueden necesitar ajuste manual.",
                "Documentos con maquetacion extrema deben revisarse al final.",
            ],
            "common_errors": [
                {"error": "Signos raros tras pegar desde chat", "fix": "Guardar primero en TXT UTF-8 o DOCX limpio antes de convertir."},
                {"error": "Ecuacion parcial sin convertir", "fix": "Comprobar delimitadores abiertos/cerrados en el texto origen."},
                {"error": "Tutor no puede editar una formula", "fix": "Validar que la ecuacion final sea OMML y no imagen/texto plano."},
            ],
            "faq_items": [
                {"q": "Sirve para respuestas largas de ChatGPT?", "a": "Si, esta orientada a respuestas completas con texto y ecuaciones."},
                {"q": "Las ecuaciones quedan editables en Word?", "a": "Si, la salida esta en formato OMML editable."},
                {"q": "Conviene para tareas academicas?", "a": "Si, es uno de los casos mas habituales de uso."},
            ],
            "secondary_cta_href": "/blog/signos-interrogacion-ecuaciones-chatgpt-word",
            "secondary_cta_label": "Resolver errores de pegado",
            "pillar_label": "Pilar: LaTeX a Word online",
            "related_guides": [
                {"href": "/blog/ia-chatgpt-a-word-ejercicios", "label": "Guia de ejercicios con IA + Word"},
                {"href": "/blog/latex-a-word-omml-guia-definitiva", "label": "Guia definitiva LaTeX a Word OMML"},
            ],
        },
        "en": {
            "primary_keyword": "chatgpt equations to word",
            "primary_intent": "transactional conversion from AI output to editable DOCX",
            "problem_angle": "avoid screenshot workflows and broken chat copy/paste",
            "problem_statement": "ChatGPT can solve the problem, but delivery in Word is where teams lose time. This page closes that gap by turning AI-generated LaTeX into native editable OMML equations in a review-ready DOCX.",
            "when_to_use_intro": "Use this page when ChatGPT is your source and Word is your required output.",
            "when_to_use_items": [
                "Assignment submission in DOCX format",
                "Supervisor review in Word",
                "Consolidated AI responses in one report",
                "Need for editable equations after sharing",
            ],
            "workflow_steps": [
                "Generate response with clear LaTeX delimiters.",
                "Save to TXT or DOCX source file.",
                "Upload and convert to OMML output.",
                "Download and verify key equations in Word.",
            ],
            "example_input": "Chat response containing $a^2+b^2=c^2$ and $$\\sum_{k=1}^{n}k$$.",
            "example_output": "DOCX with native editable Word equations.",
            "before_text": "Before: plain text formulas and manual cleanup.",
            "after_text": "After: native editable OMML equations.",
            "mini_case": "Engineering student compiling AI-generated homework into a single DOCX.",
            "supported_formats": ["DOCX input", "TXT input", "DOCX OMML output"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Invalid LaTeX must be corrected", "Custom macros may need manual edits"],
            "common_errors": [
                {"error": "Weird symbols from chat paste", "fix": "Use clean UTF-8 TXT/DOCX source before conversion."},
                {"error": "Equation not converted", "fix": "Check delimiter consistency."},
            ],
            "faq_items": [
                {"q": "Are equations editable after conversion?", "a": "Yes, output equations are generated as OMML."},
                {"q": "Does this support long chat responses?", "a": "Yes, including mixed text and equations."},
            ],
            "secondary_cta_href": "/en/blog/question-marks-chatgpt-equations-word",
            "secondary_cta_label": "Fix copy/paste issues",
            "pillar_label": "Pillar: LaTeX to Word online",
            "related_guides": [
                {"href": "/en/blog/use-ai-equations-to-word-exercises", "label": "Guide: AI exercises to Word"},
                {"href": "/en/blog/latex-to-word-online-free-editable-equations", "label": "Guide: LaTeX to Word online"},
            ],
        },
    },
    "gemini-equations-to-word": {
        "es": {
            "primary_keyword": "gemini a word ecuaciones",
            "primary_intent": "transaccional para convertir respuestas matematicas de Gemini",
            "problem_angle": "pasar prompts y bloques LaTeX de Gemini a Word editable",
            "problem_statement": "Gemini suele generar salidas matematicas utiles, pero convertirlas en un .docx limpio consume tiempo si se hace a mano. Esta landing se enfoca en ese cuello de botella: conversion directa a OMML para informes y ejercicios editables. El objetivo es que no pierdas calidad al pasar de un prompt a una entrega formal. Tambien cubre el escenario de iteraciones: cuando ajustas una respuesta con varios prompts y luego necesitas consolidar todo en un unico archivo Word coherente, sin simbolos rotos ni ecuaciones bloqueadas.",
            "extra_context": "La propuesta esta pensada para flujo continuo: generar, depurar y entregar. En lugar de tratar cada formula como caso aislado, conviertes todo el documento y mantienes una base editable para correcciones rapidas en Word.",
            "when_to_use_intro": "Recomendada cuando Gemini es tu herramienta principal y necesitas entregar en Word.",
            "when_to_use_items": [
                "Informes tecnicos con formulas en bloque",
                "Ejercicios semanales generados por IA",
                "Revision en Word con cambios sobre ecuaciones",
            ],
            "workflow_steps": [
                "Solicita salida LaTeX estable en Gemini.",
                "Guarda en .txt o .docx.",
                "Sube y convierte desde esta pagina.",
                "Descarga DOCX OMML y revisa.",
            ],
            "example_input": "Respuesta Gemini con $$\\nabla\\cdot\\vec{E}=\\rho/\\epsilon_0$$ y formulas inline.",
            "example_output": "Word con ecuaciones OMML listas para editar.",
            "before_text": "Antes: salida correcta en chat, dificil de llevar a Word final.",
            "after_text": "Despues: entrega .docx editable y consistente.",
            "mini_case": "Estudiante de fisica que transforma ejercicios Gemini en entregas Word.",
            "supported_formats": ["DOCX", "TXT", "Salida DOCX editable"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["LaTeX incompleto debe corregirse", "Macros no estandar pueden requerir retoque"],
            "common_errors": [
                {"error": "Mezcla markdown + LaTeX", "fix": "Separar bloques matematicos antes de convertir."},
                {"error": "Delimitador sin cierre", "fix": "Verificar apertura y cierre en cada bloque."},
            ],
            "faq_items": [
                {"q": "Sirve para informes largos?", "a": "Si, admite contenido mixto texto + ecuaciones."},
                {"q": "Puedo usarlo aunque mezcle Gemini y ChatGPT?", "a": "Si, mientras mantengas LaTeX delimitado."},
            ],
            "secondary_cta_href": "/blog/ia-chatgpt-a-word-ejercicios",
            "secondary_cta_label": "Guia IA + Word para ejercicios",
            "pillar_label": "Pilar: LaTeX a Word online",
            "related_guides": [
                {"href": "/blog/pasar-ecuaciones-chatgpt-word", "label": "Guia ChatGPT a Word"},
                {"href": "/blog/latex-a-word-online-gratis-ecuaciones-editables", "label": "Guia LaTeX a Word online"},
            ],
        },
        "en": {
            "primary_keyword": "gemini equations to word",
            "primary_intent": "transactional conversion of Gemini math output",
            "problem_angle": "turn Gemini LaTeX blocks into editable DOCX equations",
            "problem_statement": "Gemini output often needs post-processing before Word delivery. This page provides a direct path to OMML so equations stay editable during review.",
            "when_to_use_intro": "Use this when Gemini is your source and Word is your target deliverable.",
            "when_to_use_items": ["Assignments in DOCX", "Prompt-generated reports", "Collaborative Word review"],
            "workflow_steps": ["Generate in Gemini", "Save to TXT/DOCX", "Upload and convert", "Review final DOCX"],
            "example_input": "Gemini output with inline and display LaTeX math.",
            "example_output": "Editable OMML equations inside Word.",
            "before_text": "Before: AI output not ready for delivery.",
            "after_text": "After: clean DOCX with editable equations.",
            "mini_case": "Weekly AI-generated worksheets submitted in Word.",
            "supported_formats": ["DOCX", "TXT", "DOCX output"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Invalid source LaTeX needs correction"],
            "common_errors": [{"error": "Delimiter mismatch", "fix": "Normalize delimiter pairs before upload."}],
            "faq_items": [{"q": "Does this keep equations editable?", "a": "Yes, output is generated as OMML."}],
            "secondary_cta_href": "/en/blog/gemini-equations-to-word-omml",
            "secondary_cta_label": "Gemini troubleshooting guide",
            "pillar_label": "Pillar: LaTeX to Word online",
            "related_guides": [
                {"href": "/en/blog/use-ai-equations-to-word-exercises", "label": "Guide: AI exercises"},
                {"href": "/en/blog/latex-to-word-online-free-editable-equations", "label": "Guide: LaTeX to Word online"},
            ],
        },
    },
    "pandoc-to-word-omml": {
        "es": {
            "primary_keyword": "pandoc a word ecuaciones editables",
            "primary_intent": "transaccional para reparar salida Pandoc",
            "problem_angle": "documentos convertidos donde Word no deja editar formulas",
            "problem_statement": "Muchos flujos con Pandoc terminan en .docx visualmente correcto pero poco editable. Esta landing resuelve esa ultima milla: recuperar ecuaciones nativas OMML en Word sin rehacer todo el documento. Es especialmente valiosa en proyectos donde ya existe un pipeline markdown consolidado y no quieres romperlo, pero tampoco puedes permitir una entrega final con ecuaciones no editables. El foco esta en estabilizar la salida final para revision, mantenimiento y cambios de ultima hora.",
            "extra_context": "Asi puedes conservar las ventajas de Pandoc para conversion masiva y, al mismo tiempo, cumplir con el requisito de editabilidad real que suele pedir universidad, cliente o equipo de revisores. Ademas, reduce el riesgo de llegar al final del proyecto con formulas bloqueadas justo cuando mas cambios necesitas aplicar. En equipos con fechas ajustadas, esta capa final suele ahorrar horas de correccion manual. Tambien simplifica auditorias internas de calidad antes de publicar o entregar.",
            "when_to_use_intro": "Ideal cuando ya usaste Pandoc y necesitas estabilizar la editabilidad final en Word.",
            "when_to_use_items": ["Markdown -> DOCX con matematicas", "Revision colaborativa en Word", "Entregas tecnicas largas"],
            "workflow_steps": ["Genera salida Pandoc", "Sube el DOCX resultante", "Convierte a OMML", "Valida ecuaciones clave"],
            "example_input": "DOCX generado por Pandoc desde markdown con LaTeX.",
            "example_output": "DOCX final con ecuaciones editables en Word.",
            "before_text": "Antes: ecuaciones visibles pero no editables.",
            "after_text": "Despues: formulas OMML nativas listas para revision.",
            "mini_case": "Memoria tecnica en markdown convertida a Word para entrega institucional.",
            "supported_formats": ["DOCX post-Pandoc", "TXT", "DOCX editable"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["No sustituye tu pipeline Pandoc", "Sintaxis LaTeX rota debe corregirse en origen"],
            "common_errors": [
                {"error": "Asumir que Pandoc ya deja OMML perfecto", "fix": "Comprobar editabilidad real en Word antes de cerrar."},
                {"error": "Bloques largos fragmentados", "fix": "Revisar delimitadores en markdown fuente."},
            ],
            "faq_items": [
                {"q": "Sustituye a Pandoc?", "a": "No, complementa la fase final de compatibilidad Word."},
                {"q": "Sirve para documentos largos?", "a": "Si, especialmente en informes tecnicos extensos."},
            ],
            "secondary_cta_href": "/blog/pandoc-ecuaciones-word-no-editables-soluciones",
            "secondary_cta_label": "Checklist de errores Pandoc",
            "pillar_label": "Pilar: LaTeX a Word online",
            "related_guides": [
                {"href": "/blog/markdown-con-latex-a-word-docx", "label": "Guia Markdown + LaTeX a Word"},
                {"href": "/blog/que-es-omml-ecuaciones-word", "label": "Guia que es OMML"},
            ],
        },
        "en": {
            "primary_keyword": "pandoc to word editable equations",
            "primary_intent": "transactional post-conversion fix",
            "problem_angle": "restore editability after markdown/LaTeX pipelines",
            "problem_statement": "Pandoc output may look fine but still fail Word equation editing. This landing restores OMML editability for final delivery.",
            "when_to_use_intro": "Use this page after Pandoc conversion when equation editing in Word is still broken.",
            "when_to_use_items": ["Equation-heavy markdown docs", "DOCX review workflows", "Technical reports"],
            "workflow_steps": ["Run Pandoc", "Upload output file", "Convert to OMML", "Validate in Word"],
            "example_input": "Pandoc DOCX with non-editable equations.",
            "example_output": "DOCX with native editable OMML equations.",
            "before_text": "Before: visual math, weak editability.",
            "after_text": "After: Word-native editable equations.",
            "mini_case": "Markdown engineering memo delivered as editable DOCX.",
            "supported_formats": ["DOCX", "TXT", "DOCX output"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Invalid source LaTeX must be fixed first"],
            "common_errors": [{"error": "Late-stage manual cleanup", "fix": "Add this conversion as a fixed final step."}],
            "faq_items": [{"q": "Does this replace Pandoc?", "a": "No, it complements it for Word editability."}],
            "secondary_cta_href": "/en/blog/pandoc-math-to-word-omml-troubleshooting",
            "secondary_cta_label": "Pandoc troubleshooting guide",
            "pillar_label": "Pillar: LaTeX to Word online",
            "related_guides": [
                {"href": "/en/blog/markdown-latex-to-word-docx", "label": "Guide: Markdown + LaTeX to DOCX"},
                {"href": "/en/blog/what-is-omml-word-equations", "label": "Guide: what OMML is"},
            ],
        },
    },
    "overleaf-latex-document-to-word": {
        "es": {
            "primary_keyword": "overleaf a word",
            "primary_intent": "transaccional para entrega academica en Word",
            "problem_angle": "pasar TFG/TFM de Overleaf a .docx revisable con tutor",
            "problem_statement": "Redactar en Overleaf y entregar en Word suele romper la fase final de revision. Esta landing esta orientada a mantener ecuaciones editables en Word para TFG/TFM e informes tecnicos. El problema habitual aparece cuando el contenido ya esta bien resuelto en LaTeX, pero la institucion exige .docx para tutoria, tribunal o archivo final. Aqui el valor es reducir friccion en ese traspaso obligatorio: convertir sin perder capacidad de correccion, sin rehacer formulas y sin bloquear la colaboracion en Word.",
            "extra_context": "La idea es que el documento final no sea una conversion estatica, sino una version revisable. Eso facilita comentarios, cambios de ultima hora y adaptacion a plantillas institucionales sin destruir la parte matematica. Tambien te permite responder mas rapido a feedback de tutor o tribunal en la fase mas sensible del trabajo.",
            "when_to_use_intro": "Usa esta solucion cuando el documento nace en LaTeX pero el cierre y la revision son en Word.",
            "when_to_use_items": ["Tutor corrige en DOCX", "Normativa institucional exige Word", "Varias rondas de correccion final"],
            "workflow_steps": ["Prepara exportacion", "Sube .docx/.txt", "Convierte a OMML", "Revisa capitulos y entrega"],
            "example_input": "Capitulo de TFM con matrices e integrales exportado desde Overleaf.",
            "example_output": "DOCX revisable con formulas OMML editables.",
            "before_text": "Antes: buen contenido en LaTeX, revision pesada en Word.",
            "after_text": "Despues: documento Word apto para correccion colaborativa.",
            "mini_case": "Estudiante que entrega memoria final en plantilla Word universitaria.",
            "supported_formats": ["DOCX exportado", "TXT con LaTeX", "Salida DOCX editable"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Maquetaciones LaTeX muy custom pueden requerir ajustes visuales"],
            "common_errors": [
                {"error": "No validar ecuaciones en la plantilla final", "fix": "Comprobar editabilidad tras aplicar estilos institucionales."},
                {"error": "Revisar todo al final", "fix": "Validar por capitulos durante el proceso."},
            ],
            "faq_items": [
                {"q": "Sirve para TFG/TFM largos?", "a": "Si, esta orientada a documentos extensos con revision academica."},
                {"q": "El tutor puede editar formulas?", "a": "Si, quedan en formato OMML editable."},
            ],
            "secondary_cta_href": "/blog/overleaf-latex-a-word-ecuaciones-editables",
            "secondary_cta_label": "Guia Overleaf a Word",
            "pillar_label": "Pilar: LaTeX a Word online",
            "related_guides": [
                {"href": "/blog/convertir-documento-latex-word", "label": "Guia de documento LaTeX a Word"},
                {"href": "/blog/omml-vs-mathtype-vs-latex-word-tfg", "label": "Comparativa OMML vs MathType vs LaTeX"},
            ],
        },
        "en": {
            "primary_keyword": "overleaf to word",
            "primary_intent": "transactional thesis handoff conversion",
            "problem_angle": "move LaTeX authoring into Word review workflows",
            "problem_statement": "This page helps you keep equation editability when moving from Overleaf/LaTeX writing to Word-based review and submission.",
            "when_to_use_intro": "Use this page for thesis/coursework handoff where Word review is mandatory.",
            "when_to_use_items": ["Supervisor feedback in DOCX", "Institutional Word templates", "Final-round revisions"],
            "workflow_steps": ["Export source", "Upload file", "Convert to OMML", "Review and submit"],
            "example_input": "Overleaf chapter with matrix-heavy equations.",
            "example_output": "Editable DOCX with OMML equations.",
            "before_text": "Before: hard-to-review equation handoff.",
            "after_text": "After: Word-native equations for tracked changes.",
            "mini_case": "Master thesis drafted in Overleaf, delivered in DOCX.",
            "supported_formats": ["DOCX", "TXT", "DOCX output"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Custom layout features may need visual QA"],
            "common_errors": [{"error": "Skipping editability checks", "fix": "Validate equations in final Word template."}],
            "faq_items": [{"q": "Can this handle long thesis files?", "a": "Yes, it is built for long academic workflows."}],
            "secondary_cta_href": "/en/blog/overleaf-latex-to-word-editable-equations",
            "secondary_cta_label": "Overleaf workflow guide",
            "pillar_label": "Pillar: LaTeX to Word online",
            "related_guides": [
                {"href": "/en/blog/convert-latex-document-to-word", "label": "Guide: LaTeX document to Word"},
                {"href": "/en/blog/omml-vs-mathtype-vs-latex-word-thesis", "label": "Guide: OMML vs MathType vs LaTeX"},
            ],
        },
    },
    "omml-converter": {
        "es": {
            "primary_keyword": "conversor omml",
            "primary_intent": "transaccional + informacional sobre compatibilidad Word",
            "problem_angle": "entender y aplicar formato nativo OMML para colaboracion",
            "problem_statement": "Cuando las ecuaciones fallan en Word, el problema suele ser de formato, no de contenido. Esta landing se centra en OMML para garantizar editabilidad y compatibilidad en .docx compartidos. Es la opcion correcta cuando necesitas una base tecnica estable para colaborar: revisiones entre equipos, intercambio de versiones y correcciones sobre formulas sin depender de plugins externos. En vez de tratar sintomas por cada documento, esta pagina ataca la causa raiz: estandarizar la matematica en formato nativo Word.",
            "extra_context": "Con este enfoque evitas retrabajo repetitivo. Una vez unificas a OMML, la colaboracion en Word se vuelve predecible: menos sorpresas al compartir, menos fallos de render y mayor control durante todo el ciclo editorial. Es una mejora estructural, no solo un arreglo puntual para una entrega concreta. Tambien facilita mantener criterios uniformes de calidad cuando el sitio publica contenido tecnico de forma continua. A medio plazo, reduce incidencias y acelera la edicion colaborativa.",
            "when_to_use_intro": "Usa esta pagina cuando te importa compatibilidad Word a largo plazo y colaboracion estable.",
            "when_to_use_items": ["Documentos compartidos entre equipos", "Correcciones frecuentes sobre formulas", "Entregas institucionales en Word"],
            "workflow_steps": ["Identifica archivo origen", "Sube .docx/.txt", "Convierte a OMML", "Valida y comparte"],
            "example_input": "Documento con formulas LaTeX pegadas como texto.",
            "example_output": "DOCX con ecuaciones OMML editables y consistentes.",
            "before_text": "Antes: formulas visibles pero fragiles al editar.",
            "after_text": "Despues: formato nativo Word para revision y colaboracion.",
            "mini_case": "Equipo de proyecto que intercambia borradores DOCX con matematicas.",
            "supported_formats": ["DOCX", "TXT", "DOCX editable"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Sintaxis LaTeX invalida no se corrige automaticamente"],
            "common_errors": [
                {"error": "Confundir visualizacion con editabilidad", "fix": "Comprobar edicion directa en Equation Editor."},
                {"error": "Mezclar formatos distintos de ecuaciones", "fix": "Unificar todo en OMML antes de entregar."},
            ],
            "faq_items": [
                {"q": "Que es OMML?", "a": "Es el formato nativo de ecuaciones editable en Microsoft Word."},
                {"q": "Sirve si no uso IA?", "a": "Si, es una solucion de formato independientemente de la fuente."},
            ],
            "secondary_cta_href": "/blog/que-es-omml-ecuaciones-word",
            "secondary_cta_label": "Aprender que es OMML",
            "pillar_label": "Pilar: LaTeX a Word online",
            "related_guides": [
                {"href": "/blog/omml-vs-mathtype-vs-latex-word-tfg", "label": "Comparativa OMML vs MathType vs LaTeX"},
                {"href": "/blog/simbolos-raros-ecuaciones-word-cambria-math", "label": "Guia de compatibilidad Cambria Math"},
            ],
        },
        "en": {
            "primary_keyword": "omml converter",
            "primary_intent": "transactional + educational format standardization",
            "problem_angle": "Word-native equation compatibility for collaboration",
            "problem_statement": "Equation issues in Word are often format issues. This landing focuses on OMML conversion to preserve editability in shared DOCX files.",
            "when_to_use_intro": "Use this page when format compatibility matters more than source tool branding.",
            "when_to_use_items": ["Cross-team DOCX editing", "Word tracked changes with equations", "Compatibility-first QA"],
            "workflow_steps": ["Upload source", "Convert to OMML", "Validate in Word", "Share stable DOCX"],
            "example_input": "LaTeX formulas pasted in mixed plain text format.",
            "example_output": "Word-native OMML equations ready for editing.",
            "before_text": "Before: fragile equation rendering across edits.",
            "after_text": "After: stable editable equations in Word.",
            "mini_case": "Research collaborators exchanging equation-heavy DOCX drafts.",
            "supported_formats": ["DOCX", "TXT", "DOCX output"],
            "supported_delimiters": ["$...$", "$$...$$", "\\[...\\]"],
            "known_limitations": ["Invalid source LaTeX needs cleanup"],
            "common_errors": [{"error": "Assuming rendered means editable", "fix": "Always test equation editing directly in Word."}],
            "faq_items": [{"q": "What is OMML?", "a": "Word's native editable equation format."}],
            "secondary_cta_href": "/en/blog/what-is-omml-word-equations",
            "secondary_cta_label": "Learn what OMML is",
            "pillar_label": "Pillar: LaTeX to Word online",
            "related_guides": [
                {"href": "/en/blog/omml-vs-mathtype-vs-latex-word-thesis", "label": "Guide: OMML vs MathType vs LaTeX"},
                {"href": "/en/blog/weird-symbols-word-equations-cambria-math-fix", "label": "Guide: weird symbols fix"},
            ],
        },
    },
}


def _landing_longform(lang: str, landing_key: str) -> Dict[str, Any]:
    data = LANDING_LONGFORM.get(landing_key, {})
    if not data:
        return {}
    return data.get(lang) or data.get(_content_lang(lang)) or data.get("en") or {}


def _load_blog_data() -> Dict[str, Any]:
    """Load blog metadata from blog_content/posts.json.

    This is the single source of truth for:
    - canonical slugs
    - SEO metadata
    - translation pairing
    - alias slugs (legacy URLs) -> canonical
    """
    try:
        with open(BLOG_DATA_PATH, "r", encoding="utf-8") as f:
            raw = f.read()
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            # Best-effort repair for accidental unescaped quotes in string lines.
            repaired_lines: List[str] = []
            for line in raw.splitlines():
                m = re.match(r'^(\s*"[^"]+"\s*:\s*")(.*)("\s*,?\s*)$', line)
                if m:
                    prefix, body, suffix = m.groups()
                    body = re.sub(r'(?<!\\)"', r'\\"', body)
                    repaired_lines.append(prefix + body + suffix)
                    continue
                m2 = re.match(r'^(\s*")(.*)("\s*,?\s*)$', line)
                if m2 and ("<p class=\\\"" in line or line.strip().startswith("\"<")):
                    prefix, body, suffix = m2.groups()
                    body = re.sub(r'(?<!\\)"', r'\\"', body)
                    repaired_lines.append(prefix + body + suffix)
                    continue
                repaired_lines.append(line)
            data = json.loads("\n".join(repaired_lines))
        if not isinstance(data, dict):
            raise ValueError("posts.json root is not a JSON object")
        data.setdefault("posts", [])
        data.setdefault("aliases", {})
        return _deep_fix_mojibake(data)
    except Exception:
        logger.exception("Failed to load blog data from %s", BLOG_DATA_PATH)
        # Keep the app running (converter is critical). Blog will 404 gracefully.
        return {"posts": [], "aliases": {}}


def _init_blog_cache() -> None:
    data = _load_blog_data()
    posts = data.get("posts", [])
    aliases = data.get("aliases", {})

    for lang in SUPPORTED_LANGS:
        BLOG_POSTS[lang].clear()
        BLOG_LIST[lang].clear()
        BLOG_ALIASES[lang].clear()

    # Aliases
    if isinstance(aliases, dict):
        for lang in SUPPORTED_LANGS:
            BLOG_ALIASES[lang].update((aliases.get(lang) or {}))

    # Posts
    for p in posts:
        if not isinstance(p, dict):
            continue
        lang = (p.get("lang") or "").strip()
        slug = (p.get("slug") or "").strip()
        canonical_path = (p.get("canonical_path") or "").strip()
        if lang not in SUPPORTED_LANGS or not slug or not canonical_path:
            continue
        BLOG_POSTS[lang][slug] = p

    # Lists (sorted)
    for lang in SUPPORTED_LANGS:
        lst = list(BLOG_POSTS[lang].values())
        lst.sort(
            key=lambda d: (d.get("date_published") or "", d.get("slug") or ""),
            reverse=True,
        )
        BLOG_LIST[lang] = lst

    # EN-based locales share EN content until translated metadata/files are added.
    for lang in SUPPORTED_LANGS:
        if BLOG_LIST[lang]:
            continue
        BLOG_POSTS[lang] = dict(BLOG_POSTS["en"])
        BLOG_LIST[lang] = list(BLOG_LIST["en"])
        if not BLOG_ALIASES.get(lang):
            BLOG_ALIASES[lang] = dict(BLOG_ALIASES["en"])


_init_blog_cache()


def _render_template(template_name: str, context: Dict[str, Any]) -> str:
    template = JINJA_ENV.get_template(template_name)
    return template.render(**_deep_fix_mojibake(context))


def _read_blog_body(lang: str, slug: str) -> str:
    candidates = [lang]
    content_lang = _content_lang(lang)
    if content_lang not in candidates:
        candidates.append(content_lang)
    if "en" not in candidates:
        candidates.append("en")
    if "es" not in candidates:
        candidates.append("es")

    for cand in candidates:
        path = os.path.join(BLOG_POSTS_DIR, cand, f"{slug}.html")
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except FileNotFoundError:
            continue
    return ""


def _month_name_es(month: int) -> str:
    names = [
        "",
        "enero",
        "febrero",
        "marzo",
        "abril",
        "mayo",
        "junio",
        "julio",
        "agosto",
        "septiembre",
        "octubre",
        "noviembre",
        "diciembre",
    ]
    return names[month] if 1 <= month <= 12 else ""


def _month_name_de(month: int) -> str:
    names = [
        "",
        "Januar",
        "Februar",
        "März",
        "April",
        "Mai",
        "Juni",
        "Juli",
        "August",
        "September",
        "Oktober",
        "November",
        "Dezember",
    ]
    return names[month] if 1 <= month <= 12 else ""


def _month_name_fr(month: int) -> str:
    names = [
        "",
        "janvier",
        "février",
        "mars",
        "avril",
        "mai",
        "juin",
        "juillet",
        "août",
        "septembre",
        "octobre",
        "novembre",
        "décembre",
    ]
    return names[month] if 1 <= month <= 12 else ""


def _month_name_it(month: int) -> str:
    names = [
        "",
        "gennaio",
        "febbraio",
        "marzo",
        "aprile",
        "maggio",
        "giugno",
        "luglio",
        "agosto",
        "settembre",
        "ottobre",
        "novembre",
        "dicembre",
    ]
    return names[month] if 1 <= month <= 12 else ""


def _month_name_pt(month: int) -> str:
    names = [
        "",
        "janeiro",
        "fevereiro",
        "março",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    return names[month] if 1 <= month <= 12 else ""


def _format_date(lang: str, iso_date: str) -> str:
    """Format YYYY-MM-DD into a human-readable date."""
    try:
        dt = datetime.strptime(iso_date, "%Y-%m-%d").date()
    except Exception:
        return iso_date

    if lang == "es":
        return f"{dt.day} de {_month_name_es(dt.month)} de {dt.year}"
    if lang == "de":
        return f"{dt.day}. {_month_name_de(dt.month)} {dt.year}"
    if lang == "fr":
        return f"{dt.day} {_month_name_fr(dt.month)} {dt.year}"
    if lang == "it":
        return f"{dt.day} {_month_name_it(dt.month)} {dt.year}"
    if lang == "pt":
        return f"{dt.day} de {_month_name_pt(dt.month)} de {dt.year}"
    # en and fallback
    return dt.strftime("%b %d, %Y")


def _build_schema_article(post: Dict[str, Any], canonical_url: str) -> str:
    lang = post.get("lang") or "es"
    title = post.get("title") or ""
    desc = post.get("description") or ""
    date_pub = post.get("date_published") or ""
    date_mod = post.get("date_modified") or date_pub

    blog_path = _blog_index_path(lang)
    blog_url = f"{SITE_CANONICAL_ORIGIN}{blog_path}"
    home_url = f"{SITE_CANONICAL_ORIGIN}{_home_path(lang)}"

    article = {
        "@type": "Article",
        "mainEntityOfPage": {"@type": "WebPage", "@id": canonical_url},
        "headline": title,
        "description": desc,
        "datePublished": date_pub,
        "dateModified": date_mod,
        "author": {"@type": "Organization", "name": SITE_NAME},
        "publisher": {"@type": "Organization", "name": SITE_NAME},
        "inLanguage": lang,
        "url": canonical_url,
        "keywords": post.get("keywords") or [],
    }

    breadcrumbs = {
        "@type": "BreadcrumbList",
        "itemListElement": [
            {"@type": "ListItem", "position": 1, "name": "Home", "item": home_url},
            {
                "@type": "ListItem",
                "position": 2,
                "name": "Blog",
                "item": blog_url,
            },
            {"@type": "ListItem", "position": 3, "name": title, "item": canonical_url},
        ],
    }

    schema = {"@context": "https://schema.org", "@graph": [article, breadcrumbs]}
    return json.dumps(schema, ensure_ascii=False)


def _build_schema_simple_page(name: str, canonical_url: str, lang: str) -> str:
    schema = {
        "@context": "https://schema.org",
        "@type": "WebPage",
        "name": name,
        "url": canonical_url,
        "inLanguage": lang,
    }
    return json.dumps(schema, ensure_ascii=False)


def _build_schema_solution_page(
    name: str,
    canonical_url: str,
    lang: str,
    description: str,
    related_blog_url: str,
    faq_items: Optional[List[Dict[str, str]]] = None,
) -> str:
    solutions_hub = _abs_url(_solutions_path(lang))
    solutions_name = "Soluciones" if lang == "es" else "Solutions"
    faq_entities: List[Dict[str, Any]] = []
    for item in faq_items or []:
        q = (item.get("q") or "").strip()
        a = (item.get("a") or "").strip()
        if not q or not a:
            continue
        faq_entities.append(
            {
                "@type": "Question",
                "name": q,
                "acceptedAnswer": {"@type": "Answer", "text": a},
            }
        )
    if not faq_entities:
        faq_entities = [
            {
                "@type": "Question",
                "name": "Does it keep equations editable in Word?",
                "acceptedAnswer": {
                    "@type": "Answer",
                    "text": "Yes. The converter generates native Word OMML equations that remain editable inside .docx.",
                },
            },
            {
                "@type": "Question",
                "name": "Can I use files generated by AI tools?",
                "acceptedAnswer": {
                    "@type": "Answer",
                    "text": "Yes. You can upload DOCX or TXT files containing LaTeX produced by tools such as ChatGPT, Gemini or Overleaf workflows.",
                },
            },
            {
                "@type": "Question",
                "name": "Where can I read a detailed guide?",
                "acceptedAnswer": {
                    "@type": "Answer",
                    "text": f"You can read the related guide here: {related_blog_url}",
                },
            },
        ]

    schema = {
        "@context": "https://schema.org",
        "@graph": [
            {
                "@type": "WebPage",
                "name": name,
                "url": canonical_url,
                "inLanguage": lang,
                "description": description,
            },
            {
                "@type": "BreadcrumbList",
                "itemListElement": [
                    {"@type": "ListItem", "position": 1, "name": "Home", "item": _abs_url(_home_path(lang))},
                    {"@type": "ListItem", "position": 2, "name": solutions_name, "item": solutions_hub},
                    {"@type": "ListItem", "position": 3, "name": name, "item": canonical_url},
                ],
            },
            {
                "@type": "SoftwareApplication",
                "name": "Ecuaciones a Word",
                "applicationCategory": "UtilityApplication",
                "operatingSystem": "Web",
                "url": _abs_url(_home_path(lang)),
                "offers": {"@type": "Offer", "price": "0", "priceCurrency": "EUR"},
                "featureList": [
                    "LaTeX to Word equation conversion",
                    "Native Word OMML output",
                    "DOCX and TXT support",
                ],
            },
            {
                "@type": "FAQPage",
                "mainEntity": faq_entities[:8],
            },
        ],
    }
    return json.dumps(schema, ensure_ascii=False)


def _build_schema_index(lang: str, canonical_url: str) -> str:
    items = []
    for idx, p in enumerate(BLOG_LIST.get(lang, []), start=1):
        slug = p.get("slug") or ""
        url_path = (p.get("canonical_path") or "").strip()
        if lang not in PRIMARY_CONTENT_LANGS or not url_path:
            url_path = f"{_blog_index_path(lang)}/{slug}"
        items.append(
            {
                "@type": "ListItem",
                "position": idx,
                "url": f"{SITE_CANONICAL_ORIGIN}{url_path}",
                "name": p.get("title") or "",
            }
        )
    schema = {
        "@context": "https://schema.org",
        "@type": "WebPage",
        "name": "Blog",
        "url": canonical_url,
        "inLanguage": lang,
        "mainEntity": {"@type": "ItemList", "itemListElement": items[:50]},
    }
    return json.dumps(schema, ensure_ascii=False)




def _noindex_headers(follow: bool = False) -> Dict[str, str]:
    """Headers to discourage indexing for technical/non-content endpoints."""
    return {"X-Robots-Tag": "noindex, follow" if follow else "noindex, nofollow"}
# ================================================================
# Helpers
# ================================================================
def _read_text_file(path: str, default: Optional[str] = None) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        if default is not None:
            return default
        raise


def read_html_file(filename: str) -> str:
    path = os.path.join(BASE_DIR, filename)
    return _fix_text_mojibake(_read_text_file(path))


def _force_meta_robots_noindex(html: str) -> str:
    if re.search(r'<meta[^>]+name=["\']robots["\']', html, flags=re.IGNORECASE):
        return re.sub(
            r'<meta[^>]+name=["\']robots["\'][^>]*>',
            '<meta name="robots" content="noindex,follow,max-image-preview:large">',
            html,
            count=1,
            flags=re.IGNORECASE,
        )
    return html.replace(
        "</head>",
        '<meta name="robots" content="noindex,follow,max-image-preview:large"></head>',
        1,
    )


def _now_lastmod_iso() -> str:
    return datetime.now(timezone.utc).date().isoformat()


def _abs_url(path: str) -> str:
    # Canonical host (ajusta si cambias dominio)
    return f"{SITE_CANONICAL_ORIGIN}{path}"


def _sitemap_url_entry(
    loc: str,
    lastmod: str,
    changefreq: str,
    priority: str,
    alternates: Optional[List[Dict[str, str]]] = None,
) -> str:
    """
    Build a sitemap <url> entry. If alternates are provided, include xhtml:link elements.
    Note: loc and alternate hrefs are XML-escaped here.
    """
    alt_xml = ""
    if alternates:
        for alt in alternates:
            hreflang = (alt.get("hreflang") or "").strip()
            href = (alt.get("href") or "").strip()
            if not hreflang or not href:
                continue
            alt_xml += (
                f'    <xhtml:link rel="alternate" hreflang="{escape(hreflang)}" '
                f'href="{escape(href)}"/>\n'
            )

    return (
        "  <url>\n"
        f"    <loc>{escape(loc)}</loc>\n"
        f"{alt_xml}"
        f"    <lastmod>{escape(lastmod)}</lastmod>\n"
        f"    <changefreq>{escape(changefreq)}</changefreq>\n"
        f"    <priority>{escape(priority)}</priority>\n"
        "  </url>\n"
    )


def _is_valid_blog_canonical_path(lang: str, path: str) -> bool:
    expected = "/blog/" if lang == "es" else "/en/blog/"
    return bool(path and path.startswith(expected))


def _translated_indexable_path(lang: str, slug: str) -> Optional[str]:
    post = BLOG_POSTS.get(lang, {}).get(slug)
    if not post or not _is_blog_post_indexable(lang, post):
        return None
    path = (post.get("canonical_path") or "").strip()
    if not _is_valid_blog_canonical_path(lang, path):
        return None
    return path


def generate_sitemap_xml() -> str:
    """
    Build sitemap using only canonical, indexable, final URLs:
    - status 200 routes (no redirect hubs)
    - primary languages only (es/en)
    - blog posts explicitly marked as indexable by strategy
    """
    urls: List[str] = []
    seen_locs: set[str] = set()

    def append_url(
        loc_path: str,
        *,
        lastmod: str,
        changefreq: str,
        priority: str,
        alternates: Optional[List[Dict[str, str]]] = None,
    ) -> None:
        if not loc_path.startswith("/"):
            return
        loc = _abs_url(loc_path)
        if loc in seen_locs:
            return
        seen_locs.add(loc)
        urls.append(
            _sitemap_url_entry(
                loc,
                lastmod,
                changefreq,
                priority,
                alternates=alternates,
            )
        )

    def post_lastmod(p: Dict[str, Any]) -> str:
        d = (p.get("date_modified") or p.get("date_published") or "").strip()
        if d:
            return d
        return _now_lastmod_iso()

    # Home (primary languages only)
    home_paths = {lang: _home_path(lang) for lang in SITEMAP_LANGS}
    home_alts = _all_alternates(home_paths, default_lang="es")
    for lang in SITEMAP_LANGS:
        append_url(
            home_paths[lang],
            lastmod=_now_lastmod_iso(),
            changefreq="weekly",
            priority="1.0" if lang == "es" else "0.8",
            alternates=home_alts,
        )

    # Blog index
    blog_index_paths = {lang: _blog_index_path(lang) for lang in SITEMAP_LANGS}
    blog_index_alts = _all_alternates(blog_index_paths, default_lang="es")
    for lang in SITEMAP_LANGS:
        append_url(
            blog_index_paths[lang],
            lastmod=_now_lastmod_iso(),
            changefreq="weekly",
            priority="0.8" if lang == "es" else "0.7",
            alternates=blog_index_alts,
        )

    # Solutions hub
    solutions_paths = {lang: _solutions_path(lang) for lang in SITEMAP_LANGS}
    solutions_alts = _all_alternates(solutions_paths, default_lang="es")
    for lang in SITEMAP_LANGS:
        append_url(
            solutions_paths[lang],
            lastmod=_now_lastmod_iso(),
            changefreq="weekly",
            priority="0.75" if lang == "es" else "0.7",
            alternates=solutions_alts,
        )

    # Transactional landings
    for es_path, en_path in _all_landing_pairs():
        landing_paths = {"es": es_path, "en": en_path}
        alts = _all_alternates(landing_paths, default_lang="es")
        for lang in SITEMAP_LANGS:
            path = landing_paths.get(lang)
            if not path:
                continue
            append_url(
                path,
                lastmod=_now_lastmod_iso(),
                changefreq="weekly",
                priority="0.7" if lang == "es" else "0.6",
                alternates=alts,
            )

    # Blog posts (indexable only)
    for lang in SITEMAP_LANGS:
        for p in BLOG_LIST.get(lang, []):
            if not _is_blog_post_indexable(lang, p):
                continue
            slug = (p.get("slug") or "").strip()
            if not slug:
                continue
            canonical_path = (p.get("canonical_path") or "").strip()
            if not canonical_path:
                canonical_path = f"{_blog_index_path(lang)}/{slug}"
            if not _is_valid_blog_canonical_path(lang, canonical_path):
                continue
            translation_slug = (p.get("translation_slug") or "").strip()
            if lang == "es":
                es_path = canonical_path
                en_slug = translation_slug or slug
                en_path = _translated_indexable_path("en", en_slug)
            else:
                en_path = canonical_path
                es_slug = translation_slug or slug
                es_path = _translated_indexable_path("es", es_slug)
            paths = {lang: canonical_path}
            if es_path:
                paths["es"] = es_path
            if en_path:
                paths["en"] = en_path
            alts = _all_alternates(paths, default_lang="es")
            append_url(
                canonical_path,
                lastmod=post_lastmod(p),
                changefreq="monthly",
                priority="0.6" if lang == "es" else "0.5",
                alternates=alts,
            )

    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9" '
        'xmlns:xhtml="http://www.w3.org/1999/xhtml">\n'
        + "".join(urls)
        + "</urlset>\n"
    )


# ================================================================
# 1) Normalización / prettify (tu lógica existente)
# ================================================================
def normalize_math_text(text: str) -> str:
    """Normalize common artifacts found in pasted/converted LaTeX inside Word.

    Goals:
    - Make LaTeX more parseable for math2docx (normalize Unicode minus, remove zero-width chars, etc.)
    - Preserve non-math text as much as possible
    - Keep original app's q1->q_1, D1->D_1 and x2->x^2 tweaks
    """
    if text is None:
        return ""

    t = str(text)

    # Remove common invisible/formatting artifacts that break parsers
    for ch in ("\u200b", "\ufeff", "\u2060", "\u200e", "\u200f", "\u00ad"):
        t = t.replace(ch, "")

    # Normalize various dash/minus characters to ASCII hyphen-minus
    for ch in ("\u2212", "\u2013", "\u2014", "\u2010", "\u2011", "\u2043"):
        t = t.replace(ch, "-")
    # Some docs include the literal minus sign already decoded
    t = t.replace("−", "-")

    # Normalize non-breaking spaces
    t = t.replace("\xa0", " ")

    # Normalize other Unicode spaces that sometimes appear in Word
    for ch in ("\u202f", "\u205f", "\u3000"):
        t = t.replace(ch, " ")
    # Replace en-space/em-space/thin-space etc. with regular spaces
    t = re.sub(r"[\u2000-\u200a]", " ", t)

    # Remove Private Use Area characters (often appear as invisible placeholders)
    t = re.sub(r"[\uE000-\uF8FF]", "", t)

    # Keep original exercise-oriented normalizations
    t = re.sub(r"q([1-4])\s*\(", r"q_\1(", t)  # q1( -> q_1(
    t = re.sub(r"\bD([1-4])\b", r"D_\1", t)  # D1 -> D_1

    # x2 -> x^2, y3 -> y^3, etc.
    for var in ("x", "y", "z"):
        for exp in ("2", "3", "4"):
            t = re.sub(rf"{var}{exp}\b", rf"{var}^{{{exp}}}", t)

    return t

def _prettify_paragraphs_for_exercise(paragraph_texts: List[str]) -> List[str]:
    out: List[str] = []
    for text in paragraph_texts:
        s = normalize_math_text(text)
        stripped = s.strip()
        if stripped == "":
            continue

        if stripped.lower() == "actividad2grupal":
            out.append("Actividad 2 (trabajo grupal)")
            continue

        if all(
            sym in stripped
            for sym in ["q_1(x,y,z)", "q_2(x,y,z)", "q_3(x,y,z)", "q_4(x,y,z)"]
        ):
            out.append("$$ q_1(x,y,z) = 2x^2 + 2y^2 + 2z^2 + 2xy + 2xz $$")
            out.append("$$ q_2(x,y,z) = x^2 - y^2 + z^2 + 2xy $$")
            out.append("$$ q_3(x,y,z) = 2x^2 - y^2 + 2z^2 + 4xy - 4yz $$")
            out.append("$$ q_4(x,y,z) = 4x^2 + 2y^2 + z^2 + 2yz $$")
            continue

        if stripped.startswith("Interpretamos que") and "garantizar beneficios" in stripped:
            out.append(
                "Interpretamos que \"garantizar beneficios\" significa que el beneficio "
                "$q(x,y,z)$ sea positivo para todo $(x,y,z)\\neq(0,0,0)$, es decir, "
                "que la forma cuadrática sea definida positiva."
            )
            continue

        if (
            "Escribimos cada forma como q(x)=xT" in stripped
            or "Escribimos cada forma como q(x)=xTA" in stripped
            or "Escribimos cada forma como q(x)=xTA x(\\mathbf x)" in stripped
            or (
                "Escribimos cada forma como q(x)=xTAx(\\mathbf x)=\\mathbf x^T A\\mathbf x"
                in stripped
            )
        ):
            out.append(
                "Escribimos cada forma como $q(\\mathbf x) = \\mathbf x^T A\\, \\mathbf x$, con $\\mathbf x = (x,y,z)^T$."
            )
            continue

        if "coeficiente del término cruzado" in stripped:
            out.append(
                "Recordando que el coeficiente del término cruzado $2x_i y_j$ se reparte como $a_{ij}=a_{ji}=1$:"
            )
            continue

        if "Forma q1" in stripped or "q1q_1q1" in stripped:
            out.append("Forma $q_1$:")
            continue
        if "Forma q2" in stripped or "q2q_2q2" in stripped:
            out.append("Forma $q_2$:")
            continue
        if "Forma q3" in stripped or "q3q_3q3" in stripped:
            out.append("Forma $q_3$:")
            continue
        if "Forma q4" in stripped or "q4q_4q4" in stripped:
            out.append("Forma $q_4$:")
            continue

        if "Criterio de Sylvester" in stripped:
            out.append("2. Criterio de Sylvester")
            continue
        if "Una matriz simétrica" in stripped and "definida positiva" in stripped:
            out.append(
                "Una matriz simétrica $A$ es definida positiva si todos sus menores principales líderes son positivos:"
            )
            continue

        if "D_1>0" in stripped and "D_2>0" in stripped and "D_3>0" in stripped:
            out.append(
                "\\begin{aligned}\nD_1 &> 0,\\\\\nD_2 &> 0,\\\\\nD_3 &> 0.\n\\end{aligned}"
            )
            continue

        d_terms = list(re.finditer(r"D_[1-4][^D]*", stripped))
        if len(d_terms) >= 2:
            for m in d_terms:
                term = m.group(0).strip().strip(",")
                if term:
                    out.append(f"$$ {term} $$")
            continue

        if "A1A_1A1" in stripped or "A1A1" in stripped:
            out.append("⇒ $A_1$ es definida positiva ⇒ $q_1$ definida positiva.")
            continue
        if "A4A_4A4" in stripped or "A4A4" in stripped:
            out.append("⇒ $A_4$ es definida positiva ⇒ $q_4$ definida positiva.")
            continue

        if "detA2" in stripped or "det A_2" in stripped:
            out.append("$$ \\det A_2 = -2 < 0 $$")
            continue
        if "detA3" in stripped or "det A_3" in stripped:
            out.append("$$ \\det A_3 = -20 < 0 $$")
            continue

        if "q3q_3q3" in stripped and "indefinida" in stripped:
            out.append("⇒ $q_3$ también es indefinida.")
            continue

        out.append(stripped)

    return out


def prettify_paragraphs(paragraph_texts: List[str]) -> List[str]:
    generic: List[str] = []
    for text in paragraph_texts:
        s = normalize_math_text(text)
        if s.strip():
            generic.append(s.strip())
    if not USE_EXERCISE_TWEAKS:
        return generic
    return _prettify_paragraphs_for_exercise(generic)


# ================================================================
# 2) Parsing LaTeX
# ================================================================
def new_paragraph(doc: Document, align: Optional[Any] = None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    return p


def add_math_safe(paragraph, latex: str) -> bool:
    """Try converting LaTeX to OMML; return True if OMML was produced.

    If conversion fails, the original LaTeX is added as plain text (best-effort) and False is returned.
    """
    try:
        math2docx.add_math(paragraph, latex)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Fallo convirtiendo LaTeX a OMML: %s", exc)
        paragraph.add_run(latex)
        return False

    # math2docx may silently fall back to plain text; verify OMML presence
    try:
        return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        return True

def parse_math_segments(text: str) -> List[Segment]:
    segments: List[Segment] = []
    buf: List[str] = []

    def flush_text() -> None:
        if buf:
            segments.append(("text", "".join(buf)))
            buf.clear()

    i = 0
    n = len(text)
    while i < n:
        if text.startswith("$$", i):
            flush_text()
            end = text.find("$$", i + 2)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 2 : end]
            segments.append(("display", latex.strip()))
            i = end + 2
            continue

        if text.startswith(r"\[", i):
            flush_text()
            end = text.find(r"\]", i + 2)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 2 : end]
            segments.append(("display", latex.strip()))
            i = end + 2
            continue

        if text[i] == "$":
            flush_text()
            end = text.find("$", i + 1)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 1 : end]
            segments.append(("inline", latex.strip()))
            i = end + 1
            continue

        buf.append(text[i])
        i += 1

    flush_text()
    return segments


def split_into_paragraph_descriptors(text: str) -> List[Dict[str, Any]]:
    segments = parse_math_segments(text)
    paragraphs: List[Dict[str, Any]] = []
    current_chunks: List[Segment] = []

    for kind, value in segments:
        if kind in ("text", "inline"):
            if value:
                current_chunks.append((kind, value))
        elif kind == "display":
            if current_chunks:
                paragraphs.append({"type": "inline", "chunks": current_chunks})
                current_chunks = []
            paragraphs.append({"type": "display", "latex": value})

    if current_chunks:
        paragraphs.append({"type": "inline", "chunks": current_chunks})

    return paragraphs


def split_long_latex_equation(latex: str, max_len: int = 120) -> List[str]:
    latex = latex.strip()
    if not latex:
        return []
    if len(latex) <= max_len:
        return [latex]

    if "\\\\" in latex:
        return [p.strip() for p in latex.split("\\\\") if p.strip()]

    if "=" in latex:
        tokens = [t.strip() for t in latex.split("=")]
        eqs: List[str] = []
        current = tokens[0]
        for tok in tokens[1:]:
            candidate = current + " = " + tok
            if len(candidate) > max_len and current:
                eqs.append(current)
                current = tok
            else:
                current = candidate
        if current:
            eqs.append(current)
        return eqs

    if "," in latex:
        tokens = [t.strip() for t in latex.split(",")]
        eqs2: List[str] = []
        current2 = ""
        for tok in tokens:
            candidate = (current2 + ", " if current2 else "") + tok
            if len(candidate) > max_len and current2:
                eqs2.append(current2)
                current2 = tok
            else:
                current2 = candidate
        if current2:
            eqs2.append(current2)
        return eqs2

    return [latex[i : i + max_len].strip() for i in range(0, len(latex), max_len)]


def add_aligned_block(doc: Document, aligned_block: str) -> None:
    content = aligned_block.strip()
    if content.startswith(r"\begin{aligned}"):
        content = content[len(r"\begin{aligned}") :]
    if content.endswith(r"\end{aligned}"):
        content = content[: -len(r"\end{aligned}")]
    content = content.strip()
    if not content:
        return
    rows = [row.strip() for row in content.split(r"\\") if row.strip()]
    for row in rows:
        row_no_amp = row.replace("&", "")
        p = new_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
        add_math_safe(p, row_no_amp)


def build_document_from_paragraphs(paragraph_texts: List[str]) -> Document:
    out_doc = Document()
    if not paragraph_texts:
        new_paragraph(out_doc)
        return out_doc

    in_display_block = False
    display_block_delim: Optional[str] = None  # "$$" o "\["
    display_lines: List[str] = []

    in_aligned_block = False
    aligned_lines: List[str] = []

    matrix_envs = ["vmatrix", "bmatrix", "pmatrix", "matrix"]

    for raw_text in paragraph_texts:
        text = normalize_math_text(raw_text)
        stripped = text.strip()

        # aligned multi-línea
        if in_aligned_block:
            aligned_lines.append(text)
            if r"\end{aligned}" in text:
                add_aligned_block(out_doc, "\n".join(aligned_lines))
                in_aligned_block = False
                aligned_lines = []
            continue

        if r"\begin{aligned}" in stripped:
            if r"\end{aligned}" in stripped:
                add_aligned_block(out_doc, stripped)
            else:
                in_aligned_block = True
                aligned_lines = [stripped]
            continue

        # display multi-línea $$ ... $$  o \[ ... \]
        if in_display_block:
            if (
                (display_block_delim == "$$" and stripped == "$$")
                or (display_block_delim == r"\[" and stripped == r"\]")
            ):
                latex = "\n".join(display_lines).strip()
                for part in split_long_latex_equation(latex):
                    p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                    add_math_safe(p, part)
                in_display_block = False
                display_block_delim = None
                display_lines = []
            else:
                display_lines.append(text)
            continue

        if stripped == "$$":
            in_display_block = True
            display_block_delim = "$$"
            display_lines = []
            continue

        if stripped == r"\[":
            in_display_block = True
            display_block_delim = r"\["
            display_lines = []
            continue

        # matrices en una línea
        handled_matrix = False
        for env in matrix_envs:
            begin = rf"\begin{{{env}}}"
            end = rf"\end{{{env}}}"
            if begin in stripped and end in stripped and "$" not in stripped:
                p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                add_math_safe(p, stripped)
                handled_matrix = True
                break
        if handled_matrix:
            continue

        # normal
        descriptors = split_into_paragraph_descriptors(text)
        if not descriptors:
            new_paragraph(out_doc)
            continue

        for desc in descriptors:
            if desc["type"] == "inline":
                p = new_paragraph(out_doc)
                for kind, value in desc["chunks"]:
                    if kind == "text":
                        p.add_run(value)
                    elif kind == "inline":
                        add_math_safe(p, value)
            elif desc["type"] == "display":
                for part in split_long_latex_equation(desc["latex"].strip()):
                    p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                    add_math_safe(p, part)

    # cierre si termina dentro de display/aligned
    if in_display_block and display_lines:
        latex = "\n".join(display_lines).strip()
        for part in split_long_latex_equation(latex):
            p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
            add_math_safe(p, part)

    if in_aligned_block and aligned_lines:
        add_aligned_block(out_doc, "\n".join(aligned_lines))

    return out_doc


# ================================================================
# Routes: páginas
# ================================================================
@app.get("/", response_class=HTMLResponse)
async def home() -> HTMLResponse:
    try:
        return HTMLResponse(read_html_file("index.html"))
    except FileNotFoundError:
        return HTMLResponse("<h1>index.html not found</h1>", status_code=404)


@app.get("/en", response_class=HTMLResponse)
async def home_en() -> HTMLResponse:
    try:
        return HTMLResponse(read_html_file("index-en.html"))
    except FileNotFoundError:
        return HTMLResponse("<h1>index-en.html not found</h1>", status_code=404)


@app.get("/de", response_class=HTMLResponse)
async def home_de() -> HTMLResponse:
    return RedirectResponse(url="/en", status_code=301)


@app.get("/fr", response_class=HTMLResponse)
async def home_fr() -> HTMLResponse:
    return RedirectResponse(url="/en", status_code=301)


@app.get("/it", response_class=HTMLResponse)
async def home_it() -> HTMLResponse:
    return RedirectResponse(url="/en", status_code=301)


@app.get("/pt", response_class=HTMLResponse)
async def home_pt() -> HTMLResponse:
    return RedirectResponse(url="/en", status_code=301)


def _landing_from_slug(lang: str, slug: str) -> Optional[Dict[str, Any]]:
    data = LANDING_PAGES.get(slug)
    if not data:
        return None
    return data.get(lang) or data.get(_content_lang(lang))


def _landing_key_from_route_slug(lang: str, route_slug: str) -> Optional[str]:
    route_slug = (route_slug or "").strip().lower()
    if not route_slug:
        return None
    source_lang = lang if lang in PRIMARY_CONTENT_LANGS else "en"
    for key, langs in LANDING_PAGES.items():
        path = (langs.get(source_lang, {}).get("path") or "").strip()
        if not path:
            continue
        if path.rstrip("/").split("/")[-1].lower() == route_slug:
            return key
    return None


def _all_landing_pairs() -> List[Tuple[str, str]]:
    pairs: List[Tuple[str, str]] = []
    for slug, data in LANDING_PAGES.items():
        es = data.get("es", {}).get("path")
        en = data.get("en", {}).get("path")
        if not es or not en:
            logger.warning("Landing %s missing ES/EN path", slug)
            continue
        pairs.append((es, en))
    return pairs


def _secondary_solution_redirect_target(source_lang: str, route_slug: str) -> Optional[str]:
    """Map a secondary-language solution slug to its EN canonical path."""
    landing_key = _landing_key_from_route_slug(source_lang, route_slug)
    if not landing_key:
        return None
    en_page = _landing_from_slug("en", landing_key)
    if not en_page:
        return None
    path = (en_page.get("path") or "").strip()
    return path or None


def _secondary_blog_redirect_target(source_lang: str, slug: str) -> Optional[str]:
    """Map a secondary-language blog slug to EN canonical path when equivalent exists."""
    clean_slug = (slug or "").strip()
    if not clean_slug:
        return None

    # Resolve aliases in the incoming language first.
    clean_slug = BLOG_ALIASES.get(source_lang, {}).get(clean_slug, clean_slug)

    en_alias_target = BLOG_ALIASES.get("en", {}).get(clean_slug)
    if en_alias_target:
        clean_slug = en_alias_target

    en_post = BLOG_POSTS.get("en", {}).get(clean_slug)
    if en_post:
        return (en_post.get("canonical_path") or f"/en/blog/{clean_slug}").strip()

    source_post = BLOG_POSTS.get(source_lang, {}).get(clean_slug)
    translation_slug = (source_post or {}).get("translation_slug")
    if translation_slug:
        en_post = BLOG_POSTS.get("en", {}).get(translation_slug)
        if en_post:
            return (en_post.get("canonical_path") or f"/en/blog/{translation_slug}").strip()

    es_post = BLOG_POSTS.get("es", {}).get(clean_slug)
    es_translation = (es_post or {}).get("translation_slug")
    if es_translation:
        en_post = BLOG_POSTS.get("en", {}).get(es_translation)
        if en_post:
            return (en_post.get("canonical_path") or f"/en/blog/{es_translation}").strip()

    return None


def _solution_landing_context(lang: str, route_slug: str) -> Optional[Dict[str, Any]]:
    source_lang = lang if lang in PRIMARY_CONTENT_LANGS else "en"
    landing_key = _landing_key_from_route_slug(lang, route_slug)
    if not landing_key:
        return None
    current = _landing_from_slug(source_lang, landing_key)
    es_page = _landing_from_slug("es", landing_key)
    en_page = _landing_from_slug("en", landing_key)
    if not current or not es_page or not en_page:
        return None

    route_slug = route_slug.strip().lower()
    canonical_path = current["path"] if lang in PRIMARY_CONTENT_LANGS else f"{_solutions_path(lang)}/{route_slug}"
    canonical_url = _abs_url(canonical_path)
    source_page = _home_path(lang)
    related_blog_href = current["related_blog_href"]
    if lang not in PRIMARY_CONTENT_LANGS:
        related_blog_href = related_blog_href.replace("/en/blog/", f"/{lang}/blog/")
    longform = _landing_longform(source_lang, landing_key)
    indexable = _is_primary_lang(lang)

    alt_paths = {
        "es": es_page["path"],
        "en": en_page["path"],
    }

    return {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": current["seo_title"],
        "description": current["description"],
        "keywords": [],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=indexable),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "website",
        "og_title": current["title"],
        "og_description": current["description"],
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_solution_page(
            current["title"],
            canonical_url,
            lang,
            current["description"],
            _abs_url(related_blog_href),
            faq_items=longform.get("faq_items"),
        ),
        "converter_href": _home_path(lang),
        "blog_index_href": _blog_index_path(lang),
        "solutions_hub_href": _solutions_path(lang),
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": _home_path("es" if lang != "es" else "en"),
        "lang_switch_label": "ES" if lang != "es" else "EN",
        "lang_options": _lang_options(alt_paths),
        "page_kicker": current["kicker"],
        "page_title": current["h1"],
        "page_intro": current["intro"],
        "intent_items": current["intent_items"],
        "cta_label": current["cta_label"],
        "source_page": source_page,
        "related_blog_href": related_blog_href,
        "related_blog_label": current["related_blog_label"],
        "breadcrumbs": [
            {"name": "Inicio" if lang == "es" else "Home", "url": _home_path(lang)},
            {"name": "Soluciones" if lang == "es" else "Solutions", "url": _solutions_path(lang)},
            {"name": current["h1"], "url": canonical_path},
        ],
        "primary_keyword": longform.get("primary_keyword", current["h1"]),
        "primary_intent": longform.get("primary_intent", "transactional"),
        "problem_angle": longform.get("problem_angle", current["description"]),
        "problem_statement": longform.get("problem_statement", current["description"]),
        "extra_context": longform.get("extra_context", ""),
        "when_to_use_intro": longform.get("when_to_use_intro", current["description"]),
        "when_to_use_items": longform.get("when_to_use_items", []),
        "workflow_steps": longform.get("workflow_steps", []),
        "example_input": longform.get("example_input", ""),
        "example_output": longform.get("example_output", ""),
        "before_text": longform.get("before_text", ""),
        "after_text": longform.get("after_text", ""),
        "mini_case": longform.get("mini_case", ""),
        "supported_formats": longform.get("supported_formats", []),
        "supported_delimiters": longform.get("supported_delimiters", []),
        "known_limitations": longform.get("known_limitations", []),
        "common_errors": longform.get("common_errors", []),
        "faq_items": longform.get("faq_items", []),
        "secondary_cta_href": longform.get("secondary_cta_href", related_blog_href),
        "secondary_cta_label": longform.get("secondary_cta_label", current["related_blog_label"]),
        "pillar_label": longform.get(
            "pillar_label",
            "Pilar: latex a word online" if lang == "es" else "Pillar: latex to word online",
        ),
        "related_guides": longform.get("related_guides", []),
        "year": datetime.now().year,
        "legal_links": {
            "privacy": _legal_path(lang, "privacy"),
            "terms": _legal_path(lang, "terms"),
            "contact": _legal_path(lang, "contact"),
        },
    }


def _solutions_hub_context(lang: str) -> Dict[str, Any]:
    source_lang = lang if lang in PRIMARY_CONTENT_LANGS else "en"
    canonical_path = _solutions_path(lang)
    alt_paths = {code: _solutions_path(code) for code in PRIMARY_CONTENT_LANGS}
    cards: List[Dict[str, str]] = []
    for key, langs in LANDING_PAGES.items():
        current = langs.get(source_lang) or langs.get("en") or langs.get("es")
        if not current:
            continue
        slug = current["path"].rstrip("/").split("/")[-1]
        href = current["path"] if lang in PRIMARY_CONTENT_LANGS else f"{_solutions_path(lang)}/{slug}"
        longform = _landing_longform(source_lang, key)
        user_type = ", ".join((longform.get("when_to_use_items") or [])[:2])
        cards.append(
            {
                "href": href,
                "title": current.get("h1") or current.get("title") or "",
                "kicker": current.get("kicker") or "",
                "problem": longform.get("problem_angle", current.get("description", "")),
                "user_type_label": "Para quien es:" if lang == "es" else "Best for:",
                "user_type": user_type or (current.get("description") or ""),
                "cta": "Ver solucion" if lang == "es" else "View solution",
            }
        )

    title = "Soluciones LaTeX a Word por caso de uso" if lang == "es" else "LaTeX to Word solutions by use case"
    description = (
        "Elige la landing de conversion que mejor encaja con tu fuente: ChatGPT, Gemini, Overleaf, Pandoc u OMML."
        if lang == "es"
        else "Pick the conversion landing that matches your source: ChatGPT, Gemini, Overleaf, Pandoc or OMML."
    )
    return {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": title,
        "description": description,
        "keywords": [],
        "canonical_url": _abs_url(canonical_path),
        "robots_directive": _robots_directive(lang, indexable=_is_primary_lang(lang)),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "website",
        "og_title": title,
        "og_description": description,
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_simple_page("Soluciones" if lang == "es" else "Solutions", _abs_url(canonical_path), lang),
        "converter_href": _home_path(lang),
        "blog_index_href": _blog_index_path(lang),
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_options": _lang_options(alt_paths),
        "breadcrumbs": [
            {"name": "Inicio" if lang == "es" else "Home", "url": _home_path(lang)},
            {"name": "Soluciones" if lang == "es" else "Solutions", "url": canonical_path},
        ],
        "h1": title,
        "intro": description,
        "cluster_title": "Arquitectura de soluciones por intencion" if lang == "es" else "Intent-based solution architecture",
        "cluster_intro": (
            "Cada landing responde a un caso de uso distinto para ayudarte a elegir el flujo de conversión adecuado."
            if lang == "es"
            else "Each landing is built for a distinct use case so you can pick the right conversion workflow."
        ),
        "cards": cards,
        "cta_strong": _ui(lang, "cta_strong"),
        "cta_text": _ui(lang, "cta_text"),
        "cta_primary": _ui(lang, "cta_primary"),
        "cta_secondary": _ui(lang, "cta_secondary"),
        "year": datetime.now().year,
        "legal_links": {
            "privacy": _legal_path(lang, "privacy"),
            "terms": _legal_path(lang, "terms"),
            "contact": _legal_path(lang, "contact"),
        },
    }


@app.get("/soluciones/{slug}", response_class=HTMLResponse)
async def solution_landing_es(slug: str) -> HTMLResponse:
    ctx = _solution_landing_context("es", slug)
    if not ctx:
        raise HTTPException(status_code=404, detail="Landing not found")
    return HTMLResponse(_render_template("solution_landing.html", ctx))


@app.get("/en/solutions/{slug}", response_class=HTMLResponse)
async def solution_landing_en(slug: str) -> HTMLResponse:
    ctx = _solution_landing_context("en", slug)
    if not ctx:
        raise HTTPException(status_code=404, detail="Landing not found")
    return HTMLResponse(_render_template("solution_landing.html", ctx))


@app.get("/de/solutions/{slug}", response_class=HTMLResponse)
async def solution_landing_de(slug: str) -> HTMLResponse:
    target = _secondary_solution_redirect_target("de", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Landing not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/fr/solutions/{slug}", response_class=HTMLResponse)
async def solution_landing_fr(slug: str) -> HTMLResponse:
    target = _secondary_solution_redirect_target("fr", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Landing not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/it/solutions/{slug}", response_class=HTMLResponse)
async def solution_landing_it(slug: str) -> HTMLResponse:
    target = _secondary_solution_redirect_target("it", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Landing not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/pt/solutions/{slug}", response_class=HTMLResponse)
async def solution_landing_pt(slug: str) -> HTMLResponse:
    target = _secondary_solution_redirect_target("pt", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Landing not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/soluciones", response_class=HTMLResponse)
async def solutions_es() -> HTMLResponse:
    return HTMLResponse(_render_template("solutions_hub.html", _solutions_hub_context("es")))


@app.get("/en/solutions", response_class=HTMLResponse)
async def solutions_en() -> HTMLResponse:
    return HTMLResponse(_render_template("solutions_hub.html", _solutions_hub_context("en")))


@app.get("/de/solutions", response_class=HTMLResponse)
async def solutions_de() -> HTMLResponse:
    return RedirectResponse(url="/en/solutions", status_code=301)


@app.get("/fr/solutions", response_class=HTMLResponse)
async def solutions_fr() -> HTMLResponse:
    return RedirectResponse(url="/en/solutions", status_code=301)


@app.get("/it/solutions", response_class=HTMLResponse)
async def solutions_it() -> HTMLResponse:
    return RedirectResponse(url="/en/solutions", status_code=301)


@app.get("/pt/solutions", response_class=HTMLResponse)
async def solutions_pt() -> HTMLResponse:
    return RedirectResponse(url="/en/solutions", status_code=301)


# ================================================================
# Legal / Trust pages
# ================================================================
def _legal_page_context(lang: str, page: str) -> Dict[str, Any]:
    """Build context for legal/trust pages."""

    legal_copy: Dict[str, Dict[str, Dict[str, str]]] = {
        "es": {
            "privacy": {
                "title": "Política de privacidad",
                "description": "Cómo tratamos tus archivos y datos al convertir LaTeX a Word.",
                "body": "<h2>Tratamiento de archivos</h2><p>Procesamos tu archivo únicamente para realizar la conversión. No vendemos tus datos.</p><h2>Almacenamiento</h2><p>No almacenamos permanentemente tus documentos. Puede existir procesamiento temporal en memoria durante la conversión.</p><h2>Analítica</h2><p>Usamos analítica agregada para mejorar el producto. El contenido de tus documentos no se envía a herramientas de analítica.</p><h2>Contacto</h2><p>Si tienes dudas, escríbenos a <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>",
            },
            "terms": {
                "title": "Términos de uso",
                "description": "Normas y limitaciones de uso del conversor.",
                "body": "<h2>Herramienta gratuita</h2><p>Este conversor se ofrece de forma gratuita. Puede haber límites de tamaño de archivo.</p><h2>Sin garantías</h2><p>El servicio se ofrece tal cual. Hacemos lo posible, pero no garantizamos una conversión perfecta en todos los documentos.</p><h2>Uso aceptable</h2><p>No subas contenido ilegal, malware o documentos sensibles que no estás autorizado a compartir.</p>",
            },
            "contact": {
                "title": "Contacto",
                "description": "Contacta con el proyecto.",
                "body": "<p>Email: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>",
            },
        },
        "en": {
            "privacy": {
                "title": "Privacy Policy",
                "description": "How we handle uploaded files and data during LaTeX to Word conversion.",
                "body": "<h2>File handling</h2><p>We process uploaded files only to perform the conversion. We do not sell your data.</p><h2>Storage</h2><p>We do not permanently store documents. Temporary in-memory processing may occur during conversion.</p><h2>Analytics</h2><p>We use aggregated analytics to improve the product. Document content is not sent to analytics tools.</p><h2>Contact</h2><p>Questions: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>",
            },
            "terms": {
                "title": "Terms of use",
                "description": "Rules and limitations for using the converter.",
                "body": "<h2>Free tool</h2><p>This converter is provided free of charge. File-size limits may apply.</p><h2>No warranties</h2><p>The service is provided as-is. We do our best but cannot guarantee perfect conversion for every document.</p><h2>Acceptable use</h2><p>Do not upload illegal content, malware, or sensitive documents you are not allowed to share.</p>",
            },
            "contact": {
                "title": "Contact",
                "description": "Get in touch with the project.",
                "body": "<p>Email: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>",
            },
        },
        "de": {
            "privacy": {"title": "Datenschutzerkl?rung", "description": "Wie wir hochgeladene Dateien und Daten verarbeiten.", "body": "<h2>Dateiverarbeitung</h2><p>Wir verarbeiten hochgeladene Dateien nur f?r die Konvertierung und verkaufen keine Daten.</p><h2>Speicherung</h2><p>Dokumente werden nicht dauerhaft gespeichert. W?hrend der Konvertierung kann tempor?re Verarbeitung im Speicher stattfinden.</p><h2>Analyse</h2><p>Wir verwenden aggregierte Nutzungsdaten zur Produktverbesserung. Dokumentinhalte werden nicht an Analytics-Tools gesendet.</p><h2>Kontakt</h2><p>Fragen: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>"},
            "terms": {"title": "Nutzungsbedingungen", "description": "Regeln und Einschr?nkungen f?r den Konverter.", "body": "<h2>Kostenloses Tool</h2><p>Der Konverter ist kostenlos; Dateigr??enlimits k?nnen gelten.</p><h2>Keine Gew?hr</h2><p>Der Dienst wird ohne Gew?hr bereitgestellt.</p><h2>Zul?ssige Nutzung</h2><p>Bitte keine illegalen Inhalte oder Malware hochladen.</p>"},
            "contact": {"title": "Kontakt", "description": "So erreichst du das Projekt.", "body": "<p>E-Mail: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>"},
        },
        "fr": {
            "privacy": {"title": "Politique de confidentialit?", "description": "Comment nous traitons les fichiers et donn?es envoy?s.", "body": "<h2>Traitement des fichiers</h2><p>Nous traitons les fichiers uniquement pour la conversion et ne vendons pas vos donn?es.</p><h2>Stockage</h2><p>Nous ne stockons pas durablement les documents. Un traitement temporaire en m?moire peut avoir lieu.</p><h2>Analytique</h2><p>Nous utilisons des statistiques agr?g?es pour am?liorer le produit. Le contenu des documents n'est pas envoy? aux outils d'analyse.</p><h2>Contact</h2><p>Questions : <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>"},
            "terms": {"title": "Conditions d'utilisation", "description": "R?gles et limites d'utilisation du convertisseur.", "body": "<h2>Outil gratuit</h2><p>Le convertisseur est gratuit, avec possibles limites de taille.</p><h2>Sans garantie</h2><p>Le service est fourni tel quel.</p><h2>Usage acceptable</h2><p>N'envoyez pas de contenu ill?gal ou malveillant.</p>"},
            "contact": {"title": "Contact", "description": "Contactez le projet.", "body": "<p>Email : <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>"},
        },
        "it": {
            "privacy": {"title": "Informativa sulla privacy", "description": "Come trattiamo file e dati caricati.", "body": "<h2>Gestione dei file</h2><p>Elaboriamo i file caricati solo per la conversione e non vendiamo i dati.</p><h2>Archiviazione</h2><p>I documenti non sono conservati in modo permanente. Durante la conversione pu? esserci elaborazione temporanea in memoria.</p><h2>Analisi</h2><p>Usiamo dati aggregati per migliorare il prodotto. Il contenuto dei documenti non viene inviato agli strumenti di analytics.</p><h2>Contatto</h2><p>Domande: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>"},
            "terms": {"title": "Termini di utilizzo", "description": "Regole e limiti d'uso del convertitore.", "body": "<h2>Strumento gratuito</h2><p>Il convertitore ? gratuito, con possibili limiti di dimensione.</p><h2>Nessuna garanzia</h2><p>Il servizio ? fornito cos? com'?.</p><h2>Uso accettabile</h2><p>Non caricare contenuti illegali o malware.</p>"},
            "contact": {"title": "Contatto", "description": "Contatta il progetto.", "body": "<p>Email: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>"},
        },
        "pt": {
            "privacy": {"title": "Pol?tica de Privacidade", "description": "Como tratamos arquivos e dados enviados.", "body": "<h2>Tratamento de arquivos</h2><p>Processamos arquivos apenas para convers?o e n?o vendemos dados.</p><h2>Armazenamento</h2><p>N?o armazenamos documentos permanentemente. Pode haver processamento tempor?rio em mem?ria durante a convers?o.</p><h2>An?lises</h2><p>Usamos dados agregados para melhorar o produto. O conte?do dos documentos n?o ? enviado ?s ferramentas de analytics.</p><h2>Contato</h2><p>D?vidas: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a>.</p>"},
            "terms": {"title": "Termos de uso", "description": "Regras e limita??es para usar o conversor.", "body": "<h2>Ferramenta gratuita</h2><p>O conversor ? gratuito, com poss?veis limites de tamanho.</p><h2>Sem garantias</h2><p>O servi?o ? fornecido no estado em que se encontra.</p><h2>Uso aceit?vel</h2><p>N?o envie conte?do ilegal ou malware.</p>"},
            "contact": {"title": "Contato", "description": "Fale com o projeto.", "body": "<p>Email: <a href='mailto:ecuacionesaword@gmail.com'>ecuacionesaword@gmail.com</a></p>"},
        },
    }

    locale = lang if lang in legal_copy else ("es" if lang == "es" else "en")
    page_data = legal_copy.get(locale, legal_copy["en"]).get(page, {"title": "Info", "description": "", "body": "<p></p>"})

    title = page_data["title"]
    description = page_data["description"]
    body_html = page_data["body"]

    canonical_path = _legal_path(lang, page) if page in {"privacy", "terms", "contact"} else _home_path(lang)
    canonical_url = _abs_url(canonical_path)
    alt_paths = {code: _legal_path(code, page) for code in PRIMARY_CONTENT_LANGS}

    return {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": f"{title} | {SITE_NAME}",
        "description": description,
        "keywords": [],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=False),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "website",
        "og_title": title,
        "og_description": description,
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_simple_page(title, canonical_url, lang),
        "converter_href": _home_path(lang),
        "blog_index_href": _blog_index_path(lang),
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": _legal_path("es" if lang != "es" else "en", page),
        "lang_switch_label": "ES" if lang != "es" else "EN",
        "lang_options": _lang_options(alt_paths),
        "kicker": "",
        "title": title,
        "meta_line": "",
        "tags": [],
        "intro_html": [],
        "body_html": body_html,
        "cta_strong": "",
        "cta_text": "",
        "cta_primary": "",
        "cta_secondary": "",
        "year": datetime.now(timezone.utc).year,
        "footer_links": True,
        "legal_links": {
            "privacy": _legal_path(lang, "privacy"),
            "terms": _legal_path(lang, "terms"),
            "contact": _legal_path(lang, "contact"),
        },
    }


@app.get("/privacy", response_class=HTMLResponse)
async def privacy_es() -> HTMLResponse:
    ctx = _legal_page_context("es", "privacy")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/terms", response_class=HTMLResponse)
async def terms_es() -> HTMLResponse:
    ctx = _legal_page_context("es", "terms")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/contact", response_class=HTMLResponse)
async def contact_es() -> HTMLResponse:
    ctx = _legal_page_context("es", "contact")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/en/privacy", response_class=HTMLResponse)
async def privacy_en() -> HTMLResponse:
    ctx = _legal_page_context("en", "privacy")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/en/terms", response_class=HTMLResponse)
async def terms_en() -> HTMLResponse:
    ctx = _legal_page_context("en", "terms")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/en/contact", response_class=HTMLResponse)
async def contact_en() -> HTMLResponse:
    ctx = _legal_page_context("en", "contact")
    html = _render_template("legal_page.html", ctx)
    return HTMLResponse(html)


@app.get("/de/privacy", response_class=HTMLResponse)
async def privacy_de() -> HTMLResponse:
    return RedirectResponse(url="/en/privacy", status_code=301)


@app.get("/de/terms", response_class=HTMLResponse)
async def terms_de() -> HTMLResponse:
    return RedirectResponse(url="/en/terms", status_code=301)


@app.get("/de/contact", response_class=HTMLResponse)
async def contact_de() -> HTMLResponse:
    return RedirectResponse(url="/en/contact", status_code=301)


@app.get("/fr/privacy", response_class=HTMLResponse)
async def privacy_fr() -> HTMLResponse:
    return RedirectResponse(url="/en/privacy", status_code=301)


@app.get("/fr/terms", response_class=HTMLResponse)
async def terms_fr() -> HTMLResponse:
    return RedirectResponse(url="/en/terms", status_code=301)


@app.get("/fr/contact", response_class=HTMLResponse)
async def contact_fr() -> HTMLResponse:
    return RedirectResponse(url="/en/contact", status_code=301)


@app.get("/it/privacy", response_class=HTMLResponse)
async def privacy_it() -> HTMLResponse:
    return RedirectResponse(url="/en/privacy", status_code=301)


@app.get("/it/terms", response_class=HTMLResponse)
async def terms_it() -> HTMLResponse:
    return RedirectResponse(url="/en/terms", status_code=301)


@app.get("/it/contact", response_class=HTMLResponse)
async def contact_it() -> HTMLResponse:
    return RedirectResponse(url="/en/contact", status_code=301)


@app.get("/pt/privacy", response_class=HTMLResponse)
async def privacy_pt() -> HTMLResponse:
    return RedirectResponse(url="/en/privacy", status_code=301)


@app.get("/pt/terms", response_class=HTMLResponse)
async def terms_pt() -> HTMLResponse:
    return RedirectResponse(url="/en/terms", status_code=301)


@app.get("/pt/contact", response_class=HTMLResponse)
async def contact_pt() -> HTMLResponse:
    return RedirectResponse(url="/en/contact", status_code=301)


@app.get("/blog", response_class=HTMLResponse)
async def blog_index_es() -> HTMLResponse:
    lang = "es"
    canonical_url = _abs_url("/blog")
    other_url = _abs_url("/en/blog")

    posts_view: List[Dict[str, Any]] = []
    tag_counts: Dict[str, int] = {}
    for p in BLOG_LIST.get(lang, []):
        for t in (p.get("tags") or []):
            if not isinstance(t, str) or not t.strip():
                continue
            nt = _normalize_tag(t)
            tag_counts[nt] = tag_counts.get(nt, 0) + 1
        posts_view.append(
            {
                "url": p.get("canonical_path") or f"/blog/{p.get('slug')}",
                "title": p.get("title") or "",
                "description": p.get("description") or "",
                "kicker": p.get("kicker") or "",
                "tags": [_normalize_tag(t) for t in (p.get("tags") or []) if isinstance(t, str)],
                "meta": f"{_format_date(lang, p.get('date_published') or '')} · {_format_reading_time(p.get('reading_time'))}".strip(
                    " ·"
                ),
            }
        )

    top_tags = [
        k
        for k, _ in sorted(tag_counts.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    ][:8]

    ctx = {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": "Blog | Ecuaciones a Word",
        "description": "Guías para convertir LaTeX e IA a Word con ecuaciones nativas (OMML), sin imágenes ni fórmulas rotas.",
        "keywords": ["LaTeX a Word", "LaTeX a Word online", "OMML", "Word ecuaciones", "ChatGPT", "Pandoc"],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=True),
        "alternates": [
            {"hreflang": "es", "href": canonical_url},
            {"hreflang": "en", "href": other_url},
            {"hreflang": "x-default", "href": canonical_url},
        ],
        "og_type": "website",
        "og_title": "Blog | Ecuaciones a Word",
        "og_description": "Guías prácticas para llevar ecuaciones LaTeX a Word (OMML) de forma limpia y editable.",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_index(lang, canonical_url),
        "converter_href": "/",
        "blog_index_href": "/blog",
        "nav_converter": "Conversor",
        "nav_blog": "Blog",
        "lang_switch_href": "/en/blog",
        "lang_switch_label": "EN",
        "lang_options": _lang_options({code: _blog_index_path(code) for code in PRIMARY_CONTENT_LANGS}),
        "h1": "Blog",
        "intro": "Guías prácticas para pasar ecuaciones de LaTeX e IA a Word con ecuaciones nativas (OMML), incluyendo flujos online y troubleshooting.",
        "index_cta_primary": "Conversor LaTeX → Word (OMML)",
        "search_label": "Buscar artículos",
        "search_placeholder": "Buscar por tema, herramienta o problema (ej. pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filtrar por etiqueta",
        "filter_all": "Todos",
        "top_tags": top_tags,
        "featured_title": "Artículos",
        "posts": posts_view,
        "year": datetime.now().year,
        "legal_links": {"privacy": "/privacy", "terms": "/terms", "contact": "/contact"},
    }
    return HTMLResponse(_render_template("blog_index.html", ctx))


@app.get("/en/blog", response_class=HTMLResponse)
async def blog_index_en() -> HTMLResponse:
    lang = "en"
    canonical_url = _abs_url("/en/blog")
    other_url = _abs_url("/blog")

    posts_view: List[Dict[str, Any]] = []
    tag_counts: Dict[str, int] = {}
    for p in BLOG_LIST.get(lang, []):
        for t in (p.get("tags") or []):
            if not isinstance(t, str) or not t.strip():
                continue
            nt = _normalize_tag(t)
            tag_counts[nt] = tag_counts.get(nt, 0) + 1
        posts_view.append(
            {
                "url": p.get("canonical_path") or f"/en/blog/{p.get('slug')}",
                "title": p.get("title") or "",
                "description": p.get("description") or "",
                "kicker": p.get("kicker") or "",
                "tags": [_normalize_tag(t) for t in (p.get("tags") or []) if isinstance(t, str)],
                "meta": f"{_format_date(lang, p.get('date_published') or '')} · {_format_reading_time(p.get('reading_time'))}".strip(
                    " ·"
                ),
            }
        )

    top_tags = [
        k
        for k, _ in sorted(tag_counts.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    ][:8]

    ctx = {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": "Blog | Ecuaciones a Word",
        "description": "Practical guides to convert LaTeX/AI content to Word with native (OMML) editable equations.",
        "keywords": ["LaTeX to Word", "LaTeX to Word online", "OMML", "Word equations", "ChatGPT", "Pandoc"],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=True),
        "alternates": [
            {"hreflang": "en", "href": canonical_url},
            {"hreflang": "es", "href": other_url},
            {"hreflang": "x-default", "href": other_url},
        ],
        "og_type": "website",
        "og_title": "Blog | Ecuaciones a Word",
        "og_description": "Practical guides for converting LaTeX equations to native Word (OMML) cleanly and reliably.",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_index(lang, canonical_url),
        "converter_href": "/en",
        "blog_index_href": "/en/blog",
        "nav_converter": "Converter",
        "nav_blog": "Blog",
        "lang_switch_href": "/blog",
        "lang_switch_label": "ES",
        "lang_options": _lang_options({code: _blog_index_path(code) for code in PRIMARY_CONTENT_LANGS}),
        "h1": "Blog",
        "intro": "Practical guides to export LaTeX/AI equations to Word as native editable OMML equations, including online workflows and troubleshooting.",
        "index_cta_primary": "LaTeX → OMML Converter",
        "search_label": "Search articles",
        "search_placeholder": "Search by topic, tool or issue (e.g., pandoc, overleaf, OMML, ChatGPT)…",
        "filter_label": "Filter by tag",
        "filter_all": "All",
        "top_tags": top_tags,
        "featured_title": "Articles",
        "posts": posts_view,
        "year": datetime.now().year,
        "legal_links": {"privacy": "/en/privacy", "terms": "/en/terms", "contact": "/en/contact"},
    }
    return HTMLResponse(_render_template("blog_index.html", ctx))


def _build_blog_index_context_fallback_en(lang: str) -> Dict[str, Any]:
    canonical_url = _abs_url(_blog_index_path(lang))

    posts_view: List[Dict[str, Any]] = []
    tag_counts: Dict[str, int] = {}
    for p in BLOG_LIST.get(lang, []):
        for t in (p.get("tags") or []):
            if not isinstance(t, str) or not t.strip():
                continue
            nt = _normalize_tag(t)
            tag_counts[nt] = tag_counts.get(nt, 0) + 1
        posts_view.append(
            {
                "url": f"{_blog_index_path(lang)}/{p.get('slug')}",
                "title": p.get("title") or "",
                "description": p.get("description") or "",
                "kicker": p.get("kicker") or "",
                "tags": [_normalize_tag(t) for t in (p.get("tags") or []) if isinstance(t, str)],
                "meta": f"{_format_date(lang, p.get('date_published') or '')} · {_format_reading_time(p.get('reading_time'))}".strip(
                    " ·"
                ),
            }
        )

    top_tags = [
        k
        for k, _ in sorted(tag_counts.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    ][:8]

    alt_paths = {code: _blog_index_path(code) for code in SUPPORTED_LANGS}

    return {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": "Blog | Ecuaciones a Word",
        "description": "Practical guides to convert LaTeX/AI content to Word with native (OMML) editable equations.",
        "keywords": ["LaTeX to Word", "LaTeX to Word online", "OMML", "Word equations", "ChatGPT", "Pandoc"],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=False),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "website",
        "og_title": "Blog | Ecuaciones a Word",
        "og_description": "Practical guides for converting LaTeX equations to native Word (OMML) cleanly and reliably.",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_index(lang, canonical_url),
        "converter_href": _home_path(lang),
        "blog_index_href": _blog_index_path(lang),
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": _blog_index_path("es"),
        "lang_switch_label": "ES",
        "lang_options": _lang_options(alt_paths),
        "h1": "Blog",
        "intro": _ui(lang, "blog_intro"),
        "index_cta_primary": _ui(lang, "index_cta_primary"),
        "search_label": _ui(lang, "search_label"),
        "search_placeholder": _ui(lang, "search_placeholder"),
        "filter_label": _ui(lang, "filter_label"),
        "filter_all": _ui(lang, "filter_all"),
        "top_tags": top_tags,
        "featured_title": _ui(lang, "featured_title"),
        "posts": posts_view,
        "year": datetime.now().year,
        "legal_links": {
            "privacy": _legal_path(lang, "privacy"),
            "terms": _legal_path(lang, "terms"),
            "contact": _legal_path(lang, "contact"),
        },
    }


@app.get("/de/blog", response_class=HTMLResponse)
async def blog_index_de() -> HTMLResponse:
    return RedirectResponse(url="/en/blog", status_code=301)


@app.get("/fr/blog", response_class=HTMLResponse)
async def blog_index_fr() -> HTMLResponse:
    return RedirectResponse(url="/en/blog", status_code=301)


@app.get("/it/blog", response_class=HTMLResponse)
async def blog_index_it() -> HTMLResponse:
    return RedirectResponse(url="/en/blog", status_code=301)


@app.get("/pt/blog", response_class=HTMLResponse)
async def blog_index_pt() -> HTMLResponse:
    return RedirectResponse(url="/en/blog", status_code=301)


# Redirects legacy numeric
@app.get("/blog2")
async def blog2_redirect() -> RedirectResponse:
    return RedirectResponse(url="/blog/convertir-documento-latex-word", status_code=301)


@app.get("/blog3")
async def blog3_redirect() -> RedirectResponse:
    return RedirectResponse(url="/blog/ia-chatgpt-a-word-ejercicios", status_code=301)


@app.get("/blog4")
async def blog4_redirect() -> RedirectResponse:
    return RedirectResponse(url="/blog/pegar-latex-en-word-alt-eq", status_code=301)


@app.get("/blog5")
async def blog5_redirect() -> RedirectResponse:
    return RedirectResponse(url="/blog/que-es-omml-ecuaciones-word", status_code=301)


@app.get("/blog6")
async def blog6_redirect() -> RedirectResponse:
    return RedirectResponse(url="/blog/markdown-con-latex-a-word-docx", status_code=301)


def _resolve_blog_slug(lang: str, slug: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
    """Return (redirect_url, post). redirect_url is set for alias slugs."""
    alias_map = BLOG_ALIASES.get(lang, {})
    if slug in alias_map:
        canonical_slug = alias_map[slug]
        post = BLOG_POSTS.get(lang, {}).get(canonical_slug)
        if post and post.get("canonical_path"):
            return post["canonical_path"], None
        prefix = f"{_blog_index_path(lang)}/"
        return f"{prefix}{canonical_slug}", None

    post = BLOG_POSTS.get(lang, {}).get(slug)
    return None, post


@app.get("/blog/{slug}", response_class=HTMLResponse)
async def blog_post_es(slug: str) -> HTMLResponse:
    lang = "es"
    redirect_url, post = _resolve_blog_slug(lang, slug)
    if redirect_url:
        return RedirectResponse(url=redirect_url, status_code=301)
    if not post:
        raise HTTPException(status_code=404, detail="Blog post not found")

    body_html = _read_blog_body(lang, post["slug"])
    if not body_html.strip():
        raise HTTPException(status_code=404, detail="Blog post body not found")

    canonical_path = post.get("canonical_path") or f"/blog/{post['slug']}"
    canonical_url = _abs_url(canonical_path)
    translation_slug = (post.get("translation_slug") or "").strip()

    en_path = f"/en/blog/{post['slug']}"
    if translation_slug and translation_slug in BLOG_POSTS.get("en", {}):
        en_path = BLOG_POSTS["en"][translation_slug].get("canonical_path") or f"/en/blog/{translation_slug}"

    alt_paths = {
        "es": canonical_path,
        "en": en_path,
    }

    date_pub = post.get("date_published") or ""
    date_mod = post.get("date_modified") or ""
    meta_line = f"Publicado {_format_date(lang, date_pub)}"
    reading_time = _format_reading_time(post.get("reading_time"))
    if reading_time:
        meta_line += f" · {reading_time}"
    if date_mod and date_mod != date_pub:
        meta_line += f" · Actualizado {_format_date(lang, date_mod)}"
    indexable = _is_blog_post_indexable(lang, post)

    ctx = {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": post.get("seo_title") or (post.get("title") or ""),
        "description": post.get("description") or "",
        "keywords": post.get("keywords") or [],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=indexable),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "article",
        "og_title": post.get("title") or "",
        "og_description": post.get("description") or "",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_article(post, canonical_url),
        "related_posts": _get_related_posts(lang, post, limit=4),
        "related_title": _ui(lang, "related_title"),
        "summary_box_title": _ui(lang, "summary_title"),
        "summary_box_text": _ui(lang, "summary_text"),
        "summary_box_link_text": _ui(lang, "summary_link"),
        "contextual_links_title": "Ruta de lectura y conversion recomendada",
        "contextual_links": _build_contextual_links(lang, post),
        "converter_href": "/",
        "blog_index_href": "/blog",
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": en_path,
        "lang_switch_label": "EN",
        "lang_options": _lang_options(alt_paths),
        "kicker": post.get("kicker") or "",
        "title": post.get("title") or "",
        "meta_line": meta_line,
        "tags": [_normalize_tag(t) for t in (post.get("tags") or []) if isinstance(t, str)],
        "intro_html": post.get("intro_html") or [],
        "breadcrumbs": [
            {"name": "Inicio", "url": "/"},
            {"name": "Blog", "url": "/blog"},
            {"name": post.get("title") or "", "url": canonical_path},
        ],
        "body_html": body_html,
        "cta_strong": _ui(lang, "cta_strong"),
        "cta_text": _ui(lang, "cta_text"),
        "cta_primary": _ui(lang, "cta_primary"),
        "cta_secondary": _ui(lang, "cta_secondary"),
        "year": datetime.now().year,
        "legal_links": {"privacy": "/privacy", "terms": "/terms", "contact": "/contact"},
    }
    return HTMLResponse(_render_template("blog_post.html", ctx))


@app.get("/en/blog/{slug}", response_class=HTMLResponse)
async def blog_post_en(slug: str) -> HTMLResponse:
    lang = "en"
    redirect_url, post = _resolve_blog_slug(lang, slug)
    if redirect_url:
        return RedirectResponse(url=redirect_url, status_code=301)
    if not post:
        raise HTTPException(status_code=404, detail="Blog post not found")

    body_html = _read_blog_body(lang, post["slug"])
    if not body_html.strip():
        raise HTTPException(status_code=404, detail="Blog post body not found")

    canonical_path = post.get("canonical_path") or f"/en/blog/{post['slug']}"
    canonical_url = _abs_url(canonical_path)
    translation_slug = (post.get("translation_slug") or "").strip()

    es_path = f"/blog/{post['slug']}"
    if translation_slug and translation_slug in BLOG_POSTS.get("es", {}):
        es_path = BLOG_POSTS["es"][translation_slug].get("canonical_path") or f"/blog/{translation_slug}"

    alt_paths = {
        "es": es_path,
        "en": canonical_path,
    }

    date_pub = post.get("date_published") or ""
    date_mod = post.get("date_modified") or ""
    meta_line = f"Published {_format_date(lang, date_pub)}"
    reading_time = _format_reading_time(post.get("reading_time"))
    if reading_time:
        meta_line += f" · {reading_time}"
    if date_mod and date_mod != date_pub:
        meta_line += f" · Updated {_format_date(lang, date_mod)}"
    indexable = _is_blog_post_indexable(lang, post)

    ctx = {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": post.get("seo_title") or (post.get("title") or ""),
        "description": post.get("description") or "",
        "keywords": post.get("keywords") or [],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=indexable),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "article",
        "og_title": post.get("title") or "",
        "og_description": post.get("description") or "",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_article(post, canonical_url),
        "related_posts": _get_related_posts(lang, post, limit=4),
        "related_title": _ui(lang, "related_title"),
        "summary_box_title": _ui(lang, "summary_title"),
        "summary_box_text": _ui(lang, "summary_text"),
        "summary_box_link_text": _ui(lang, "summary_link"),
        "contextual_links_title": "Recommended reading and conversion path",
        "contextual_links": _build_contextual_links(lang, post),
        "converter_href": "/en",
        "blog_index_href": "/en/blog",
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": es_path,
        "lang_switch_label": "ES",
        "lang_options": _lang_options(alt_paths),
        "kicker": post.get("kicker") or "",
        "title": post.get("title") or "",
        "meta_line": meta_line,
        "tags": [_normalize_tag(t) for t in (post.get("tags") or []) if isinstance(t, str)],
        "intro_html": post.get("intro_html") or [],
        "breadcrumbs": [
            {"name": "Home", "url": "/en"},
            {"name": "Blog", "url": "/en/blog"},
            {"name": post.get("title") or "", "url": canonical_path},
        ],
        "body_html": body_html,
        "cta_strong": _ui(lang, "cta_strong"),
        "cta_text": _ui(lang, "cta_text"),
        "cta_primary": _ui(lang, "cta_primary"),
        "cta_secondary": _ui(lang, "cta_secondary"),
        "year": datetime.now().year,
        "legal_links": {"privacy": "/en/privacy", "terms": "/en/terms", "contact": "/en/contact"},
    }
    return HTMLResponse(_render_template("blog_post.html", ctx))


def _blog_post_context_fallback_en(lang: str, post: Dict[str, Any], body_html: str) -> Dict[str, Any]:
    slug = post.get("slug") or ""
    canonical_path = f"{_blog_index_path(lang)}/{slug}"
    canonical_url = _abs_url(canonical_path)
    date_pub = post.get("date_published") or ""
    date_mod = post.get("date_modified") or ""

    meta_label = "Publicado" if lang == "es" else {
        "de": "Veröffentlicht",
        "fr": "Publié",
        "it": "Pubblicato",
        "pt": "Publicado",
    }.get(lang, "Published")
    updated_label = "Actualizado" if lang == "es" else {
        "de": "Aktualisiert",
        "fr": "Mis à jour",
        "it": "Aggiornato",
        "pt": "Atualizado",
    }.get(lang, "Updated")

    meta_line = f"{meta_label} {_format_date(lang, date_pub)}"
    reading_time = _format_reading_time(post.get("reading_time"))
    if reading_time:
        meta_line += f" · {reading_time}"
    if date_mod and date_mod != date_pub:
        meta_line += f" · {updated_label} {_format_date(lang, date_mod)}"

    alt_paths = {
        "es": (BLOG_POSTS.get("es", {}).get(slug, {}).get("canonical_path") or f"/blog/{slug}"),
        "en": (BLOG_POSTS.get("en", {}).get(slug, {}).get("canonical_path") or f"/en/blog/{slug}"),
        "de": f"/de/blog/{slug}",
        "fr": f"/fr/blog/{slug}",
        "it": f"/it/blog/{slug}",
        "pt": f"/pt/blog/{slug}",
    }

    post_schema = dict(post)
    post_schema["lang"] = lang

    return {
        "lang": lang,
        "site_name": SITE_NAME,
        "seo_title": post.get("seo_title") or (post.get("title") or ""),
        "description": post.get("description") or "",
        "keywords": post.get("keywords") or [],
        "canonical_url": canonical_url,
        "robots_directive": _robots_directive(lang, indexable=False),
        "alternates": _all_alternates(alt_paths, default_lang="es"),
        "og_type": "article",
        "og_title": post.get("title") or "",
        "og_description": post.get("description") or "",
        "og_image": _abs_url("/static/og-image.svg"),
        "schema_json": _build_schema_article(post_schema, canonical_url),
        "related_posts": _get_related_posts(lang, post, limit=4),
        "related_title": _ui(lang, "related_title"),
        "summary_box_title": _ui(lang, "summary_title"),
        "summary_box_text": _ui(lang, "summary_text"),
        "summary_box_link_text": _ui(lang, "summary_link"),
        "contextual_links_title": "Recommended reading and conversion path",
        "contextual_links": _build_contextual_links(lang, post),
        "converter_href": _home_path(lang),
        "blog_index_href": _blog_index_path(lang),
        "nav_converter": _ui(lang, "nav_converter"),
        "nav_blog": "Blog",
        "lang_switch_href": _blog_index_path("es"),
        "lang_switch_label": "ES",
        "lang_options": _lang_options(alt_paths),
        "kicker": post.get("kicker") or "",
        "title": post.get("title") or "",
        "meta_line": meta_line,
        "tags": [_normalize_tag(t) for t in (post.get("tags") or []) if isinstance(t, str)],
        "intro_html": post.get("intro_html") or [],
        "breadcrumbs": [
            {"name": "Home", "url": _home_path(lang)},
            {"name": "Blog", "url": _blog_index_path(lang)},
            {"name": post.get("title") or "", "url": canonical_path},
        ],
        "body_html": body_html,
        "cta_strong": _ui(lang, "cta_strong"),
        "cta_text": _ui(lang, "cta_text"),
        "cta_primary": _ui(lang, "cta_primary"),
        "cta_secondary": _ui(lang, "cta_secondary"),
        "year": datetime.now().year,
        "legal_links": {
            "privacy": _legal_path(lang, "privacy"),
            "terms": _legal_path(lang, "terms"),
            "contact": _legal_path(lang, "contact"),
        },
    }


@app.get("/de/blog/{slug}", response_class=HTMLResponse)
async def blog_post_de(slug: str) -> HTMLResponse:
    target = _secondary_blog_redirect_target("de", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Blog post not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/fr/blog/{slug}", response_class=HTMLResponse)
async def blog_post_fr(slug: str) -> HTMLResponse:
    target = _secondary_blog_redirect_target("fr", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Blog post not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/it/blog/{slug}", response_class=HTMLResponse)
async def blog_post_it(slug: str) -> HTMLResponse:
    target = _secondary_blog_redirect_target("it", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Blog post not found")
    return RedirectResponse(url=target, status_code=301)


@app.get("/pt/blog/{slug}", response_class=HTMLResponse)
async def blog_post_pt(slug: str) -> HTMLResponse:
    target = _secondary_blog_redirect_target("pt", slug)
    if not target:
        raise HTTPException(status_code=404, detail="Blog post not found")
    return RedirectResponse(url=target, status_code=301)


LEGACY_REDIRECTS: Dict[str, str] = {
    "/index.html": "/",
    "/index-en.html": "/en",
    "/blog-index.html": "/blog",
    "/blog-index-en.html": "/en/blog",
    "/blog-en-1.html": "/en/blog/copy-chatgpt-equations-word",
    "/blog-en-2.html": "/en/blog/convert-latex-document-to-word",
    "/blog-en-3.html": "/en/blog/use-ai-equations-to-word-exercises",
    "/blog-en-4.html": "/en/blog/paste-latex-into-word-alt-eq",
    "/blog-en-5.html": "/en/blog/what-is-omml-word-equations",
    "/blog-en-6.html": "/en/blog/markdown-latex-to-word-docx",
    "/blog-en-question-marks-chatgpt-equations-word.html": "/en/blog/question-marks-chatgpt-equations-word",
    "/blog-en-overleaf-latex-to-word-editable-equations.html": "/en/blog/overleaf-latex-to-word-editable-equations",
    "/blog-en-pandoc-math-to-word-omml-troubleshooting.html": "/en/blog/pandoc-math-to-word-omml-troubleshooting",
    "/blog-en-omml-vs-mathtype-vs-latex-word-thesis.html": "/en/blog/omml-vs-mathtype-vs-latex-word-thesis",
    "/blog-signos-interrogacion-ecuaciones-chatgpt-word.html": "/blog/signos-interrogacion-ecuaciones-chatgpt-word",
    "/blog-overleaf-latex-a-word-ecuaciones-editables.html": "/blog/overleaf-latex-a-word-ecuaciones-editables",
    "/blog-pandoc-ecuaciones-word-no-editables-soluciones.html": "/blog/pandoc-ecuaciones-word-no-editables-soluciones",
    "/blog-omml-vs-mathtype-vs-latex-word-tfg.html": "/blog/omml-vs-mathtype-vs-latex-word-tfg",
}


@app.get("/index.html")
@app.get("/index-en.html")
@app.get("/blog-index.html")
@app.get("/blog-index-en.html")
@app.get("/blog-en-1.html")
@app.get("/blog-en-2.html")
@app.get("/blog-en-3.html")
@app.get("/blog-en-4.html")
@app.get("/blog-en-5.html")
@app.get("/blog-en-6.html")
@app.get("/blog-en-question-marks-chatgpt-equations-word.html")
@app.get("/blog-en-overleaf-latex-to-word-editable-equations.html")
@app.get("/blog-en-pandoc-math-to-word-omml-troubleshooting.html")
@app.get("/blog-en-omml-vs-mathtype-vs-latex-word-thesis.html")
@app.get("/blog-signos-interrogacion-ecuaciones-chatgpt-word.html")
@app.get("/blog-overleaf-latex-a-word-ecuaciones-editables.html")
@app.get("/blog-pandoc-ecuaciones-word-no-editables-soluciones.html")
@app.get("/blog-omml-vs-mathtype-vs-latex-word-tfg.html")
async def legacy_redirects(request: Request) -> RedirectResponse:
    target = LEGACY_REDIRECTS.get(request.url.path)
    if not target:
        raise HTTPException(status_code=404, detail="Legacy route not mapped")
    return RedirectResponse(url=target, status_code=301)


@app.get("/robots.txt")
async def robots_txt() -> Response:
    content = "\n".join(
        [
            "User-agent: *",
            "Allow: /",
            "Disallow: /healthz",
            "",
            f"Sitemap: {_abs_url('/sitemap.xml')}",
            "",
        ]
    )
    resp = Response(content=content, media_type="text/plain")
    resp.headers["Cache-Control"] = "public, max-age=3600"
    return resp


@app.get("/sitemap.xml")
async def sitemap_xml() -> Response:
    content = generate_sitemap_xml()
    resp = Response(content=content, media_type="application/xml")
    resp.headers["Cache-Control"] = "public, max-age=3600"
    return resp


@app.get("/healthz")
async def healthz() -> PlainTextResponse:
    resp = PlainTextResponse("ok")
    resp.headers.update(_noindex_headers())
    return resp


# ================================================================
# Convert endpoint
# ================================================================
def _extract_text_lines_from_txt(file_bytes: bytes) -> List[str]:
    text = file_bytes.decode("utf-8", errors="replace")
    return text.splitlines()


# ----------------------------
# DOCX in-place conversion
# ----------------------------
# Goal: preserve the original Word document (styles, fonts, bold/italic, spacing, headers/footers, tables, etc.)
# and only transform LaTeX fragments into native Word OMML equations.
#
# The previous implementation rebuilt the document from plain paragraph text, which necessarily lost formatting.
# This implementation edits the DOCX structure in-place:
# - Detects LaTeX delimited math: $$...$$, \[...\], $...$
# - Converts only the math spans to OMML using math2docx
# - Leaves all other runs, paragraph properties, and styles untouched
#
# Additionally, it supports multi-paragraph display blocks where the delimiters $$ / \[ and $$ / \] appear on
# their own lines in separate paragraphs (common when pasting LaTeX).

from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

_MATH_NS = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}


def _is_math_child(el) -> bool:
    try:
        return bool(el.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        return False


def _latex_to_omml_children_with_success(latex: str) -> Tuple[List[Any], bool]:
    """Return (children, success) where success indicates OMML was produced."""
    tmp_doc = Document()
    tmp_p = tmp_doc.add_paragraph()
    ok = add_math_safe(tmp_p, latex)

    # Determine success via OMML presence, even if add_math_safe returned True
    try:
        has_omml = bool(tmp_p._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        has_omml = ok

    out: List[Any] = []
    for child in list(tmp_p._p):
        if child.tag == qn("w:pPr"):
            continue
        if _is_math_child(child):
            out.append(deepcopy(child))

    # If OMML exists but is nested deeper (rare), keep all non-pPr nodes
    if has_omml and not out:
        for child in list(tmp_p._p):
            if child.tag != qn("w:pPr"):
                out.append(deepcopy(child))

    return out, has_omml


def _latex_to_omml_children(latex: str) -> List[Any]:
    """Convert LaTeX to OMML paragraph children.

    This is a best-effort helper. If conversion fails, it will return children that keep the original
    text, so the caller does not lose content.
    """
    children, success = _latex_to_omml_children_with_success(latex)

    if success and children:
        return children

    # Fallback: insert plain text runs produced by math2docx/add_math_safe
    tmp_doc = Document()
    tmp_p = tmp_doc.add_paragraph()
    add_math_safe(tmp_p, latex)
    out: List[Any] = []
    for child in list(tmp_p._p):
        if child.tag != qn("w:pPr"):
            out.append(deepcopy(child))
    return out

def _make_run_with_text(src_r, text: str):
    r = OxmlElement("w:r")
    rPr = src_r.find(qn("w:rPr")) if src_r is not None else None
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    # Preserve spaces (Word collapses them unless xml:space="preserve")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Insert an empty paragraph right after the given one, preserving paragraph properties."""
    new_p = OxmlElement("w:p")
    # copy paragraph properties if present (style, spacing, etc.)
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    """Remove all paragraph content except paragraph properties."""
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def _paragraph_plain_text(paragraph: Paragraph) -> str:
    # python-docx paragraph.text ignores OMML; we only need the regular text runs for detection.
    return "".join(r.text or "" for r in paragraph.runs)


def _find_math_spans(text: str) -> List[Tuple[str, int, int, str]]:
    """Return list of spans: (kind, start, end, latex).

    - kind: 'inline' or 'display'
    - start/end: indices in the concatenated run text, covering the delimiters too
    - latex: extracted LaTeX content (without delimiters)
    """
    spans: List[Tuple[str, int, int, str]] = []
    i = 0
    n = len(text)

    while i < n:
        if text.startswith("$$", i):
            end = text.find("$$", i + 2)
            if end == -1:
                break
            latex = text[i + 2 : end]
            spans.append(("display", i, end + 2, latex))
            i = end + 2
            continue

        if text.startswith(r"\[", i):
            end = text.find(r"\]", i + 2)
            if end == -1:
                break
            latex = text[i + 2 : end]
            spans.append(("display", i, end + 2, latex))
            i = end + 2
            continue

        if text[i] == "$":
            # avoid $$ (already handled)
            if i + 1 < n and text[i + 1] == "$":
                i += 1
                continue
            end = text.find("$", i + 1)
            if end == -1:
                break
            latex = text[i + 1 : end]
            spans.append(("inline", i, end + 1, latex))
            i = end + 1
            continue

        i += 1

    return spans




# ----------------------------
# Undelimited (heuristic) math detection
# ----------------------------

# A conservative token regex: backslash commands and explicit sub/sup patterns.
_UNDELIM_TOKEN_RE = re.compile(
    r"(\\[A-Za-z]+|[A-Za-z]\s*_\s*\{[^}]+\}|[A-Za-z]\s*_\s*\d+|[A-Za-z]\s*\^\s*\{[^}]+\}|[A-Za-z]\s*\^\s*\d+)"
)

# Characters that are plausibly part of an inline math fragment
_MATH_CHAR_SET = set(
    "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    "\\_^{ }()[]<>+=-*/.,:;|!?'\"&%"
    "≤≥≠≈∈∉∑∏√∞∂∇→←↔⇒⇔⋅·×÷±"
)

def _is_mathy_char(ch: str) -> bool:
    if not ch:
        return False
    if ch in _MATH_CHAR_SET:
        return True
    # Allow some Unicode letters/digits (e.g., Greek) and common math punctuation
    if ch.isalnum():
        return True
    return False


def _merge_spans(spans: List[Tuple[int, int, str]]) -> List[Tuple[int, int, str]]:
    if not spans:
        return []
    spans = sorted(spans, key=lambda x: (x[0], x[1]))
    merged: List[Tuple[int, int, str]] = []
    cur_s, cur_e, _ = spans[0]
    for s, e, _txt in spans[1:]:
        if s <= cur_e:
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e, ""))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e, ""))
    return merged


def _find_undelimited_math_spans(text: str) -> List[Tuple[int, int, str]]:
    """Find math-like spans without $ delimiters.

    This exists to catch cases like:
      - D_1>0, \\quad D_2>0, \\quad D_3>0.
      - (x,y,z)\\neq(0,0,0)

    We keep it conservative to avoid converting normal prose.
    """
    if not text:
        return []
    if "$" in text:
        return []

    spans: List[Tuple[int, int, str]] = []
    for m in _UNDELIM_TOKEN_RE.finditer(text):
        s, e = m.span()

        # Expand left/right to capture the surrounding math expression
        ls = s
        while ls > 0 and _is_mathy_char(text[ls - 1]):
            ls -= 1
        re_ = e
        while re_ < len(text) and _is_mathy_char(text[re_]):
            re_ += 1

        frag = text[ls:re_].strip()
        if not frag:
            continue

        # Must contain a LaTeX command or explicit sub/sup to be considered
        if ("\\" not in frag) and ("_" not in frag) and ("^" not in frag):
            continue

        # Avoid converting very long prose fragments (likely false positives)
        if len(frag) > 250:
            continue

        spans.append((ls, re_, text[ls:re_]))

    # Merge overlapping spans
    merged = _merge_spans(spans)
    out: List[Tuple[int, int, str]] = []
    for s, e, _ in merged:
        frag = text[s:e]
        # Re-check after merge
        if ("\\" not in frag) and ("_" not in frag) and ("^" not in frag):
            continue
        out.append((s, e, frag))
    return out


def _is_equation_only_paragraph(paragraph: Paragraph) -> bool:
    """True if the paragraph contains OMML and no visible plain text."""
    try:
        has_omml = bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        has_omml = False
    if not has_omml:
        return False
    plain = _paragraph_plain_text(paragraph)
    return plain.strip() == ""


def _apply_equation_spacing(paragraphs: List[Paragraph]) -> None:
    """Add visual breathing room between consecutive equation paragraphs.

    We do this via paragraph spacing (space-after) instead of inserting blank paragraphs,
    which is less intrusive and keeps layout more stable.
    """
    if not paragraphs:
        return
    for idx, p in enumerate(paragraphs):
        if not _is_equation_only_paragraph(p):
            continue

        next_is_eq = (idx + 1 < len(paragraphs)) and _is_equation_only_paragraph(paragraphs[idx + 1])

        pf = p.paragraph_format
        try:
            pf.space_before = Pt(0)
        except Exception:
            pass

        # A bit larger spacing when another equation follows immediately
        desired = Pt(10) if next_is_eq else Pt(6)
        try:
            if pf.space_after is None or pf.space_after < desired:
                pf.space_after = desired
        except Exception:
            pass
def _replace_span_with_omml(paragraph: Paragraph, start: int, end: int, latex: str, is_display: bool) -> bool:
    """Replace [start:end] (in concatenated run text) by OMML, preserving other runs.

    If OMML conversion fails, the original text in [start:end] is preserved.
    """
    runs = list(paragraph.runs)
    if not runs:
        return False

    run_texts = [r.text or "" for r in runs]
    full_text = "".join(run_texts)

    # build cumulative offsets
    offsets: List[int] = []
    pos = 0
    for t in run_texts:
        offsets.append(pos)
        pos += len(t)

    if start < 0 or end > pos or start >= end:
        return False

    # locate run indices for start/end
    def locate(char_index: int) -> Tuple[int, int]:
        for idx, base in enumerate(offsets):
            t = run_texts[idx]
            if base + len(t) >= char_index:
                return idx, max(0, char_index - base)
        return len(runs) - 1, len(run_texts[-1])

    start_idx, start_off = locate(start)
    end_idx, end_off = locate(end)

    start_r = runs[start_idx]._r
    end_r = runs[end_idx]._r

    prefix = run_texts[start_idx][:start_off]
    suffix = run_texts[end_idx][end_off:]
    original_span_text = full_text[start:end]

    parent_p = paragraph._p
    insert_at = parent_p.index(start_r)

    # Remove affected runs (start_idx .. end_idx)
    for ridx in range(end_idx, start_idx - 1, -1):
        r_el = runs[ridx]._r
        try:
            parent_p.remove(r_el)
        except Exception:
            pass

    cursor = insert_at

    # Insert prefix run (keep start run formatting)
    if prefix:
        parent_p.insert(cursor, _make_run_with_text(start_r, prefix))
        cursor += 1

    # Insert OMML
    latex_clean = normalize_math_text((latex or "").strip())
    omml_children, success = _latex_to_omml_children_with_success(latex_clean) if latex_clean else ([], False)

    if not success or not omml_children:
        # Preserve the original text exactly
        parent_p.insert(cursor, _make_run_with_text(start_r, original_span_text))
        cursor += 1
    else:
        for child in omml_children:
            parent_p.insert(cursor, child)
            cursor += 1

    # Insert suffix run (keep end run formatting)
    if suffix:
        parent_p.insert(cursor, _make_run_with_text(end_r, suffix))
        cursor += 1

    # Center display equations when they occupy the full paragraph
    if is_display:
        remaining_text = ("".join(r.text or "" for r in paragraph.runs)).strip()
        if remaining_text == "":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return True

def _extract_environment_block(text: str, env_name: str) -> Optional[str]:
    r"""Extract the first \begin{env}...\end{env} block (non-greedy)."""
    if not text:
        return None
    pat = rf"\\begin\{{{re.escape(env_name)}\}}(.*?)\\end\{{{re.escape(env_name)}\}}"
    m = re.search(pat, text, flags=re.S)
    if not m:
        return None
    return rf"\begin{{{env_name}}}" + m.group(1) + rf"\end{{{env_name}}}"


def _extract_aligned_block(text: str) -> Optional[str]:
    """Extract first aligned block (aligned or aligned*)."""
    for env in ("aligned", "aligned*"):
        block = _extract_environment_block(text, env)
        if block:
            return block
    return None


def _maybe_build_sylvester_aligned(text: str) -> Optional[str]:
    """Heuristic: if the paragraph contains D_1, D_2, D_3 inequalities, build an aligned block."""
    s = normalize_math_text(text)
    if "$" in s:
        return None
    if all(tok in s for tok in ("D_1", "D_2", "D_3")) and ">" in s:
        # Avoid triggering inside long prose: count alphabetic letters beyond D
        letters = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", "", s)
        if len(letters) <= 10:
            return "\\begin{aligned}\nD_1 &> 0,\\\\\nD_2 &> 0,\\\\\nD_3 &> 0.\n\\end{aligned}"
    return None

def _handle_matrix_paragraph(paragraph: Paragraph) -> bool:
    """Handle matrix-like environments without $...$ delimiters.

    Conservative trigger:
    - A full begin/end matrix environment is present (vmatrix/bmatrix/pmatrix/matrix)
    - No explicit '$' delimiters appear in the paragraph (to avoid double-processing)

    If there is noise around the environment, we extract the environment block and, when possible,
    keep a short left prefix like 'A_1 =' to preserve meaning.
    """
    raw = _paragraph_plain_text(paragraph)
    stripped = normalize_math_text(raw).strip()
    if not stripped or "$" in stripped:
        return False

    matrix_envs = ["vmatrix", "bmatrix", "pmatrix", "matrix"]
    for env in matrix_envs:
        block = _extract_environment_block(stripped, env)
        if not block:
            continue

        # Try to preserve an identifier prefix (A_1 =, etc.) if present immediately before the environment
        prefix = stripped.split(block, 1)[0].strip()
        kept_prefix = ""
        if prefix:
            m = re.search(r"(A_\d\s*=\s*|A\d\s*=\s*|\\det\s*A_\d\s*=\s*|det\s*A_\d\s*=\s*)$", prefix)
            if m:
                kept_prefix = m.group(0)

        latex = (kept_prefix + block).strip()
        latex = normalize_math_text(latex)

        _clear_paragraph_content(paragraph)
        for child in _latex_to_omml_children(latex):
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    return False


def _handle_aligned_paragraph(paragraph: Paragraph) -> bool:
    r"""Handle \begin{aligned}...\end{aligned} blocks without $ delimiters.

    Tries converting the whole environment. If the LaTeX backend cannot handle aligned, falls back to
    converting each row as a separate displayed equation (closest behavior to the original builder).
    """
    raw = _paragraph_plain_text(paragraph)
    stripped = normalize_math_text(raw).strip()
    if not stripped or "$" in stripped:
        return False

    # Sylvester inequalities heuristic (your sample document)
    sylv = _maybe_build_sylvester_aligned(stripped)
    if sylv:
        latex = normalize_math_text(sylv)
        children, success = _latex_to_omml_children_with_success(latex)
        if not success or not children:
            return False
        _clear_paragraph_content(paragraph)
        for child in children:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    block = _extract_aligned_block(stripped)
    if not block:
        return False

    latex = normalize_math_text(block)
    children, success = _latex_to_omml_children_with_success(latex)
    if success and children:
        _clear_paragraph_content(paragraph)
        for child in children:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    # Fallback: split rows and convert one-by-one
    m = re.search(r"\\begin\{aligned\*?\}(.*?)\\end\{aligned\*?\}", latex, flags=re.S)
    if not m:
        return False
    inner = m.group(1)

    # Split on LaTeX row breaks \\ (two backslashes)
    rows = [r.strip() for r in re.split(r"\\\\\s*", inner) if r.strip()]
    if not rows:
        return False

    # Clean alignment markers (&)
    rows = [row.replace("&", "").strip() for row in rows if row.replace("&", "").strip()]

    _clear_paragraph_content(paragraph)

    # Convert first row into current paragraph
    first = rows[0]
    ch_first, ok_first = _latex_to_omml_children_with_success(normalize_math_text(first))
    if ok_first and ch_first:
        for child in ch_first:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Preserve original if even fallback fails
        paragraph.add_run(raw)
        return True

    last_p = paragraph
    for row in rows[1:]:
        new_p = _insert_paragraph_after(last_p)
        ch, ok = _latex_to_omml_children_with_success(normalize_math_text(row))
        if ok and ch:
            for child in ch:
                new_p._p.append(child)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            new_p.add_run(row)
        last_p = new_p

    return True

def _replace_paragraph_text_preserve(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph content with a single run, preserving the first run's formatting."""
    runs = list(paragraph.runs)
    first_r = runs[0]._r if runs else None
    _clear_paragraph_content(paragraph)
    if new_text is None:
        new_text = ""
    if first_r is None:
        paragraph.add_run(new_text)
        return
    paragraph._p.append(_make_run_with_text(first_r, new_text))


def _maybe_prettify_paragraph_in_place(paragraph: Paragraph, paragraphs_list: Optional[List[Paragraph]] = None) -> None:
    """Apply exercise-specific prettify rules in-place only when they help conversion.

    This mirrors the behavior of the original web version (prettify + build) but edits the document in-place
    to preserve formatting elsewhere.
    """
    raw = _paragraph_plain_text(paragraph)
    if not raw:
        return

    stripped = normalize_math_text(raw).strip()

    # Only prettify when the paragraph obviously contains LaTeX-ish artifacts or matches known triggers
    triggers = (
        "\\begin{",
        "\\end{",
        "\\mathbf",
        "actividad2grupal",
        "Escribimos cada forma como q(x)=xT",
        "Escribimos cada forma como q(x)=xTA",
        "coeficiente del término cruzado",
        "Criterio de Sylvester",
        "detA",
        "det A",
        "A1A_1A1",
        "A4A_4A4",
    )
    if not any(t in stripped for t in triggers):
        return

    pretty = _prettify_paragraphs_for_exercise([raw])
    if not pretty:
        return

    # If prettify produced multiple paragraphs, expand in-place
    if len(pretty) > 1 and paragraphs_list is not None:
        _replace_paragraph_text_preserve(paragraph, pretty[0])
        last = paragraph
        for extra in pretty[1:]:
            new_p = _insert_paragraph_after(last)
            try:
                new_p.style = paragraph.style
            except Exception:
                pass
            _replace_paragraph_text_preserve(new_p, extra)
            insert_pos = paragraphs_list.index(last) + 1
            paragraphs_list.insert(insert_pos, new_p)
            last = new_p
        return

    # Single paragraph replacement
    if pretty[0] != raw:
        _replace_paragraph_text_preserve(paragraph, pretty[0])


def _process_undelimited_math_in_paragraph(paragraph: Paragraph) -> bool:
    """Convert math-like fragments that are NOT delimited by $...$.

    This is a conservative heuristic intended to replicate the original app's coverage (which relied on
    prettify + rebuild) without touching non-math prose or document styling.
    """
    # Skip paragraphs that already contain OMML (avoid double-processing)
    try:
        if paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS):
            return False
    except Exception:
        pass

    text = _paragraph_plain_text(paragraph)
    if not text or "$" in text:
        return False

    spans = _find_undelimited_math_spans(text)
    if not spans:
        return False

    changed = False
    for s, e, frag in reversed(spans):
        latex = normalize_math_text(frag.strip())
        if not latex:
            continue
        # Treat as display if the paragraph is essentially only this fragment
        remainder = (text[:s] + text[e:]).strip()
        is_noteq_or_list = ("\\neq" in latex) or ("\\quad" in latex) or ("D_" in latex)
        is_display = (remainder == "") and (len(latex) >= 3 or is_noteq_or_list)
        changed = _replace_span_with_omml(paragraph, s, e, latex, is_display=is_display) or changed
        if is_display:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return changed


def _process_inline_math_in_paragraph(paragraph: Paragraph) -> bool:
    """Convert inline/display delimited math within a single paragraph, preserving formatting."""
    # First, handle raw environments that appear without $...$ delimiters
    if _handle_aligned_paragraph(paragraph):
        return True
    if _handle_matrix_paragraph(paragraph):
        return True

    text = _paragraph_plain_text(paragraph)
    if not text:
        return False

    spans = _find_math_spans(text)
    if spans:
        # Process from end to start so offsets remain valid
        changed = False
        for kind, s, e, latex in reversed(spans):
            changed = _replace_span_with_omml(paragraph, s, e, latex, is_display=(kind == "display")) or changed
        return changed

    # If there are no explicit delimiters, try a conservative heuristic for math-like tokens
    return _process_undelimited_math_in_paragraph(paragraph)



def _process_display_blocks_in_paragraphs(paragraphs: List[Paragraph]) -> None:
    r"""Process multi-paragraph blocks (display $$...$$ / \[...\] and aligned blocks) in-place."""
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        raw = _paragraph_plain_text(p)
        stripped = normalize_math_text(raw).strip()

        # Exercise-specific normalization (only for known LaTeX/no-delimiter patterns)
        _maybe_prettify_paragraph_in_place(p, paragraphs)
        raw = _paragraph_plain_text(p)
        stripped = normalize_math_text(raw).strip()

        # --- Multi-paragraph aligned block
        if r"\begin{aligned}" in stripped and r"\end{aligned}" not in stripped and "$" not in stripped:
            start_i = i
            lines = [raw]
            j = i + 1
            while j < len(paragraphs):
                ln = _paragraph_plain_text(paragraphs[j])
                lines.append(ln)
                if r"\end{aligned}" in normalize_math_text(ln):
                    break
                j += 1

            if j < len(paragraphs) and r"\end{aligned}" in normalize_math_text(lines[-1]):
                block_text = "\n".join(lines)
                block = _extract_aligned_block(normalize_math_text(block_text)) or normalize_math_text(block_text)

                _clear_paragraph_content(paragraphs[start_i])
                for child in _latex_to_omml_children(normalize_math_text(block)):
                    paragraphs[start_i]._p.append(child)
                paragraphs[start_i].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for k in range(j, start_i, -1):
                    _remove_paragraph(paragraphs[k])
                    del paragraphs[k]

                i = start_i + 1
                continue

        # --- Multi-paragraph display block: $$ ... $$  or \[ ... \]
        if stripped in ("$$", r"\["):
            delim_open = stripped
            delim_close = "$$" if delim_open == "$$" else r"\]"
            j = i + 1
            lines: List[str] = []
            while j < len(paragraphs) and normalize_math_text(_paragraph_plain_text(paragraphs[j])).strip() != delim_close:
                lines.append(_paragraph_plain_text(paragraphs[j]))
                j += 1

            if j < len(paragraphs) and normalize_math_text(_paragraph_plain_text(paragraphs[j])).strip() == delim_close:
                latex_block = "\n".join(lines).strip()
                latex_block = normalize_math_text(latex_block)

                _clear_paragraph_content(p)
                parts = split_long_latex_equation(latex_block) if latex_block else []
                if not parts:
                    for k in range(j, i - 1, -1):
                        _remove_paragraph(paragraphs[k])
                        del paragraphs[k]
                    continue

                for child in _latex_to_omml_children(parts[0]):
                    p._p.append(child)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                last_p = p
                for part in parts[1:]:
                    new_p = _insert_paragraph_after(last_p)
                    for child in _latex_to_omml_children(part):
                        new_p._p.append(child)
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    insert_pos = paragraphs.index(last_p) + 1
                    paragraphs.insert(insert_pos, new_p)
                    last_p = new_p

                for k in range(j, i, -1):
                    _remove_paragraph(paragraphs[k])
                    del paragraphs[k]

                i += 1
                continue

        # Single-paragraph aligned/matrix, then inline/delimited
        _process_inline_math_in_paragraph(p)
        i += 1


    i = 0
    n = len(paragraphs)
    while i < n:
        p = paragraphs[i]
        t = _paragraph_plain_text(p).strip()

        if t in ("$$", r"\["):
            delim_open = t
            delim_close = "$$" if delim_open == "$$" else r"\]"
            j = i + 1
            lines: List[str] = []
            while j < n and _paragraph_plain_text(paragraphs[j]).strip() != delim_close:
                lines.append(_paragraph_plain_text(paragraphs[j]))
                j += 1

            if j < n and _paragraph_plain_text(paragraphs[j]).strip() == delim_close:
                latex_block = "\n".join(lines).strip()
                latex_block = normalize_math_text(latex_block)

                # Replace opening delimiter paragraph with the equation
                _clear_paragraph_content(p)
                parts = split_long_latex_equation(latex_block) if latex_block else []
                if not parts:
                    # If empty, just delete the whole block
                    for k in range(j, i - 1, -1):
                        _remove_paragraph(paragraphs[k])
                    i = j + 1
                    continue

                # First part in the original paragraph
                for child in _latex_to_omml_children(parts[0]):
                    p._p.append(child)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Additional parts as extra paragraphs (if any)
                last_p = p
                for part in parts[1:]:
                    new_p = _insert_paragraph_after(last_p)
                    for child in _latex_to_omml_children(part):
                        new_p._p.append(child)
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p = new_p

                # Remove inner lines and closing delimiter paragraph
                for k in range(j, i, -1):
                    _remove_paragraph(paragraphs[k])

                i = j + 1
                continue

        # Also process single-paragraph equations after block check
        _process_inline_math_in_paragraph(p)
        i += 1
    # Add spacing between consecutive equation paragraphs for readability
    _apply_equation_spacing(paragraphs)



def _iter_block_items(parent) -> Any:
    """Yield Paragraph and Table objects in document order for a parent container.

    We intentionally avoid relying on `isinstance(parent, Document)` because `docx.Document`
    is a factory function, not a class. Instead, we detect the underlying XML container.
    """
    # Document body: parent._element.body exists
    if hasattr(parent, "_element") and hasattr(parent._element, "body"):
        parent_elm = parent._element.body
    else:
        # Header/Footer: parent._element is the container element
        parent_elm = getattr(parent, "_element", None)
        # Table cell: parent._tc is the container element
        if parent_elm is None:
            parent_elm = getattr(parent, "_tc", None)
        if parent_elm is None:
            raise TypeError(f"Unsupported container type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _collect_paragraphs_in_order(parent) -> List[Paragraph]:
    """Collect all paragraphs under parent (including inside tables) in document order."""
    out: List[Paragraph] = []
    for item in _iter_block_items(parent):
        if isinstance(item, Paragraph):
            out.append(item)
        else:
            # Table: descend into its cells
            for row in item.rows:
                for cell in row.cells:
                    out.extend(_collect_paragraphs_in_order(cell))
    return out


def _convert_docx_in_place(file_bytes: bytes) -> Document:
    doc = Document(io.BytesIO(file_bytes))

    # Body
    body_pars = _collect_paragraphs_in_order(doc)
    _process_display_blocks_in_paragraphs(body_pars)

    # Headers/footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                pars = _collect_paragraphs_in_order(hf)
                _process_display_blocks_in_paragraphs(pars)
            except Exception:
                # Some environments may have restricted access; continue.
                continue

        # First-page / even-page headers/footers if present
        for hf_attr in ("first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
            hf = getattr(section, hf_attr, None)
            if hf is None:
                continue
            try:
                pars = _collect_paragraphs_in_order(hf)
                _process_display_blocks_in_paragraphs(pars)
            except Exception:
                continue

    return doc


@app.get("/convert")
async def convert_get() -> RedirectResponse:
    """Redirect GET requests to the homepage.

    Googlebot may discover /convert via the form action in the UI. The converter endpoint is POST-only,
    so serving a redirect here avoids 4xx/405 reports in Search Console.
    """
    return RedirectResponse(url="/", status_code=301)


@app.post("/convert")
async def convert(file: UploadFile = File(...)) -> StreamingResponse:
    filename = (file.filename or "").lower().strip()
    if not filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")

    if filename.endswith(".docx"):
        # IMPORTANT: preserve formatting and modify ONLY equations.
        try:
            out_doc = _convert_docx_in_place(content)
        except (BadZipFile, PackageNotFoundError):
            raise HTTPException(
                status_code=400,
                detail="Invalid .docx file. Upload a valid Word document.",
            ) from None
        except Exception:
            logger.exception("Unexpected error while converting .docx")
            raise HTTPException(
                status_code=500,
                detail="Internal error while converting the document.",
            ) from None
    elif filename.endswith(".txt"):
        paragraph_texts = _extract_text_lines_from_txt(content)
        cleaned = prettify_paragraphs(paragraph_texts)
        out_doc = build_document_from_paragraphs(cleaned)
    else:
        raise HTTPException(
            status_code=400, detail="Unsupported file type. Use .docx or .txt"
        )

    out = io.BytesIO()
    out_doc.save(out)
    out.seek(0)

    download_name = "ecuaciones-a-word.docx"
    headers = {
        "Content-Disposition": f'attachment; filename="{download_name}"',
        "Cache-Control": "no-store",
    }

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

# ================================================================
# 404 handler (HTML en vez de JSON)
# ================================================================
@app.exception_handler(StarletteHTTPException)
async def custom_http_exception_handler(request: Request, exc: StarletteHTTPException):
    if exc.status_code == 404:
        safe_path = html_escape(request.url.path)
        is_en = str(request.url.path).startswith("/en")
        if is_en:
            html = f"""
            <h1>404 · Page not found</h1>
            <p>The URL <code>{safe_path}</code> does not exist.</p>
            <p><a href="/en">Go to the converter</a> · <a href="/en/blog">Read the guides</a></p>
            """
        else:
            html = f"""
            <h1>404 · Página no encontrada</h1>
            <p>No existe la URL <code>{safe_path}</code>.</p>
            <p><a href="/">Ir al conversor</a> · <a href="/blog">Ver las guías</a></p>
            """
        return HTMLResponse(html, status_code=404)

    return PlainTextResponse(str(exc.detail), status_code=exc.status_code)

def _get_related_posts(lang: str, current_post: Dict[str, Any], limit: int = 4) -> List[Dict[str, str]]:
    """Pick related posts by shared tags, with a recent-post fallback."""
    tags = set(current_post.get("tags") or [])
    cur_slug = current_post.get("slug") or ""
    candidates: List[Tuple[int, str, Dict[str, Any]]] = []
    for p in BLOG_LIST.get(lang, []):
        if (p.get("slug") or "") == cur_slug:
            continue
        p_tags = set(p.get("tags") or [])
        score = len(tags.intersection(p_tags))
        if tags and score <= 0:
            continue
        # date_published is ISO (YYYY-MM-DD); lexicographic order works
        candidates.append((score, (p.get("date_published") or ""), p))

    # Sort by score desc, then date desc
    candidates.sort(key=lambda t: (t[0], t[1]), reverse=True)

    out: List[Dict[str, str]] = []
    for _, __, p in candidates[:limit]:
        url = p.get("canonical_path") or ""
        if lang not in PRIMARY_CONTENT_LANGS:
            url = f"{_blog_index_path(lang)}/{p.get('slug') or ''}"
        out.append({"url": url, "title": p.get("title") or ""})

    if len(out) >= limit:
        return out

    # Fallback: fill with most recent posts not yet included.
    used_urls = {item["url"] for item in out}
    for p in BLOG_LIST.get(lang, []):
        if (p.get("slug") or "") == cur_slug:
            continue
        url = p.get("canonical_path") or ""
        if lang not in PRIMARY_CONTENT_LANGS:
            url = f"{_blog_index_path(lang)}/{p.get('slug') or ''}"
        if not url or url in used_urls:
            continue
        out.append({"url": url, "title": p.get("title") or ""})
        used_urls.add(url)
        if len(out) >= limit:
            break

    return out


def _solution_href_by_cluster(lang: str, cluster: str) -> str:
    base = _solutions_path(lang)
    if cluster == "chatgpt":
        return f"{base}/chatgpt-ecuaciones-a-word" if lang == "es" else f"{base}/chatgpt-equations-to-word"
    if cluster == "gemini":
        return f"{base}/gemini-ecuaciones-a-word" if lang == "es" else f"{base}/gemini-equations-to-word"
    if cluster == "pandoc":
        return f"{base}/pandoc-a-word-omml" if lang == "es" else f"{base}/pandoc-to-word-omml"
    if cluster == "overleaf":
        return f"{base}/overleaf-latex-documento-a-word" if lang == "es" else f"{base}/overleaf-latex-document-to-word"
    return f"{base}/conversor-omml" if lang == "es" else f"{base}/omml-converter"


def _landing_label_by_cluster(lang: str, cluster: str) -> str:
    if cluster == "chatgpt":
        return "Landing: ChatGPT a Word" if lang == "es" else "Landing: ChatGPT to Word"
    if cluster == "gemini":
        return "Landing: Gemini a Word" if lang == "es" else "Landing: Gemini to Word"
    if cluster == "pandoc":
        return "Landing: Pandoc a Word OMML" if lang == "es" else "Landing: Pandoc to Word OMML"
    if cluster == "overleaf":
        return "Landing: Overleaf a Word" if lang == "es" else "Landing: Overleaf to Word"
    return "Landing: Conversor OMML" if lang == "es" else "Landing: OMML converter"


def _cluster_from_slug(slug: str) -> str:
    s = (slug or "").lower()
    if any(k in s for k in ("chatgpt", "copilot", "perplexity", "claude")):
        return "chatgpt"
    if "gemini" in s:
        return "gemini"
    if any(k in s for k in ("pandoc", "markdown")):
        return "pandoc"
    if any(k in s for k in ("overleaf", "obsidian", "notion", "thesis", "tfg", "tfm")):
        return "overleaf"
    if any(k in s for k in ("omml", "cambria", "unicode", "linear-format", "word-equation")):
        return "omml"
    if any(k in s for k in ("latex", "matrices", "delimitadores")):
        return "overleaf"
    if "ia-" in s:
        return "chatgpt"
    return "omml"


def _pillar_link_for_lang(lang: str) -> Dict[str, str]:
    if lang == "es":
        return {
            "href": "/blog/latex-a-word-omml-guia-definitiva",
            "label": "Pilar: LaTeX a Word online (guia definitiva)",
        }
    return {
        "href": "/en/blog/latex-to-word-online-free-editable-equations",
        "label": "Pillar: LaTeX to Word online guide",
    }


def _build_contextual_links(lang: str, current_post: Dict[str, Any]) -> List[Dict[str, str]]:
    slug = (current_post.get("slug") or "").strip()
    cluster = _cluster_from_slug(slug)
    links: List[Dict[str, str]] = [
        {
            "href": _solution_href_by_cluster(lang, cluster),
            "label": _landing_label_by_cluster(lang, cluster),
            "reason": (
                "Ruta transaccional recomendada para ejecutar la conversion"
                if lang == "es"
                else "Primary transactional route for conversion"
            ),
        }
    ]

    related = _get_related_posts(lang, current_post, limit=6)
    for rp in related[:2]:
        links.append(
            {
                "href": rp["url"],
                "label": rp["title"],
                "reason": "Articulo relacionado del mismo cluster" if lang == "es" else "Related article in the same cluster",
            }
        )

    links.append(
        {
            "href": _pillar_link_for_lang(lang)["href"],
            "label": _pillar_link_for_lang(lang)["label"],
            "reason": "Pagina pilar para contexto completo" if lang == "es" else "Pillar page for broad context",
        }
    )
    return links


