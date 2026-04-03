from __future__ import annotations

import io
import logging
import os
import json
import re
from collections import defaultdict, deque
from pathlib import Path
from html import escape as html_escape
from threading import Lock
from time import monotonic
from urllib.parse import quote
from zipfile import BadZipFile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    HTMLResponse,
    PlainTextResponse,
    Response,
    StreamingResponse,
    RedirectResponse,
)
from fastapi.staticfiles import StaticFiles

from starlette.exceptions import HTTPException as StarletteHTTPException
from starlette.middleware.gzip import GZipMiddleware
from starlette.requests import Request

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response as StarletteResponse


class CanonicalHostRedirectMiddleware(BaseHTTPMiddleware):
    """Redirect non-canonical host/scheme to canonical values.

    - Skips localhost/127.0.0.1
    - Uses X-Forwarded-Proto / X-Forwarded-Host when present (common behind Render/CDNs)
    - Controlled via env vars:
        CANONICAL_HOST (e.g. www.ecuacionesaword.com)
        CANONICAL_SCHEME (e.g. https)
    """

    def __init__(self, app, canonical_host: str, canonical_scheme: str = "https"):
        super().__init__(app)
        self.canonical_host = (canonical_host or "").strip().lower()
        self.canonical_scheme = (canonical_scheme or "https").strip().lower()

    async def dispatch(self, request: Request, call_next):
        # Determine effective host/scheme considering reverse proxy headers
        host = (request.headers.get("x-forwarded-host") or request.headers.get("host") or "").split(",")[0].strip().lower()
        scheme = (request.headers.get("x-forwarded-proto") or request.url.scheme or "").split(",")[0].strip().lower()

        # Skip local/dev
        if host.startswith("127.0.0.1") or host.startswith("localhost"):
            return await call_next(request)

        # If not configured, do nothing
        if not self.canonical_host:
            return await call_next(request)

        if host != self.canonical_host or (self.canonical_scheme and scheme != self.canonical_scheme):
            # Preserve path + query
            url = str(request.url)
            # Rebuild
            path_q = request.url.path
            if request.url.query:
                path_q += "?" + request.url.query
            target = f"{self.canonical_scheme}://{self.canonical_host}{path_q}"
            return RedirectResponse(url=target, status_code=301)

        return await call_next(request)


def _add_common_headers(resp: StarletteResponse) -> StarletteResponse:
    # SEO/UX-safe defaults
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    resp.headers.setdefault("Permissions-Policy", "interest-cohort=()")
    resp.headers.setdefault("X-Frame-Options", "DENY")
    resp.headers.setdefault(
        "Content-Security-Policy",
        "; ".join(
            [
                "default-src 'self'",
                "base-uri 'self'",
                "object-src 'none'",
                "frame-ancestors 'none'",
                "img-src 'self' data: https:",
                "style-src 'self' 'unsafe-inline'",
                "font-src 'self' data:",
                "script-src 'self' 'unsafe-inline' https://www.googletagmanager.com https://www.google-analytics.com",
                "connect-src 'self' https://www.googletagmanager.com https://www.google-analytics.com",
                "form-action 'self'",
            ]
        ),
    )
    return resp

import math2docx


# ================================================================
# Config
# ================================================================
APP_TITLE = "Ecuaciones a Word (LaTeX → Word OMML)"
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

logger = logging.getLogger("ecuacionesaword")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_UPLOAD_CONTENT_LENGTH_BYTES = MAX_FILE_SIZE_BYTES + 64 * 1024
RATE_LIMIT_WINDOW_SECONDS = int(os.getenv("CONVERT_RATE_LIMIT_WINDOW_SECONDS", "60"))
RATE_LIMIT_MAX_REQUESTS = int(os.getenv("CONVERT_RATE_LIMIT_MAX_REQUESTS", "20"))
Segment = Tuple[str, str]  # ("text" | "inline" | "display", contenido)

# Puedes desactivar reglas específicas del "ejercicio" si no las quieres:
USE_EXERCISE_TWEAKS = os.getenv("USE_EXERCISE_TWEAKS", "0").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}

_RATE_LIMIT_BUCKETS: Dict[str, deque[float]] = defaultdict(deque)
_RATE_LIMIT_LOCK = Lock()


def _mojibake_score(value: str) -> int:
    if not value:
        return 0
    suspect_chars = {"\u00c2", "\u00c3", "\u00e2", "\ufffd"}
    score = sum(ch in suspect_chars for ch in value)
    score += value.count("â") * 2
    return score


def _decode_mojibake_once(value: str) -> str:
    candidates = [value]
    for source_encoding in ("latin-1", "cp1252"):
        try:
            candidates.append(value.encode(source_encoding).decode("utf-8"))
        except Exception:
            continue
    return min(candidates, key=lambda candidate: (_mojibake_score(candidate), len(candidate)))


def _fix_text_mojibake(text: str) -> str:
    """Best-effort repair for UTF-8/Latin-1 mojibake fragments."""
    if not text:
        return text

    fixed = text.lstrip("\ufeff")
    original_score = _mojibake_score(fixed)
    candidate = fixed
    for _ in range(4):
        new_candidate = _decode_mojibake_once(candidate)
        if new_candidate == candidate:
            break
        if _mojibake_score(new_candidate) <= _mojibake_score(candidate):
            candidate = new_candidate
            continue
        break
    fixed = candidate

    replacements = {
        "â€¦": "…",
        "â€œ": "“",
        "â€": "”",
        "â€˜": "‘",
        "â€™": "’",
        "â€“": "–",
        "â€”": "—",
        "â€¢": "•",
        "â†’": "→",
        "â†": "←",
        "â†”": "↔",
        "â‡’": "⇒",
        "â‡”": "⇔",
        "âˆ‘": "∑",
        "âˆ": "∏",
        "âˆš": "√",
        "âˆž": "∞",
        "âˆ‚": "∂",
        "âˆ‡": "∇",
        "â‰¤": "≤",
        "â‰¥": "≥",
        "â‰ ": "≠",
        "â‰ˆ": "≈",
        "Â©": "©",
        "Â·": "·",
        "Â¿": "¿",
        "Â¡": "¡",
        "Âº": "º",
        "Âª": "ª",
        "\u00a0": " ",
    }
    for bad, good in replacements.items():
        fixed = fixed.replace(bad, good)

    if _mojibake_score(fixed) > original_score:
        return text
    return fixed


def _deep_fix_mojibake(value: Any) -> Any:
    if isinstance(value, str):
        return _fix_text_mojibake(value)
    if isinstance(value, list):
        return [_deep_fix_mojibake(v) for v in value]
    if isinstance(value, tuple):
        return tuple(_deep_fix_mojibake(v) for v in value)
    if isinstance(value, dict):
        return {k: _deep_fix_mojibake(v) for k, v in value.items()}
    return value


def _request_is_https(request: Request) -> bool:
    scheme = (
        request.headers.get("x-forwarded-proto")
        or request.url.scheme
        or ""
    ).split(",")[0].strip().lower()
    return scheme == "https"


def _client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for", "")
    if forwarded:
        return forwarded.split(",")[0].strip()
    if request.client and request.client.host:
        return request.client.host
    return "unknown"


def _check_convert_rate_limit(request: Request) -> bool:
    now = monotonic()
    key = f"{_client_ip(request)}:convert"
    with _RATE_LIMIT_LOCK:
        bucket = _RATE_LIMIT_BUCKETS[key]
        while bucket and now - bucket[0] > RATE_LIMIT_WINDOW_SECONDS:
            bucket.popleft()
        if len(bucket) >= RATE_LIMIT_MAX_REQUESTS:
            return False
        bucket.append(now)
    return True


def _parse_allowed_origins(value: str) -> List[str]:
    return [o.strip() for o in value.split(",") if o.strip()]


def get_allowed_origins() -> List[str]:
    env_val = os.getenv("ALLOWED_ORIGINS", "").strip()
    if env_val:
        return _parse_allowed_origins(env_val)
    return [
        "https://www.ecuacionesaword.com",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
    ]


app = FastAPI(title=APP_TITLE)

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)
app.add_middleware(GZipMiddleware, minimum_size=800)

# Canonical redirects (host + scheme). Configure in production via env vars.
CANONICAL_HOST = os.getenv("CANONICAL_HOST", "").strip()
CANONICAL_SCHEME = os.getenv("CANONICAL_SCHEME", "https").strip()
if CANONICAL_HOST:
    app.add_middleware(CanonicalHostRedirectMiddleware, canonical_host=CANONICAL_HOST, canonical_scheme=CANONICAL_SCHEME)


@app.middleware("http")
async def add_common_headers_mw(request: Request, call_next):
    if request.url.path == "/convert" and request.method == "POST":
        content_length = request.headers.get("content-length", "").strip()
        if content_length.isdigit() and int(content_length) > MAX_UPLOAD_CONTENT_LENGTH_BYTES:
            resp = PlainTextResponse("File too large (max 5MB)", status_code=413)
            return _add_common_headers(resp)
        if not _check_convert_rate_limit(request):
            resp = PlainTextResponse("Too many conversion requests. Please try again shortly.", status_code=429)
            return _add_common_headers(resp)

    resp = await call_next(request)
    if _request_is_https(request):
        resp.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    return _add_common_headers(resp)


# Static
_static_dir = os.path.join(BASE_DIR, "static")
if os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")


# ================================================================
def _noindex_headers(follow: bool = False) -> Dict[str, str]:
    """Headers to discourage indexing for technical/non-content endpoints."""
    return {"X-Robots-Tag": "noindex, follow" if follow else "noindex, nofollow"}
# ================================================================
# Helpers
# ================================================================
def normalize_math_text(text: str) -> str:
    """Normalize common artifacts found in pasted/converted LaTeX inside Word.

    Goals:
    - Make LaTeX more parseable for math2docx (normalize Unicode minus, remove zero-width chars, etc.)
    - Preserve non-math text as much as possible
    - Keep original app's q1->q_1, D1->D_1 and x2->x^2 tweaks
    """
    if text is None:
        return ""

    t = str(text)

    # Remove common invisible/formatting artifacts that break parsers
    for ch in ("\u200b", "\ufeff", "\u2060", "\u200e", "\u200f", "\u00ad"):
        t = t.replace(ch, "")

    # Normalize various dash/minus characters to ASCII hyphen-minus
    for ch in ("\u2212", "\u2013", "\u2014", "\u2010", "\u2011", "\u2043"):
        t = t.replace(ch, "-")
    # Some docs include the literal minus sign already decoded
    t = t.replace("−", "-")

    # Normalize non-breaking spaces
    t = t.replace("\xa0", " ")

    # Normalize other Unicode spaces that sometimes appear in Word
    for ch in ("\u202f", "\u205f", "\u3000"):
        t = t.replace(ch, " ")
    # Replace en-space/em-space/thin-space etc. with regular spaces
    t = re.sub(r"[\u2000-\u200a]", " ", t)

    # Remove Private Use Area characters (often appear as invisible placeholders)
    t = re.sub(r"[\uE000-\uF8FF]", "", t)

    # Keep original exercise-oriented normalizations
    t = re.sub(r"q([1-4])\s*\(", r"q_\1(", t)  # q1( -> q_1(
    t = re.sub(r"\bD([1-4])\b", r"D_\1", t)  # D1 -> D_1

    # x2 -> x^2, y3 -> y^3, etc.
    for var in ("x", "y", "z"):
        for exp in ("2", "3", "4"):
            t = re.sub(rf"{var}{exp}\b", rf"{var}^{{{exp}}}", t)

    return t

def _prettify_paragraphs_for_exercise(paragraph_texts: List[str]) -> List[str]:
    out: List[str] = []
    for text in paragraph_texts:
        s = normalize_math_text(text)
        stripped = s.strip()
        if stripped == "":
            continue

        if stripped.lower() == "actividad2grupal":
            out.append("Actividad 2 (trabajo grupal)")
            continue

        if all(
            sym in stripped
            for sym in ["q_1(x,y,z)", "q_2(x,y,z)", "q_3(x,y,z)", "q_4(x,y,z)"]
        ):
            out.append("$$ q_1(x,y,z) = 2x^2 + 2y^2 + 2z^2 + 2xy + 2xz $$")
            out.append("$$ q_2(x,y,z) = x^2 - y^2 + z^2 + 2xy $$")
            out.append("$$ q_3(x,y,z) = 2x^2 - y^2 + 2z^2 + 4xy - 4yz $$")
            out.append("$$ q_4(x,y,z) = 4x^2 + 2y^2 + z^2 + 2yz $$")
            continue

        if stripped.startswith("Interpretamos que") and "garantizar beneficios" in stripped:
            out.append(
                "Interpretamos que \"garantizar beneficios\" significa que el beneficio "
                "$q(x,y,z)$ sea positivo para todo $(x,y,z)\\neq(0,0,0)$, es decir, "
                "que la forma cuadrática sea definida positiva."
            )
            continue

        if (
            "Escribimos cada forma como q(x)=xT" in stripped
            or "Escribimos cada forma como q(x)=xTA" in stripped
            or "Escribimos cada forma como q(x)=xTA x(\\mathbf x)" in stripped
            or (
                "Escribimos cada forma como q(x)=xTAx(\\mathbf x)=\\mathbf x^T A\\mathbf x"
                in stripped
            )
        ):
            out.append(
                "Escribimos cada forma como $q(\\mathbf x) = \\mathbf x^T A\\, \\mathbf x$, con $\\mathbf x = (x,y,z)^T$."
            )
            continue

        if "coeficiente del término cruzado" in stripped:
            out.append(
                "Recordando que el coeficiente del término cruzado $2x_i y_j$ se reparte como $a_{ij}=a_{ji}=1$:"
            )
            continue

        if "Forma q1" in stripped or "q1q_1q1" in stripped:
            out.append("Forma $q_1$:")
            continue
        if "Forma q2" in stripped or "q2q_2q2" in stripped:
            out.append("Forma $q_2$:")
            continue
        if "Forma q3" in stripped or "q3q_3q3" in stripped:
            out.append("Forma $q_3$:")
            continue
        if "Forma q4" in stripped or "q4q_4q4" in stripped:
            out.append("Forma $q_4$:")
            continue

        if "Criterio de Sylvester" in stripped:
            out.append("2. Criterio de Sylvester")
            continue
        if "Una matriz simétrica" in stripped and "definida positiva" in stripped:
            out.append(
                "Una matriz simétrica $A$ es definida positiva si todos sus menores principales líderes son positivos:"
            )
            continue

        if "D_1>0" in stripped and "D_2>0" in stripped and "D_3>0" in stripped:
            out.append(
                "\\begin{aligned}\nD_1 &> 0,\\\\\nD_2 &> 0,\\\\\nD_3 &> 0.\n\\end{aligned}"
            )
            continue

        d_terms = list(re.finditer(r"D_[1-4][^D]*", stripped))
        if len(d_terms) >= 2:
            for m in d_terms:
                term = m.group(0).strip().strip(",")
                if term:
                    out.append(f"$$ {term} $$")
            continue

        if "A1A_1A1" in stripped or "A1A1" in stripped:
            out.append("⇒ $A_1$ es definida positiva ⇒ $q_1$ definida positiva.")
            continue
        if "A4A_4A4" in stripped or "A4A4" in stripped:
            out.append("⇒ $A_4$ es definida positiva ⇒ $q_4$ definida positiva.")
            continue

        if "detA2" in stripped or "det A_2" in stripped:
            out.append("$$ \\det A_2 = -2 < 0 $$")
            continue
        if "detA3" in stripped or "det A_3" in stripped:
            out.append("$$ \\det A_3 = -20 < 0 $$")
            continue

        if "q3q_3q3" in stripped and "indefinida" in stripped:
            out.append("⇒ $q_3$ también es indefinida.")
            continue

        out.append(stripped)

    return out


def prettify_paragraphs(paragraph_texts: List[str]) -> List[str]:
    generic: List[str] = []
    for text in paragraph_texts:
        s = normalize_math_text(text)
        if s.strip():
            generic.append(s.strip())
    if not USE_EXERCISE_TWEAKS:
        return generic
    return _prettify_paragraphs_for_exercise(generic)


# ================================================================
# 2) Parsing LaTeX
# ================================================================
def new_paragraph(doc: Document, align: Optional[Any] = None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    return p


def add_math_safe(paragraph, latex: str) -> bool:
    """Try converting LaTeX to OMML; return True if OMML was produced.

    If conversion fails, the original LaTeX is added as plain text (best-effort) and False is returned.
    """
    try:
        math2docx.add_math(paragraph, latex)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Fallo convirtiendo LaTeX a OMML: %s", exc)
        paragraph.add_run(latex)
        return False

    # math2docx may silently fall back to plain text; verify OMML presence
    try:
        return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        return True


def _is_escaped(text: str, idx: int) -> bool:
    backslashes = 0
    j = idx - 1
    while j >= 0 and text[j] == "\\":
        backslashes += 1
        j -= 1
    return bool(backslashes % 2)


def _find_unescaped(text: str, token: str, start: int) -> int:
    idx = text.find(token, start)
    while idx != -1:
        if not _is_escaped(text, idx):
            return idx
        idx = text.find(token, idx + 1)
    return -1


def _looks_like_math_fragment(fragment: str) -> bool:
    candidate = (fragment or "").strip()
    if not candidate:
        return False
    if "\\" in candidate or any(ch in candidate for ch in "^_{}[]=<>"):
        return True
    if re.search(r"[A-Za-z]", candidate) and re.search(r"[0-9]", candidate):
        return True
    if re.search(r"[A-Za-z]", candidate) and any(op in candidate for op in ("+", "-", "*", "/", "=", "(", ")")):
        return True
    if re.fullmatch(r"[0-9().,+\-*/=\s]+", candidate) and any(op in candidate for op in ("+", "-", "*", "/", "=")):
        return True
    return False


def _is_probable_currency_span(
    text: str,
    start: int,
    end: int,
    fragment: Optional[str] = None,
) -> bool:
    if start < 0 or end <= start or text[start] != "$":
        return False
    next_char = text[start + 1] if start + 1 < len(text) else ""
    after_char = text[end] if end < len(text) else ""
    if fragment is None:
        fragment = text[start + 1 : end]
    candidate = (fragment or "").strip()
    if not next_char.isdigit():
        return False
    if after_char.isdigit():
        return True
    if not _looks_like_math_fragment(candidate):
        return True
    if candidate.endswith(("-", "–", "—", ",")):
        return True
    return False

def parse_math_segments(text: str) -> List[Segment]:
    segments: List[Segment] = []
    buf: List[str] = []

    def flush_text() -> None:
        if buf:
            segments.append(("text", "".join(buf)))
            buf.clear()

    i = 0
    n = len(text)
    while i < n:
        if text.startswith("$$", i) and not _is_escaped(text, i):
            flush_text()
            end = _find_unescaped(text, "$$", i + 2)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                segments.append(("display", latex.strip()))
            else:
                buf.append(text[i : end + 2])
            i = end + 2
            continue

        if text.startswith(r"\[", i) and not _is_escaped(text, i):
            flush_text()
            end = _find_unescaped(text, r"\]", i + 2)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                segments.append(("display", latex.strip()))
            else:
                buf.append(text[i : end + 2])
            i = end + 2
            continue

        if text.startswith(r"\(", i) and not _is_escaped(text, i):
            flush_text()
            end = _find_unescaped(text, r"\)", i + 2)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                segments.append(("inline", latex.strip()))
                i = end + 2
                continue
            buf.append(text[i : end + 2])
            i = end + 2
            continue

        if text[i] == "$" and not _is_escaped(text, i):
            if i + 1 < n and text[i + 1] == "$":
                buf.append(text[i])
                i += 1
                continue
            flush_text()
            end = _find_unescaped(text, "$", i + 1)
            if end == -1:
                buf.append(text[i:])
                break
            latex = text[i + 1 : end]
            if _is_probable_currency_span(text, i, end + 1, latex):
                buf.append(text[i : end + 1])
                i = end + 1
                continue
            if not _looks_like_math_fragment(latex):
                buf.append(text[i : end + 1])
                i = end + 1
                continue
            segments.append(("inline", latex.strip()))
            i = end + 1
            continue

        buf.append(text[i])
        i += 1

    flush_text()
    merged: List[Segment] = []
    for kind, value in segments:
        if kind == "text" and merged and merged[-1][0] == "text":
            merged[-1] = ("text", merged[-1][1] + value)
        else:
            merged.append((kind, value))
    return merged


def split_into_paragraph_descriptors(text: str) -> List[Dict[str, Any]]:
    segments = parse_math_segments(text)
    paragraphs: List[Dict[str, Any]] = []
    current_chunks: List[Segment] = []

    for kind, value in segments:
        if kind in ("text", "inline"):
            if value:
                current_chunks.append((kind, value))
        elif kind == "display":
            if current_chunks:
                paragraphs.append({"type": "inline", "chunks": current_chunks})
                current_chunks = []
            paragraphs.append({"type": "display", "latex": value})

    if current_chunks:
        paragraphs.append({"type": "inline", "chunks": current_chunks})

    return paragraphs


def split_long_latex_equation(latex: str, max_len: int = 120) -> List[str]:
    latex = latex.strip()
    if not latex:
        return []
    if len(latex) <= max_len:
        return [latex]

    if "\\\\" in latex:
        return [p.strip() for p in latex.split("\\\\") if p.strip()]

    if "=" in latex:
        tokens = [t.strip() for t in latex.split("=")]
        eqs: List[str] = []
        current = tokens[0]
        for tok in tokens[1:]:
            candidate = current + " = " + tok
            if len(candidate) > max_len and current:
                eqs.append(current)
                current = tok
            else:
                current = candidate
        if current:
            eqs.append(current)
        return eqs

    if "," in latex:
        tokens = [t.strip() for t in latex.split(",")]
        eqs2: List[str] = []
        current2 = ""
        for tok in tokens:
            candidate = (current2 + ", " if current2 else "") + tok
            if len(candidate) > max_len and current2:
                eqs2.append(current2)
                current2 = tok
            else:
                current2 = candidate
        if current2:
            eqs2.append(current2)
        return eqs2

    return [latex[i : i + max_len].strip() for i in range(0, len(latex), max_len)]


def add_aligned_block(doc: Document, aligned_block: str) -> None:
    content = aligned_block.strip()
    if content.startswith(r"\begin{aligned}"):
        content = content[len(r"\begin{aligned}") :]
    if content.endswith(r"\end{aligned}"):
        content = content[: -len(r"\end{aligned}")]
    content = content.strip()
    if not content:
        return
    rows = [row.strip() for row in content.split(r"\\") if row.strip()]
    for row in rows:
        row_no_amp = row.replace("&", "")
        p = new_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
        add_math_safe(p, row_no_amp)


def build_document_from_paragraphs(paragraph_texts: List[str]) -> Document:
    out_doc = Document()
    if not paragraph_texts:
        new_paragraph(out_doc)
        return out_doc

    in_display_block = False
    display_block_delim: Optional[str] = None  # "$$" o "\["
    display_lines: List[str] = []

    in_aligned_block = False
    aligned_lines: List[str] = []

    matrix_envs = ["vmatrix", "bmatrix", "pmatrix", "matrix"]

    for raw_text in paragraph_texts:
        text = normalize_math_text(raw_text)
        stripped = text.strip()

        # aligned multi-línea
        if in_aligned_block:
            aligned_lines.append(text)
            if r"\end{aligned}" in text:
                add_aligned_block(out_doc, "\n".join(aligned_lines))
                in_aligned_block = False
                aligned_lines = []
            continue

        if r"\begin{aligned}" in stripped:
            if r"\end{aligned}" in stripped:
                add_aligned_block(out_doc, stripped)
            else:
                in_aligned_block = True
                aligned_lines = [stripped]
            continue

        # display multi-línea $$ ... $$  o \[ ... \]
        if in_display_block:
            if (
                (display_block_delim == "$$" and stripped == "$$")
                or (display_block_delim == r"\[" and stripped == r"\]")
            ):
                latex = "\n".join(display_lines).strip()
                for part in split_long_latex_equation(latex):
                    p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                    add_math_safe(p, part)
                in_display_block = False
                display_block_delim = None
                display_lines = []
            else:
                display_lines.append(text)
            continue

        if stripped == "$$":
            in_display_block = True
            display_block_delim = "$$"
            display_lines = []
            continue

        if stripped == r"\[":
            in_display_block = True
            display_block_delim = r"\["
            display_lines = []
            continue

        # matrices en una línea
        handled_matrix = False
        for env in matrix_envs:
            begin = rf"\begin{{{env}}}"
            end = rf"\end{{{env}}}"
            if begin in stripped and end in stripped and "$" not in stripped:
                p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                add_math_safe(p, stripped)
                handled_matrix = True
                break
        if handled_matrix:
            continue

        # normal
        descriptors = split_into_paragraph_descriptors(text)
        if not descriptors:
            new_paragraph(out_doc)
            continue

        for desc in descriptors:
            if desc["type"] == "inline":
                p = new_paragraph(out_doc)
                for kind, value in desc["chunks"]:
                    if kind == "text":
                        p.add_run(value)
                    elif kind == "inline":
                        add_math_safe(p, value)
            elif desc["type"] == "display":
                for part in split_long_latex_equation(desc["latex"].strip()):
                    p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
                    add_math_safe(p, part)

    # cierre si termina dentro de display/aligned
    if in_display_block and display_lines:
        latex = "\n".join(display_lines).strip()
        for part in split_long_latex_equation(latex):
            p = new_paragraph(out_doc, WD_ALIGN_PARAGRAPH.CENTER)
            add_math_safe(p, part)

    if in_aligned_block and aligned_lines:
        add_aligned_block(out_doc, "\n".join(aligned_lines))

    return out_doc


# ================================================================

# ================================================================
# Convert endpoint
# ================================================================
def _extract_text_lines_from_txt(file_bytes: bytes) -> List[str]:
    text = file_bytes.decode("utf-8", errors="replace")
    return text.splitlines()


# ----------------------------
# DOCX in-place conversion
# ----------------------------
# Goal: preserve the original Word document (styles, fonts, bold/italic, spacing, headers/footers, tables, etc.)
# and only transform LaTeX fragments into native Word OMML equations.
#
# The previous implementation rebuilt the document from plain paragraph text, which necessarily lost formatting.
# This implementation edits the DOCX structure in-place:
# - Detects LaTeX delimited math: $$...$$, \[...\], $...$
# - Converts only the math spans to OMML using math2docx
# - Leaves all other runs, paragraph properties, and styles untouched
#
# Additionally, it supports multi-paragraph display blocks where the delimiters $$ / \[ and $$ / \] appear on
# their own lines in separate paragraphs (common when pasting LaTeX).

from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

_MATH_NS = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}


def _is_math_child(el) -> bool:
    try:
        return bool(el.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        return False


def _latex_to_omml_children_with_success(latex: str) -> Tuple[List[Any], bool]:
    """Return (children, success) where success indicates OMML was produced."""
    tmp_doc = Document()
    tmp_p = tmp_doc.add_paragraph()
    ok = add_math_safe(tmp_p, latex)

    # Determine success via OMML presence, even if add_math_safe returned True
    try:
        has_omml = bool(tmp_p._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        has_omml = ok

    out: List[Any] = []
    for child in list(tmp_p._p):
        if child.tag == qn("w:pPr"):
            continue
        if _is_math_child(child):
            out.append(deepcopy(child))

    # If OMML exists but is nested deeper (rare), keep all non-pPr nodes
    if has_omml and not out:
        for child in list(tmp_p._p):
            if child.tag != qn("w:pPr"):
                out.append(deepcopy(child))

    return out, has_omml


def _latex_to_omml_children(latex: str) -> List[Any]:
    """Convert LaTeX to OMML paragraph children.

    This is a best-effort helper. If conversion fails, it will return children that keep the original
    text, so the caller does not lose content.
    """
    children, success = _latex_to_omml_children_with_success(latex)

    if success and children:
        return children

    # Fallback: insert plain text runs produced by math2docx/add_math_safe
    tmp_doc = Document()
    tmp_p = tmp_doc.add_paragraph()
    add_math_safe(tmp_p, latex)
    out: List[Any] = []
    for child in list(tmp_p._p):
        if child.tag != qn("w:pPr"):
            out.append(deepcopy(child))
    return out

def _make_run_with_text(src_r, text: str):
    r = OxmlElement("w:r")
    rPr = src_r.find(qn("w:rPr")) if src_r is not None else None
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    # Preserve spaces (Word collapses them unless xml:space="preserve")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Insert an empty paragraph right after the given one, preserving paragraph properties."""
    new_p = OxmlElement("w:p")
    # copy paragraph properties if present (style, spacing, etc.)
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    """Remove all paragraph content except paragraph properties."""
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def _paragraph_plain_text(paragraph: Paragraph) -> str:
    # python-docx paragraph.text ignores OMML; we only need the regular text runs for detection.
    return "".join(r.text or "" for r in paragraph.runs)


def _find_math_spans(text: str) -> List[Tuple[str, int, int, str]]:
    """Return list of spans: (kind, start, end, latex).

    - kind: 'inline' or 'display'
    - start/end: indices in the concatenated run text, covering the delimiters too
    - latex: extracted LaTeX content (without delimiters)
    """
    spans: List[Tuple[str, int, int, str]] = []
    i = 0
    n = len(text)

    while i < n:
        if text.startswith("$$", i) and not _is_escaped(text, i):
            end = _find_unescaped(text, "$$", i + 2)
            if end == -1:
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                spans.append(("display", i, end + 2, latex))
                i = end + 2
                continue
            i += 2
            continue

        if text.startswith(r"\[", i) and not _is_escaped(text, i):
            end = _find_unescaped(text, r"\]", i + 2)
            if end == -1:
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                spans.append(("display", i, end + 2, latex))
                i = end + 2
                continue
            i += 2
            continue

        if text.startswith(r"\(", i) and not _is_escaped(text, i):
            end = _find_unescaped(text, r"\)", i + 2)
            if end == -1:
                break
            latex = text[i + 2 : end]
            if _looks_like_math_fragment(latex):
                spans.append(("inline", i, end + 2, latex))
                i = end + 2
                continue
            i += 2
            continue

        if text[i] == "$" and not _is_escaped(text, i):
            if i + 1 < n and text[i + 1] == "$":
                i += 1
                continue
            end = _find_unescaped(text, "$", i + 1)
            if end == -1:
                break
            latex = text[i + 1 : end]
            if _looks_like_math_fragment(latex) and not _is_probable_currency_span(text, i, end):
                spans.append(("inline", i, end + 1, latex))
                i = end + 1
                continue
            i += 1
            continue

        i += 1

    return spans




# ----------------------------
# Undelimited (heuristic) math detection
# ----------------------------

# A conservative token regex: backslash commands and explicit sub/sup patterns.
_UNDELIM_TOKEN_RE = re.compile(
    r"(\\[A-Za-z]+|[A-Za-z]\s*_\s*\{[^}]+\}|[A-Za-z]\s*_\s*\d+|[A-Za-z]\s*\^\s*\{[^}]+\}|[A-Za-z]\s*\^\s*\d+)"
)

# Characters that are plausibly part of an inline math fragment
_MATH_CHAR_SET = set(
    "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    "\\_^{ }()[]<>+=-*/.,:;|!?'\"&%"
    "≤≥≠≈∈∉∑∏√∞∂∇→←↔⇒⇔⋅·×÷±"
)

def _is_mathy_char(ch: str) -> bool:
    if not ch:
        return False
    if ch in _MATH_CHAR_SET:
        return True
    # Allow some Unicode letters/digits (e.g., Greek) and common math punctuation
    if ch.isalnum():
        return True
    return False


def _merge_spans(spans: List[Tuple[int, int, str]]) -> List[Tuple[int, int, str]]:
    if not spans:
        return []
    spans = sorted(spans, key=lambda x: (x[0], x[1]))
    merged: List[Tuple[int, int, str]] = []
    cur_s, cur_e, _ = spans[0]
    for s, e, _txt in spans[1:]:
        if s <= cur_e:
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e, ""))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e, ""))
    return merged


def _find_undelimited_math_spans(text: str) -> List[Tuple[int, int, str]]:
    """Find math-like spans without $ delimiters.

    This exists to catch cases like:
      - D_1>0, \\quad D_2>0, \\quad D_3>0.
      - (x,y,z)\\neq(0,0,0)

    We keep it conservative to avoid converting normal prose.
    """
    if not text:
        return []
    if "$" in text:
        return []

    spans: List[Tuple[int, int, str]] = []
    for m in _UNDELIM_TOKEN_RE.finditer(text):
        s, e = m.span()

        # Expand left/right to capture the surrounding math expression
        ls = s
        while ls > 0 and _is_mathy_char(text[ls - 1]):
            ls -= 1
        re_ = e
        while re_ < len(text) and _is_mathy_char(text[re_]):
            re_ += 1

        frag = text[ls:re_].strip()
        if not frag:
            continue

        # Must contain a LaTeX command or explicit sub/sup to be considered
        if ("\\" not in frag) and ("_" not in frag) and ("^" not in frag):
            continue

        # Avoid converting very long prose fragments (likely false positives)
        if len(frag) > 250:
            continue

        spans.append((ls, re_, text[ls:re_]))

    # Merge overlapping spans
    merged = _merge_spans(spans)
    out: List[Tuple[int, int, str]] = []
    for s, e, _ in merged:
        frag = text[s:e]
        # Re-check after merge
        if ("\\" not in frag) and ("_" not in frag) and ("^" not in frag):
            continue
        out.append((s, e, frag))
    return out


def _is_equation_only_paragraph(paragraph: Paragraph) -> bool:
    """True if the paragraph contains OMML and no visible plain text."""
    try:
        has_omml = bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS))
    except Exception:
        has_omml = False
    if not has_omml:
        return False
    plain = _paragraph_plain_text(paragraph)
    return plain.strip() == ""


def _apply_equation_spacing(paragraphs: List[Paragraph]) -> None:
    """Add visual breathing room between consecutive equation paragraphs.

    We do this via paragraph spacing (space-after) instead of inserting blank paragraphs,
    which is less intrusive and keeps layout more stable.
    """
    if not paragraphs:
        return
    for idx, p in enumerate(paragraphs):
        if not _is_equation_only_paragraph(p):
            continue

        next_is_eq = (idx + 1 < len(paragraphs)) and _is_equation_only_paragraph(paragraphs[idx + 1])

        pf = p.paragraph_format
        try:
            pf.space_before = Pt(0)
        except Exception:
            pass

        # A bit larger spacing when another equation follows immediately
        desired = Pt(10) if next_is_eq else Pt(6)
        try:
            if pf.space_after is None or pf.space_after < desired:
                pf.space_after = desired
        except Exception:
            pass
def _replace_span_with_omml(paragraph: Paragraph, start: int, end: int, latex: str, is_display: bool) -> bool:
    """Replace [start:end] (in concatenated run text) by OMML, preserving other runs.

    If OMML conversion fails, the original text in [start:end] is preserved.
    """
    runs = list(paragraph.runs)
    if not runs:
        return False

    run_texts = [r.text or "" for r in runs]
    full_text = "".join(run_texts)

    # build cumulative offsets
    offsets: List[int] = []
    pos = 0
    for t in run_texts:
        offsets.append(pos)
        pos += len(t)

    if start < 0 or end > pos or start >= end:
        return False

    # locate run indices for start/end
    def locate(char_index: int) -> Tuple[int, int]:
        for idx, base in enumerate(offsets):
            t = run_texts[idx]
            if base + len(t) >= char_index:
                return idx, max(0, char_index - base)
        return len(runs) - 1, len(run_texts[-1])

    start_idx, start_off = locate(start)
    end_idx, end_off = locate(end)

    start_r = runs[start_idx]._r
    end_r = runs[end_idx]._r

    prefix = run_texts[start_idx][:start_off]
    suffix = run_texts[end_idx][end_off:]
    original_span_text = full_text[start:end]

    parent_p = paragraph._p
    insert_at = parent_p.index(start_r)

    # Remove affected runs (start_idx .. end_idx)
    for ridx in range(end_idx, start_idx - 1, -1):
        r_el = runs[ridx]._r
        try:
            parent_p.remove(r_el)
        except Exception:
            pass

    cursor = insert_at

    # Insert prefix run (keep start run formatting)
    if prefix:
        parent_p.insert(cursor, _make_run_with_text(start_r, prefix))
        cursor += 1

    # Insert OMML
    latex_clean = normalize_math_text((latex or "").strip())
    omml_children, success = _latex_to_omml_children_with_success(latex_clean) if latex_clean else ([], False)

    if not success or not omml_children:
        # Preserve the original text exactly
        parent_p.insert(cursor, _make_run_with_text(start_r, original_span_text))
        cursor += 1
    else:
        for child in omml_children:
            parent_p.insert(cursor, child)
            cursor += 1

    # Insert suffix run (keep end run formatting)
    if suffix:
        parent_p.insert(cursor, _make_run_with_text(end_r, suffix))
        cursor += 1

    # Center display equations when they occupy the full paragraph
    if is_display:
        remaining_text = ("".join(r.text or "" for r in paragraph.runs)).strip()
        if remaining_text == "":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return True

def _extract_environment_block(text: str, env_name: str) -> Optional[str]:
    r"""Extract the first \begin{env}...\end{env} block (non-greedy)."""
    if not text:
        return None
    pat = rf"\\begin\{{{re.escape(env_name)}\}}(.*?)\\end\{{{re.escape(env_name)}\}}"
    m = re.search(pat, text, flags=re.S)
    if not m:
        return None
    return rf"\begin{{{env_name}}}" + m.group(1) + rf"\end{{{env_name}}}"


def _extract_aligned_block(text: str) -> Optional[str]:
    """Extract first aligned block (aligned or aligned*)."""
    for env in ("aligned", "aligned*"):
        block = _extract_environment_block(text, env)
        if block:
            return block
    return None


def _maybe_build_sylvester_aligned(text: str) -> Optional[str]:
    """Heuristic: if the paragraph contains D_1, D_2, D_3 inequalities, build an aligned block."""
    s = normalize_math_text(text)
    if "$" in s:
        return None
    if all(tok in s for tok in ("D_1", "D_2", "D_3")) and ">" in s:
        # Avoid triggering inside long prose: count alphabetic letters beyond D
        letters = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", "", s)
        if len(letters) <= 10:
            return "\\begin{aligned}\nD_1 &> 0,\\\\\nD_2 &> 0,\\\\\nD_3 &> 0.\n\\end{aligned}"
    return None

def _handle_matrix_paragraph(paragraph: Paragraph) -> bool:
    """Handle matrix-like environments without $...$ delimiters.

    Conservative trigger:
    - A full begin/end matrix environment is present (vmatrix/bmatrix/pmatrix/matrix)
    - No explicit '$' delimiters appear in the paragraph (to avoid double-processing)

    If there is noise around the environment, we extract the environment block and, when possible,
    keep a short left prefix like 'A_1 =' to preserve meaning.
    """
    raw = _paragraph_plain_text(paragraph)
    stripped = normalize_math_text(raw).strip()
    if not stripped or "$" in stripped:
        return False

    matrix_envs = ["vmatrix", "bmatrix", "pmatrix", "matrix"]
    for env in matrix_envs:
        block = _extract_environment_block(stripped, env)
        if not block:
            continue

        # Try to preserve an identifier prefix (A_1 =, etc.) if present immediately before the environment
        prefix = stripped.split(block, 1)[0].strip()
        kept_prefix = ""
        if prefix:
            m = re.search(r"(A_\d\s*=\s*|A\d\s*=\s*|\\det\s*A_\d\s*=\s*|det\s*A_\d\s*=\s*)$", prefix)
            if m:
                kept_prefix = m.group(0)

        latex = (kept_prefix + block).strip()
        latex = normalize_math_text(latex)

        _clear_paragraph_content(paragraph)
        for child in _latex_to_omml_children(latex):
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    return False


def _handle_aligned_paragraph(paragraph: Paragraph) -> bool:
    r"""Handle \begin{aligned}...\end{aligned} blocks without $ delimiters.

    Tries converting the whole environment. If the LaTeX backend cannot handle aligned, falls back to
    converting each row as a separate displayed equation (closest behavior to the original builder).
    """
    raw = _paragraph_plain_text(paragraph)
    stripped = normalize_math_text(raw).strip()
    if not stripped or "$" in stripped:
        return False

    # Sylvester inequalities heuristic (your sample document)
    sylv = _maybe_build_sylvester_aligned(stripped)
    if sylv:
        latex = normalize_math_text(sylv)
        children, success = _latex_to_omml_children_with_success(latex)
        if not success or not children:
            return False
        _clear_paragraph_content(paragraph)
        for child in children:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    block = _extract_aligned_block(stripped)
    if not block:
        return False

    latex = normalize_math_text(block)
    children, success = _latex_to_omml_children_with_success(latex)
    if success and children:
        _clear_paragraph_content(paragraph)
        for child in children:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    # Fallback: split rows and convert one-by-one
    m = re.search(r"\\begin\{aligned\*?\}(.*?)\\end\{aligned\*?\}", latex, flags=re.S)
    if not m:
        return False
    inner = m.group(1)

    # Split on LaTeX row breaks \\ (two backslashes)
    rows = [r.strip() for r in re.split(r"\\\\\s*", inner) if r.strip()]
    if not rows:
        return False

    # Clean alignment markers (&)
    rows = [row.replace("&", "").strip() for row in rows if row.replace("&", "").strip()]

    _clear_paragraph_content(paragraph)

    # Convert first row into current paragraph
    first = rows[0]
    ch_first, ok_first = _latex_to_omml_children_with_success(normalize_math_text(first))
    if ok_first and ch_first:
        for child in ch_first:
            paragraph._p.append(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Preserve original if even fallback fails
        paragraph.add_run(raw)
        return True

    last_p = paragraph
    for row in rows[1:]:
        new_p = _insert_paragraph_after(last_p)
        ch, ok = _latex_to_omml_children_with_success(normalize_math_text(row))
        if ok and ch:
            for child in ch:
                new_p._p.append(child)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            new_p.add_run(row)
        last_p = new_p

    return True

def _replace_paragraph_text_preserve(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph content with a single run, preserving the first run's formatting."""
    runs = list(paragraph.runs)
    first_r = runs[0]._r if runs else None
    _clear_paragraph_content(paragraph)
    if new_text is None:
        new_text = ""
    if first_r is None:
        paragraph.add_run(new_text)
        return
    paragraph._p.append(_make_run_with_text(first_r, new_text))


def _maybe_prettify_paragraph_in_place(paragraph: Paragraph, paragraphs_list: Optional[List[Paragraph]] = None) -> None:
    """Apply exercise-specific prettify rules in-place only when they help conversion.

    This mirrors the behavior of the original web version (prettify + build) but edits the document in-place
    to preserve formatting elsewhere.
    """
    raw = _paragraph_plain_text(paragraph)
    if not raw:
        return

    stripped = normalize_math_text(raw).strip()

    # Only prettify when the paragraph obviously contains LaTeX-ish artifacts or matches known triggers
    triggers = (
        "\\begin{",
        "\\end{",
        "\\mathbf",
        "actividad2grupal",
        "Escribimos cada forma como q(x)=xT",
        "Escribimos cada forma como q(x)=xTA",
        "coeficiente del término cruzado",
        "Criterio de Sylvester",
        "detA",
        "det A",
        "A1A_1A1",
        "A4A_4A4",
    )
    if not any(t in stripped for t in triggers):
        return

    pretty = _prettify_paragraphs_for_exercise([raw])
    if not pretty:
        return

    # If prettify produced multiple paragraphs, expand in-place
    if len(pretty) > 1 and paragraphs_list is not None:
        _replace_paragraph_text_preserve(paragraph, pretty[0])
        last = paragraph
        for extra in pretty[1:]:
            new_p = _insert_paragraph_after(last)
            try:
                new_p.style = paragraph.style
            except Exception:
                pass
            _replace_paragraph_text_preserve(new_p, extra)
            insert_pos = paragraphs_list.index(last) + 1
            paragraphs_list.insert(insert_pos, new_p)
            last = new_p
        return

    # Single paragraph replacement
    if pretty[0] != raw:
        _replace_paragraph_text_preserve(paragraph, pretty[0])


def _process_undelimited_math_in_paragraph(paragraph: Paragraph) -> bool:
    """Convert math-like fragments that are NOT delimited by $...$.

    This is a conservative heuristic intended to replicate the original app's coverage (which relied on
    prettify + rebuild) without touching non-math prose or document styling.
    """
    # Skip paragraphs that already contain OMML (avoid double-processing)
    try:
        if paragraph._p.xpath(".//m:oMath | .//m:oMathPara", namespaces=_MATH_NS):
            return False
    except Exception:
        pass

    text = _paragraph_plain_text(paragraph)
    if not text or "$" in text:
        return False

    spans = _find_undelimited_math_spans(text)
    if not spans:
        return False

    changed = False
    for s, e, frag in reversed(spans):
        latex = normalize_math_text(frag.strip())
        if not latex:
            continue
        # Treat as display if the paragraph is essentially only this fragment
        remainder = (text[:s] + text[e:]).strip()
        is_noteq_or_list = ("\\neq" in latex) or ("\\quad" in latex) or ("D_" in latex)
        is_display = (remainder == "") and (len(latex) >= 3 or is_noteq_or_list)
        changed = _replace_span_with_omml(paragraph, s, e, latex, is_display=is_display) or changed
        if is_display:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return changed


def _process_inline_math_in_paragraph(paragraph: Paragraph) -> bool:
    """Convert inline/display delimited math within a single paragraph, preserving formatting."""
    # First, handle raw environments that appear without $...$ delimiters
    if _handle_aligned_paragraph(paragraph):
        return True
    if _handle_matrix_paragraph(paragraph):
        return True

    text = _paragraph_plain_text(paragraph)
    if not text:
        return False

    spans = _find_math_spans(text)
    if spans:
        # Process from end to start so offsets remain valid
        changed = False
        for kind, s, e, latex in reversed(spans):
            changed = _replace_span_with_omml(paragraph, s, e, latex, is_display=(kind == "display")) or changed
        return changed

    # If there are no explicit delimiters, try a conservative heuristic for math-like tokens
    return _process_undelimited_math_in_paragraph(paragraph)



def _process_display_blocks_in_paragraphs(paragraphs: List[Paragraph]) -> None:
    r"""Process multi-paragraph blocks (display $$...$$ / \[...\] and aligned blocks) in-place."""
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        raw = _paragraph_plain_text(p)
        stripped = normalize_math_text(raw).strip()

        # Exercise-specific normalization (only for known LaTeX/no-delimiter patterns)
        _maybe_prettify_paragraph_in_place(p, paragraphs)
        raw = _paragraph_plain_text(p)
        stripped = normalize_math_text(raw).strip()

        # --- Multi-paragraph aligned block
        if r"\begin{aligned}" in stripped and r"\end{aligned}" not in stripped and "$" not in stripped:
            start_i = i
            lines = [raw]
            j = i + 1
            while j < len(paragraphs):
                ln = _paragraph_plain_text(paragraphs[j])
                lines.append(ln)
                if r"\end{aligned}" in normalize_math_text(ln):
                    break
                j += 1

            if j < len(paragraphs) and r"\end{aligned}" in normalize_math_text(lines[-1]):
                block_text = "\n".join(lines)
                block = _extract_aligned_block(normalize_math_text(block_text)) or normalize_math_text(block_text)

                _clear_paragraph_content(paragraphs[start_i])
                for child in _latex_to_omml_children(normalize_math_text(block)):
                    paragraphs[start_i]._p.append(child)
                paragraphs[start_i].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for k in range(j, start_i, -1):
                    _remove_paragraph(paragraphs[k])
                    del paragraphs[k]

                i = start_i + 1
                continue

        # --- Multi-paragraph display block: $$ ... $$  or \[ ... \]
        if stripped in ("$$", r"\["):
            delim_open = stripped
            delim_close = "$$" if delim_open == "$$" else r"\]"
            j = i + 1
            lines: List[str] = []
            while j < len(paragraphs) and normalize_math_text(_paragraph_plain_text(paragraphs[j])).strip() != delim_close:
                lines.append(_paragraph_plain_text(paragraphs[j]))
                j += 1

            if j < len(paragraphs) and normalize_math_text(_paragraph_plain_text(paragraphs[j])).strip() == delim_close:
                latex_block = "\n".join(lines).strip()
                latex_block = normalize_math_text(latex_block)

                _clear_paragraph_content(p)
                parts = split_long_latex_equation(latex_block) if latex_block else []
                if not parts:
                    for k in range(j, i - 1, -1):
                        _remove_paragraph(paragraphs[k])
                        del paragraphs[k]
                    continue

                for child in _latex_to_omml_children(parts[0]):
                    p._p.append(child)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                last_p = p
                for part in parts[1:]:
                    new_p = _insert_paragraph_after(last_p)
                    for child in _latex_to_omml_children(part):
                        new_p._p.append(child)
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    insert_pos = paragraphs.index(last_p) + 1
                    paragraphs.insert(insert_pos, new_p)
                    last_p = new_p

                for k in range(j, i, -1):
                    _remove_paragraph(paragraphs[k])
                    del paragraphs[k]

                i += 1
                continue

        # Single-paragraph aligned/matrix, then inline/delimited
        _process_inline_math_in_paragraph(p)
        i += 1


    i = 0
    n = len(paragraphs)
    while i < n:
        p = paragraphs[i]
        t = _paragraph_plain_text(p).strip()

        if t in ("$$", r"\["):
            delim_open = t
            delim_close = "$$" if delim_open == "$$" else r"\]"
            j = i + 1
            lines: List[str] = []
            while j < n and _paragraph_plain_text(paragraphs[j]).strip() != delim_close:
                lines.append(_paragraph_plain_text(paragraphs[j]))
                j += 1

            if j < n and _paragraph_plain_text(paragraphs[j]).strip() == delim_close:
                latex_block = "\n".join(lines).strip()
                latex_block = normalize_math_text(latex_block)

                # Replace opening delimiter paragraph with the equation
                _clear_paragraph_content(p)
                parts = split_long_latex_equation(latex_block) if latex_block else []
                if not parts:
                    # If empty, just delete the whole block
                    for k in range(j, i - 1, -1):
                        _remove_paragraph(paragraphs[k])
                    i = j + 1
                    continue

                # First part in the original paragraph
                for child in _latex_to_omml_children(parts[0]):
                    p._p.append(child)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Additional parts as extra paragraphs (if any)
                last_p = p
                for part in parts[1:]:
                    new_p = _insert_paragraph_after(last_p)
                    for child in _latex_to_omml_children(part):
                        new_p._p.append(child)
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p = new_p

                # Remove inner lines and closing delimiter paragraph
                for k in range(j, i, -1):
                    _remove_paragraph(paragraphs[k])

                i = j + 1
                continue

        # Also process single-paragraph equations after block check
        _process_inline_math_in_paragraph(p)
        i += 1
    # Add spacing between consecutive equation paragraphs for readability
    _apply_equation_spacing(paragraphs)



def _iter_block_items(parent) -> Any:
    """Yield Paragraph and Table objects in document order for a parent container.

    We intentionally avoid relying on `isinstance(parent, Document)` because `docx.Document`
    is a factory function, not a class. Instead, we detect the underlying XML container.
    """
    # Document body: parent._element.body exists
    if hasattr(parent, "_element") and hasattr(parent._element, "body"):
        parent_elm = parent._element.body
    else:
        # Header/Footer: parent._element is the container element
        parent_elm = getattr(parent, "_element", None)
        # Table cell: parent._tc is the container element
        if parent_elm is None:
            parent_elm = getattr(parent, "_tc", None)
        if parent_elm is None:
            raise TypeError(f"Unsupported container type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _collect_paragraphs_in_order(parent) -> List[Paragraph]:
    """Collect all paragraphs under parent (including inside tables) in document order."""
    out: List[Paragraph] = []
    for item in _iter_block_items(parent):
        if isinstance(item, Paragraph):
            out.append(item)
        else:
            # Table: descend into its cells
            for row in item.rows:
                for cell in row.cells:
                    out.extend(_collect_paragraphs_in_order(cell))
    return out


def _convert_docx_in_place(file_bytes: bytes) -> Document:
    doc = Document(io.BytesIO(file_bytes))

    # Body
    body_pars = _collect_paragraphs_in_order(doc)
    _process_display_blocks_in_paragraphs(body_pars)

    # Headers/footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                pars = _collect_paragraphs_in_order(hf)
                _process_display_blocks_in_paragraphs(pars)
            except Exception:
                # Some environments may have restricted access; continue.
                continue

        # First-page / even-page headers/footers if present
        for hf_attr in ("first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
            hf = getattr(section, hf_attr, None)
            if hf is None:
                continue
            try:
                pars = _collect_paragraphs_in_order(hf)
                _process_display_blocks_in_paragraphs(pars)
            except Exception:
                continue

    return doc


@app.get("/convert")
async def convert_get() -> RedirectResponse:
    """Redirect GET requests to the homepage.

    Googlebot may discover /convert via the form action in the UI. The converter endpoint is POST-only,
    so serving a redirect here avoids 4xx/405 reports in Search Console.
    """
    return RedirectResponse(url="/", status_code=301)


def _download_filename(lang: str) -> str:
    return "equations-to-word.docx" if lang == "en" else "ecuaciones-a-word.docx"


@app.post("/convert")
async def convert(file: UploadFile = File(...), lang: str = Form("es")) -> StreamingResponse:
    lang = "en" if (lang or "").strip().lower() == "en" else "es"
    filename = (file.filename or "").lower().strip()
    if not filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")

    if filename.endswith(".docx"):
        # IMPORTANT: preserve formatting and modify ONLY equations.
        try:
            out_doc = _convert_docx_in_place(content)
        except (BadZipFile, PackageNotFoundError):
            raise HTTPException(
                status_code=400,
                detail="Invalid .docx file. Upload a valid Word document.",
            ) from None
        except Exception:
            logger.exception("Unexpected error while converting .docx")
            raise HTTPException(
                status_code=500,
                detail="Internal error while converting the document.",
            ) from None
    elif filename.endswith(".txt"):
        paragraph_texts = _extract_text_lines_from_txt(content)
        cleaned = prettify_paragraphs(paragraph_texts)
        out_doc = build_document_from_paragraphs(cleaned)
    else:
        raise HTTPException(
            status_code=400, detail="Unsupported file type. Use .docx or .txt"
        )

    out = io.BytesIO()
    out_doc.save(out)
    out.seek(0)

    download_name = _download_filename(lang)
    headers = {
        "Content-Disposition": f'attachment; filename="{download_name}"; filename*=UTF-8\'\'{quote(download_name)}',
        "Cache-Control": "no-store",
    }

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

# ================================================================
# 404 handler (HTML en vez de JSON)
# ================================================================
@app.exception_handler(StarletteHTTPException)
async def custom_http_exception_handler(request: Request, exc: StarletteHTTPException):
    if exc.status_code == 404:
        safe_path = html_escape(request.url.path)
        is_en = str(request.url.path).startswith("/en")
        title = "404 | Page not found" if is_en else "404 | Página no encontrada"
        h1 = "404 · Page not found" if is_en else "404 · Página no encontrada"
        message = (
            f'The URL <code>{safe_path}</code> does not exist.'
            if is_en
            else f'No existe la URL <code>{safe_path}</code>.'
        )
        links = (
            '<a href="/en">Go to the converter</a> · <a href="/en/blog">Read the guides</a>'
            if is_en
            else '<a href="/">Ir al conversor</a> · <a href="/blog">Ver las guías</a>'
        )
        html = f"""<!doctype html>
<html lang="{('en' if is_en else 'es')}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="robots" content="noindex, nofollow">
  <title>{title}</title>
</head>
<body>
  <main style="max-width:760px;margin:48px auto;padding:0 20px;font:16px/1.6 system-ui,sans-serif;color:#0f172a;">
    <h1>{h1}</h1>
    <p>{message}</p>
    <p>{links}</p>
  </main>
</body>
</html>
"""
        return HTMLResponse(html, status_code=404, headers=_noindex_headers())

    return PlainTextResponse(str(exc.detail), status_code=exc.status_code)

